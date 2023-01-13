# -*- coding: utf-8 -*-kindOfInvoice
"""
Gauranga 4.1 - Python 3.11.1
Welcome to Gauranga ERP free software.
Gauranga comes with ABSOLUTELY NO WARRANTY.
2022 - Michael Tschoepe - mich.tscho@gmail.com


"""

# Create Windows .exe file
#pyinstaller -F Gauranga_4_1_ERP_free_software.pyw


# Gauranga ERP free software
# Software für Buchverlag im Bereich Warenwirtschaft
 
# Die Standalone Software wurde 2022 für einen Anwender entwickelt. Die Eingabe erfolgt über die grafische Benutzeroberfläche und Excel Importdateien. Ausgegeben wird über die grafische Benutzeroberfläche, Excel- und Worddateien.

# Folgende Aufgaben werden von der Software erfüllt:
# •	Bücher-, Verteiler-, und Saldenverwaltung
# •	Bestellungsannahme
# •	Rechnungs-, Lieferschein-, Gutschrift- und Berichterstellung

# Systemanforderungen:
# •	Windows 10
# •	Excel 365
# •	Word 365
# •	Python 3.11.1
# o	PySimpleGUI
# o	sqlite3
# o	as
# o	pandas
# o	asyncio.windows_events
# o	datetime
# o	sys
# o	os
# o	os.path
# o	pandas
# o	pathlib
# o	shutil
# o	random
# o	openpyxl
# o	time
# o	docx
# o	re
# o	math



from ast import expr_context
from asyncio.windows_events import NULL
import sqlite3
# from datetime import datetime
from datetime import date
from datetime import datetime
import PySimpleGUI as sg
import sys, os
import os.path
import pandas as pd
from pathlib import Path
import shutil
import random
import openpyxl
import time
import datetime
#pip3 install python-docx 
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import math

donwload = "no"

importAnz = "no"

ersteBestellZeileOK = "no"

zweiteBestellZeileOK = "no"

anzahlBüCheck = ""

büchLieferCheck = ""

bookToLieferDict = {}


sumBi = 0

def open_db():
    global connection
    global cursor
    connection = sqlite3.connect("sankirtan.db")
    cursor = connection.cursor()


def db_backup():
    dtNw = datetime.datetime.now()
    dtNw = dtNw.strftime("%d_%m_%Y_%H_%M_%S")
    dbBackupName = dtNw + '_sankirtan.db'
    dbBackup = os.path.join(dbBackupDirPath, dbBackupName)
    shutil.copyfile('sankirtan.db', dbBackup)
    #print(dtNw)




def clear_Ditri():
    try:
        window1['outVerteiler_IDDistri'].update('')
        window1['Verteiler_ID'].update('')
        window1['outVerteiler_IDDistri'].update('')
        window1['Verteiler_ID'].update('')
        window1['outVerteiler_IDDistri'].update('')
        window1['Email'].update('')
        window1['outEmailDistri'].update('')
        window1['outLandDistri'].update('')
        window1['Land'].update('')
        window1['outNameDistri'].update('')
        window1['Name'].update('')
        window1['outStraßeDistri'].update('')
        window1['Straße'].update('')
        window1['outPostleitzahl_OrtDistri'].update('')
        window1['Postleitzahl_Ort'].update('')
        #window1['outOrtDistri'].update('')
        #window1['Ort'].update('')
        window1['outStraßeDistri'].update('')
        window1['Straße'].update('')
        #window1['outInfoDistri'].update('')
        window1['Telefonnummer'].update('')
        window1['outTelefonnummerDistri'].update('')
        window1['UST_ID'].update('')
        window1['outUST_ID'].update('')

    except:
        pass

def clear_Buch():
    try:
        #window1['anzahlBu'].update('')
        window1['outIDBuch'].update('')
        window1['ID'].update('')
        window1['outLanguageBuch'].update('')
        window1['language'].update('')
        window1['outwarehouseBuch'].update('')
        window1['warehouse'].update('')
        window1['name_of_item'].update('')
        window1['outNameOfItemBuch'].update('')
        window1['outtypeBuch'].update('')
        window1['type'].update('')
        window1['outBBT_priceBuch'].update('')
        window1['BBT_price'].update('')
        window1['outSPBuch'].update('')
        window1['SP'].update('')
        window1['outP1Buch'].update('')
        window1['P1'].update('')
        window1['outP2Buch'].update('')
        window1['P2'].update('')
        window1['outP3_40Buch'].update('')
        window1['P3_40'].update('')
        window1['outP3_30Buch'].update('')
        window1['P3_30'].update('')
        window1['outEnd_PBuch'].update('')
        window1['End_P'].update('')
        window1['outInventory_Name'].update('')
        window1['Inventory_Name'].update('')
        window1['outBestandBuch'].update('')
        window1['Bestand'].update('')
        window1['outLa_Preis'].update('')
        window1['La_Preis'].update('')
        window1['outName_Preis'].update('')
        window1['Preis_Liste'].update('')


    except:
        pass


def delConTmp():
    for f in os.listdir(repDirPath):
        try:
            os.remove(os.path.join(repDirPath, f))
        except:
            pass


scriptPath = os.path.abspath(os.path.dirname(sys.argv[0]))
dbName = "sankirtan.db"
#print(scriptPath)

#  Go to script foulder.

os.chdir(scriptPath)


# Create Foulder if the doesent exst

repDir = "tmp"
Path(repDir).mkdir(parents=True, exist_ok=True)
repDirPath = os.path.join(scriptPath, repDir)

dbBackupDir = "sankirtan_db_backup"
Path(dbBackupDir).mkdir(parents=True, exist_ok=True)
dbBackupDirPath = os.path.join(scriptPath, dbBackupDir)

rechLiDir = "rechnungen_lieferscheine_gutschriften"
Path(rechLiDir).mkdir(parents=True, exist_ok=True)
rechDirPath = os.path.join(scriptPath, rechLiDir)




# Check if Database is createted.

if not os.path.isfile(dbName):

    layout0 = [
    [sg.Text("")],
    [sg.Text('Please copy the actual "' + dbName + '" in the foulder "' + scriptPath + '" and restart the programe.' )]
    ]

    window0 = sg.Window("No database", layout0)
    event, values = window0.read()
    time.sleep(5)
    sys.exit(0)


#connection = sqlite3.connect("sankirtan.db")
#cursor = connection.cursor()


#################################################
########## this table is in use 1.5.2022 ########
#################################################
#cursor.execute("CREATE TABLE distributer (Verteiler_ID TEXT PRIMARY KEY, Name TEXT, Straße TEXT, Postleitzahl_Ort TEXT, Land TEXT, Email TEXT, Telefonnummer TEXT)")
#cursor.execute("ALTER TABLE distributer ADD COLUMN UST_ID TEXT")


#################################################
########## this table is in use 1.5.2022 ########
#################################################
#cursor.execute("DROP TABLE bücherbewegung;")
#cursor.execute("CREATE TABLE bücherbewegung (bbID TEXT PRIMARY KEY, book TEXT, date TEXT, amount INTEGER, Entnahme_Zuführung TEXT, Vorgangsnummer TEXT, Verteiler TEXT, name_of_item #TEXT, amount_old TEXT)")




#################################################
######### this table is in use 26.05.2022 #######
#################################################
#cursor.execute("DROP TABLE vorgang;")
#cursor.execute("CREATE TABLE vorgang (vorgangsnummer TEXT PRIMARY KEY, buchdatensatz TEXT, verteilerdatensatz TEXT, summe INTEGER)")


#cursor.execute("DROP TABLE books;")date(Datum)

# cursor.execute("CREATE TABLE dsa_kontobewegung (number INTEGER PRIMARY KEY, Datum TEXT, Eingang REAL, Ausgang REAL)")






#cursor.execute("DROP TABLE dsa_kontostand;")
#cursor.execute("CREATE TABLE dsa_kontostand (DAS_Kontostand REAL PRIMARY KEY)")

#cursor.execute("CREATE TABLE kunden_saldo (Verteiler_ID TEXT PRIMARY KEY, kunden_saldo REAL)")

#################################################
########## this table is in use 23.8.2022 ########
#################################################
#cursor.execute("DROP TABLE kundenbewegung;")

#cursor.execute("CREATE TABLE kundenbewegung (vbID TEXT PRIMARY KEY, Verteiler_ID TEXT, Name TEXT, Vorgang TEXT, Vorgangsnummer TEXT, Datum TEXT, Datum_Zahlungsziel TEXT, Tage_Zahlungsziel INTEGER, Tage_Ablauf_Zahlungsziel INTEGER, Betrag REAL, Forderung TEXT, Saldo REAL)")


#
# cursor.execute("CREATE TABLE kunden_konto (Verteiler_ID TEXT PRIMARY KEY, Kontostand REAL)")




#cursor.execute("DROP TABLE forderung;")

#cursor.execute("CREATE TABLE forderung (uzID TEXT PRIMARY KEY, Verteiler_ID TEXT, Name TEXT, Vorgang TEXT, Vorgangsnummer TEXT, Datum TEXT, Datum_Zahlungsziel TEXT, Tage_Zahlungsziel INTEGER, Tage_Ablauf_Zahlungsziel INTEGER, Betrag REAL, Forderung TEXT, Saldo REAL)")


#
# cursor.execute("CREATE TABLE kunden_konto (Verteiler_ID TEXT PRIMARY KEY, Kontostand REAL)")




#cursor.execute("DROP TABLE books;")
#################################################
########## this table is in use 1.5.2022 ########
#################################################
#cursor.execute("CREATE TABLE books (ID Text PRIMARY KEY, name_of_item TEXT, Info TEXT, type TEXT, language TEXT, warehouse TEXT, quantity INTEGER, Info_2 TEXT, BBT_price REAL, SP REAL, P1 REAL, P2 REAL, P3_40 REAL, P3_30 REAL, End_P REAL, total_BBT_price_quantity REAL, Info_3 TEXT)")

# Für den Infentorry Report mussten noch mehr Spalten hinzugefügt werden:

# cursor.execute("ALTER TABLE books ADD COLUMN name TEXT")

# cursor.execute("ALTER TABLE books ADD COLUMN number INTEGER")

# cursor.execute("ALTER TABLE books ADD COLUMN total_BBT_price_number REAL")

#cursor.execute("ALTER TABLE books ADD COLUMN name_of_article TEXT")

#cursor.execute("ALTER TABLE books ADD COLUMN warehouses TEXT")


# Das ist die ganze DB mit allen Feldnern die gebraucht werden
#cursor.execute("CREATE TABLE books (ID Text PRIMARY KEY, name_of_item TEXT, type TEXT, language TEXT, warehouse TEXT, quantity INTEGER, BBT_price REAL, SP REAL, P1 REAL, P2 REAL, P3_40 REAL, P3_30 REAL, End_P REAL, total_BBT_price_quantity REAL, TEXT, Info_3 TEXT, name TEXT, number INTEGER, total_BBT_price_number REAL, Info_2 TEXT, name_of_article TEXT, warehouses TEXT)")

################################################
################# ENDE books ###################
################################################

open_db()

#     #NOT NULL PRIMARY KEY
#     #show tables
cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
#print(cursor.fetchall())






# Wenn aus kundenbewegung gelöscht wird stimmt das saldo nicht mehr
# def del_kndn_büchr_bwgng():
#     days = datetime.timedelta(1825)
#     today = date.today()
#     delDay = today - days
#     open_db()
#     cursor.execute('DELETE FROM kundenbewegung WHERE Datum <= ?', (delDay,))
#     cursor.execute('DELETE FROM bücherbewegung WHERE date <= ?', (delDay,))
#     connection.commit()
#     connection.close()


#Create the window.

#sg.ChangeLookAndFeel('GreenTan')
sg.theme('SandyBeach')
#sg.theme('TealMono')
#sg.theme('Light Blue 3')
direct = "menu"
#sg.theme_previewer()





while direct == "menu":

    delConTmp()



    open_db()


    content = [distributer[0] for distributer in cursor.execute("SELECT Verteiler_ID FROM distributer")]
    #print(content)
    #print('content')
    for kunde in content:

        #print(kunde)

        qtyBe = [qtyBe[0] for qtyBe in cursor.execute("SELECT Betrag FROM kundenbewegung WHERE Verteiler_ID=? ", (kunde,))]
        #print('SSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS ' + str(qtyBe))
        qtyBe = sum(qtyBe)
        qtyBe = round(qtyBe, 2)


        #print('SSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS ' + str(qtyBe))



        cursor.execute("UPDATE kundenbewegung SET Saldo=? WHERE Verteiler_ID=?", (qtyBe, kunde))


    content = [kundenbewegung[0] for kundenbewegung in cursor.execute("SELECT Vorgangsnummer FROM kundenbewegung")]

    for vorgangsnummer in content:


        vorganag1 = cursor.execute("SELECT Vorgang FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()
        #print(datumZahlungsziel)


        for t in vorganag1:
                for x in t:
                    vorganag1 = x

         #print(vorganag1)               


        if vorgangsnummer is not None and vorganag1 is not "G":
            #print('eeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeee')
            #print(vorgangsnummer)

            datumZahlungsziel = cursor.execute("SELECT Datum_Zahlungsziel FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()
            #print(datumZahlungsziel)


            for t in datumZahlungsziel:
                    for x in t:
                        datumZahlungsziel = x






            if datumZahlungsziel is not None: 

                #print(datumZahlungsziel)           

                currentDate =  datetime.datetime.now()

                #print(datumZahlungsziel)

                datumZahlungsziel = str(datumZahlungsziel)

                datumZahlungsziel = datumZahlungsziel.replace(' 00:00:00', '')


                #print(datumZahlungsziel)

                datumZahlungsziel = datetime.datetime.strptime(datumZahlungsziel, '%Y-%m-%d')



                taTageAblaufZahlungsziel = datumZahlungsziel - currentDate





                cursor.execute("UPDATE kundenbewegung SET Tage_Ablauf_Zahlungsziel=? WHERE Vorgangsnummer=?", (str(taTageAblaufZahlungsziel), vorgangsnummer))


                connection.commit()




            saldo = cursor.execute("SELECT Saldo FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()
            #print(saldo)

            for t in saldo:
                for x in t:
                    saldo = x

            saldo = float(saldo)

            saldo = round(saldo, 2)

            if saldo >= 0.00:

                #cursor.execute("UPDATE kundenbewegung SET Forderung=? WHERE Vorgangsnummer=?", ('bezahlt', vorgangsnummer))
                cursor.execute("UPDATE kundenbewegung SET Forderung=? WHERE Vorgangsnummer=? AND Vorgang=?", ('bezahlt', vorgangsnummer, 'L'))

                connection.commit()




            bezahlung = cursor.execute("SELECT Forderung FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()

            for t in bezahlung:
                for x in t:
                    bezahlung = x
            #print('uuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuub')
            #print(bezahlung)

            if '-' in str(taTageAblaufZahlungsziel):

                if not 'bezahlt' in str(bezahlung):

                    cursor.execute("UPDATE kundenbewegung SET Forderung=? WHERE Vorgangsnummer=?", ('überfällig', vorgangsnummer))

                    connection.commit()




    rows = cursor.execute("SELECT Name, Vorgang, Vorgangsnummer, Datum, Datum_Zahlungsziel, Tage_Zahlungsziel, Tage_Ablauf_Zahlungsziel, Betrag, Forderung FROM kundenbewegung WHERE Forderung=? ORDER BY Name, date(Datum_Zahlungsziel) ASC", ('überfällig',)).fetchall()

    countRows = len(rows)

    connection.close()



    layout0 = [
    [sg.Text("")],
    [sg.Button("Versand", size=(9,1)), sg.Button("Berichte", size=(9,1)), sg.Button("Laksmi", size=(9,1)), sg.Button("Bücher", size=(9,1)), sg.Button("Verteiler", size=(9,1)), sg.Button("Backup", size=(9,1)), sg.Button('Exit', size=(9,1))],
    [sg.Text(key="menuText")],
]

    window0 = sg.Window("Gauranga 4.1", layout0, finalize=True)

    #print(countRows)

    if countRows == 1:

        window0['menuText'].update('Eine Forderunge ist überfällig.', text_color='Red')

    elif countRows >= 1:

        window0['menuText'].update('Es sind ' + str(countRows) + ' Forderungen überfällig.', text_color='Red')


    while direct == "menu":

        event, values = window0.read()




        if event == "Backup" or event == sg.WIN_CLOSED:
            db_backup()


        if event == "Bücher" or event == sg.WIN_CLOSED:

            direct = "add"
            window0.close()


        if event == "Berichte" or event == sg.WIN_CLOSED:

            direct = "bbtReports"
            window0.close()


        if event == "Laksmi" or event == sg.WIN_CLOSED:

            direct = "laksmi"
            window0.close()


        if event == "Verteiler" or event == sg.WIN_CLOSED:

            direct = "Kunden"
            window0.close()

        if event == "Versand" or event == sg.WIN_CLOSED:

            delConTmp()
            direct = "versand"
            window0.close()


        if event == "Exit" or event == sg.WIN_CLOSED:
            db_backup()
            direct = "exit"
            window0.close()
            break



    while direct == "bbtReports":

        delConTmp()

        dtNew = date.today()
        dtNew = dtNew.strftime('%Y-%m-%d')

        days = datetime.timedelta(365)
        today = date.today()
        dtStart = today - days

        layout111 = [
            [sg.Text("")],
            #[sg.Text("", size=(79,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
            [sg.Button("Verteilerliste", size=(15,1)), sg.Text("", size=(55,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
            [sg.Text("")],
            [sg.Button("Inventory", size=(15,1)), sg.Button("Preisliste Excel", size=(15,1)), sg.Button("Alles", size=(15,1)), sg.Button("Kalpa Taru", size=(15,1))],
            [sg.Text("")],
            [sg.CalendarButton(' Von Datum', target='-vBCAL-', size=(15, 1), pad=None, key='_CALENDAR_', format=('%Y-%m-%d')), sg.In(dtStart, key='-vBCAL-', readonly=True, size=(17, 1)), sg.CalendarButton(' Bis Datum', target='-bBCAL-', size=(15, 1), pad=None, key='_CALENDAR_', format=('%Y-%m-%d')), sg.In(dtNew, key='-bBCAL-', readonly=True, size=(18, 1)), sg.Button("Bücherbewegung", size=(15,1))],
            [sg.Text("")],
            [sg.Button("Zahlungsverzug", size=(15,1)), sg.Button("Alle Salden", size=(15,1))],
            [sg.Text("")],
            [sg.CalendarButton(' Von Datum', target='-vVCAL-', size=(15, 1), pad=None, key='_CALENDAR_', format=('%Y-%m-%d')), sg.In(dtStart, key='-vVCAL-', readonly=True, size=(17, 1)), sg.CalendarButton(' Bis Datum', target='-bVCAL-', size=(15, 1), pad=None, key='_CALENDAR_', format=('%Y-%m-%d')), sg.In(dtNew, key='-bVCAL-', readonly=True, size=(18, 1)), sg.Button("Vorgänge u. Salden", size=(15,1))],
            [sg.Text("", size=(66,1))],
            [sg.Text(key="outBericht")]

        ]


        window1 = sg.Window("Erstellung von Berichten", layout111)

        # Datensatz hinzüfgen

        while direct == "bbtReports":

            event, values = window1.read()


            if event == 'Verteilerliste' or event == sg.WIN_CLOSED:


                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Verteilerliste_" + dtNw + "_" + str(random_number) + ".xlsx"


                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'



                title = "Verteilerliste " + dtNw
                koExc = os.path.join(repDirPath, koExc)

                window1['outBericht'].update(asbKoExc)

                open_db()

                rows = cursor.execute("SELECT * FROM distributer ORDER BY Name").fetchall()
                colNames = cursor.execute("SELECT * FROM distributer")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows

                df = pd.DataFrame(tableBooks)




                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='Verteilerliste', index=False)

                #workbook  = writer.book
                worksheet = writer.sheets['Verteilerliste']


                #worksheet.write(4, 10, title)


                writer.save()



                wb = openpyxl.load_workbook(koExc)
                sheet = wb['Verteilerliste']

                sheet.delete_rows(1)
                sheet.insert_rows(idx=0, amount=3)
                sheet.cell(row=2, column=1).value = title
                wb.save(koExc)





                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')

                connection.close()




            if event == 'Vorgänge u. Salden' or event == sg.WIN_CLOSED:



                vonVerteilerCAL = values['-vVCAL-']
                bisVerteilerCAL = values['-bVCAL-']



                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Verteilerbewegung_" + dtNw + "_" + str(random_number) + ".xlsx"

                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'

                title = "Verteilerbewegung " + dtNw
                koExc = os.path.join(repDirPath, koExc)
                window1['outBericht'].update(asbKoExc)

                open_db()


                content = [distributer[0] for distributer in cursor.execute("SELECT Verteiler_ID FROM distributer")]
                #print(content)
                #print('content')
                for kunde in content:

                    #print(kunde)

                    qtyBe = [qtyBe[0] for qtyBe in cursor.execute("SELECT Betrag FROM kundenbewegung WHERE Verteiler_ID=? ", (kunde,))]
                    #print('SSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS ' + str(qtyBe))
                    qtyBe = sum(qtyBe)
                    qtyBe = round(qtyBe, 2)


                    #print('SSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS ' + str(qtyBe))



                    cursor.execute("UPDATE kundenbewegung SET Saldo=? WHERE Verteiler_ID=?", (qtyBe, kunde))


                content = [kundenbewegung[0] for kundenbewegung in cursor.execute("SELECT Vorgangsnummer FROM kundenbewegung")]

                for vorgangsnummer in content:


                    vorganag1 = cursor.execute("SELECT Vorgang FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()
                    #print(datumZahlungsziel)


                    for t in vorganag1:
                            for x in t:
                                vorganag1 = x



                    if vorgangsnummer is not None and vorganag1 is not "G":
                        #print('eeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeee')
                        #print(vorgangsnummer)

                        datumZahlungsziel = cursor.execute("SELECT Datum_Zahlungsziel FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()
                        #print(datumZahlungsziel)

                        for t in datumZahlungsziel:
                                for x in t:
                                    datumZahlungsziel = x



                        if datumZahlungsziel is not None:

                            currentDate =  datetime.datetime.now()

                            #print(datumZahlungsziel)

                            datumZahlungsziel = str(datumZahlungsziel)

                            datumZahlungsziel = datumZahlungsziel.replace(' 00:00:00', '')


                            #print(datumZahlungsziel)

                            datumZahlungsziel = datetime.datetime.strptime(datumZahlungsziel, '%Y-%m-%d')



                            taTageAblaufZahlungsziel = datumZahlungsziel - currentDate


                            cursor.execute("UPDATE kundenbewegung SET Tage_Ablauf_Zahlungsziel=? WHERE Vorgangsnummer=?", (str(taTageAblaufZahlungsziel), vorgangsnummer))
                            connection.commit()

                        saldo = cursor.execute("SELECT Saldo FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()
                        #print(saldo)

                        for t in saldo:
                            for x in t:
                                saldo = x

                        saldo = float(saldo)

                        saldo = round(saldo, 2)

                        if saldo >= 0.00:

                            cursor.execute("UPDATE kundenbewegung SET Forderung=? WHERE Vorgangsnummer=? AND Vorgang=?", ('bezahlt', vorgangsnummer, 'L'))


                            connection.commit()


                        bezahlung = cursor.execute("SELECT Forderung FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()

                        for t in bezahlung:
                            for x in t:
                                bezahlung = x
                        #print('uuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuub')
                        #print(bezahlung)
                        if '-' in str(taTageAblaufZahlungsziel):

                            if not 'bezahlt' in str(bezahlung):

                                cursor.execute("UPDATE kundenbewegung SET Forderung=? WHERE Vorgangsnummer=?", ('überfällig', vorgangsnummer))

                                connection.commit()



                rows = cursor.execute("SELECT Name, Vorgang, Vorgangsnummer, Datum, Datum_Zahlungsziel, Tage_Zahlungsziel, Tage_Ablauf_Zahlungsziel, Betrag, Forderung, Saldo FROM kundenbewegung WHERE Datum>= ? AND Datum<= ? ORDER BY Name, date(Datum) ASC", (vonVerteilerCAL, bisVerteilerCAL)).fetchall()




                colNames = cursor.execute("SELECT Name, Vorgang, Vorgangsnummer, Datum, Datum_Zahlungsziel, Tage_Zahlungsziel, Tage_Ablauf_Zahlungsziel, Betrag, Forderung, Saldo FROM kundenbewegung")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows

                #print(colNames)

                df = pd.DataFrame(tableBooks)

                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='Verteilerbewegung', index=False)

                worksheet = writer.sheets['Verteilerbewegung']
                writer.save()

                wb = openpyxl.load_workbook(koExc)
                sheet = wb['Verteilerbewegung']

                sheet.delete_rows(1)
                sheet.insert_rows(idx=0, amount=3)
                sheet.cell(row=2, column=1).value = title
                wb.save(koExc)

                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')

                connection.close()





            if event == 'Alle Salden' or event == sg.WIN_CLOSED:



                vonVerteilerCAL = values['-vVCAL-']
                bisVerteilerCAL = values['-bVCAL-']



                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Alle_Salden_" + dtNw + "_" + str(random_number) + ".xlsx"

                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'

                title = "Alle Salden " + dtNw
                koExc = os.path.join(repDirPath, koExc)
                window1['outBericht'].update(asbKoExc)

                open_db()


                content = [distributer[0] for distributer in cursor.execute("SELECT Verteiler_ID FROM distributer")]
                #print(content)
                #print('content')
                for kunde in content:

                    #print(kunde)

                    qtyBe = [qtyBe[0] for qtyBe in cursor.execute("SELECT Betrag FROM kundenbewegung WHERE Verteiler_ID=? ", (kunde,))]
                    #print('SSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS ' + str(qtyBe))
                    qtyBe = sum(qtyBe)
                    qtyBe = round(qtyBe, 2)


                    #print('SSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS ' + str(qtyBe))



                    cursor.execute("UPDATE kundenbewegung SET Saldo=? WHERE Verteiler_ID=?", (qtyBe, kunde))


                rows = cursor.execute("SELECT DISTINCT Name, Saldo FROM kundenbewegung").fetchall()


                colNames = cursor.execute("SELECT Name, Saldo FROM kundenbewegung")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows

                #print(colNames)

                df = pd.DataFrame(tableBooks)

                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='Alle Salden', index=False)

                worksheet = writer.sheets['Alle Salden']
                writer.save()

                wb = openpyxl.load_workbook(koExc)
                sheet = wb['Alle Salden']

                sheet.delete_rows(1)
                sheet.insert_rows(idx=0, amount=3)
                sheet.cell(row=2, column=1).value = title
                wb.save(koExc)

                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')

                connection.close()






            if event == 'Zahlungsverzug' or event == sg.WIN_CLOSED:







                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Forderungen_" + dtNw + "_" + str(random_number) + ".xlsx"

                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'

                title = "Forderung " + dtNw
                koExc = os.path.join(repDirPath, koExc)
                window1['outBericht'].update(asbKoExc)



                open_db()


                content = [distributer[0] for distributer in cursor.execute("SELECT Verteiler_ID FROM distributer")]
                #print(content)
                #print('content')
                for kunde in content:

                    #print(kunde)

                    qtyBe = [qtyBe[0] for qtyBe in cursor.execute("SELECT Betrag FROM kundenbewegung WHERE Verteiler_ID=? ", (kunde,))]
                    #print('SSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS ' + str(qtyBe))
                    qtyBe = sum(qtyBe)
                    qtyBe = round(qtyBe, 2)


                    #print('SSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS ' + str(qtyBe))



                    cursor.execute("UPDATE kundenbewegung SET Saldo=? WHERE Verteiler_ID=?", (qtyBe, kunde))


                content = [kundenbewegung[0] for kundenbewegung in cursor.execute("SELECT Vorgangsnummer FROM kundenbewegung")]

                for vorgangsnummer in content:


                    
                    vorganag1 = cursor.execute("SELECT Vorgang FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()
                    #print(datumZahlungsziel)


                    for t in vorganag1:
                            for x in t:
                                vorganag1 = x



                    if vorgangsnummer is not None and vorganag1 is not "G":

                        try:
                            #print('eeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeee')
                            #print(vorgangsnummer)

                                datumZahlungsziel = cursor.execute("SELECT Datum_Zahlungsziel FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()
                                #print(datumZahlungsziel)

                                for t in datumZahlungsziel:
                                        for x in t:
                                            datumZahlungsziel = x

                                currentDate =  datetime.datetime.now()

                                #print(datumZahlungsziel)

                                datumZahlungsziel = str(datumZahlungsziel)

                                datumZahlungsziel = datumZahlungsziel.replace(' 00:00:00', '')


                                #print(datumZahlungsziel)

                                datumZahlungsziel = datetime.datetime.strptime(datumZahlungsziel, '%Y-%m-%d')



                                taTageAblaufZahlungsziel = datumZahlungsziel - currentDate





                                cursor.execute("UPDATE kundenbewegung SET Tage_Ablauf_Zahlungsziel=? WHERE Vorgangsnummer=?", (str(taTageAblaufZahlungsziel), vorgangsnummer))


                                connection.commit()

                        except:
                            pass
                    
                        saldo = cursor.execute("SELECT Saldo FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()
                        #print(saldo)

                        for t in saldo:
                            for x in t:
                                saldo = x

                        saldo = float(saldo)

                        saldo = round(saldo, 2)

                        if saldo >= 0.00:

                            #cursor.execute("UPDATE kundenbewegung SET Forderung=? WHERE Vorgangsnummer=?", ('bezahlt', vorgangsnummer))
                            cursor.execute("UPDATE kundenbewegung SET Forderung=? WHERE Vorgangsnummer=? AND Vorgang=?", ('bezahlt', vorgangsnummer, 'L'))

                            connection.commit()



                        bezahlung = cursor.execute("SELECT Forderung FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()

                        for t in bezahlung:
                            for x in t:
                                bezahlung = x
                        #print('uuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuub')
                        #print(bezahlung)

                        if '-' in str(taTageAblaufZahlungsziel):

                            if not 'bezahlt' in str(bezahlung):

                                cursor.execute("UPDATE kundenbewegung SET Forderung=? WHERE Vorgangsnummer=?", ('überfällig', vorgangsnummer))

                                connection.commit()




                rows = cursor.execute("SELECT Name, Vorgang, Vorgangsnummer, Datum, Datum_Zahlungsziel, Tage_Zahlungsziel, Tage_Ablauf_Zahlungsziel, Betrag, Forderung, Saldo FROM kundenbewegung WHERE Forderung=? ORDER BY Name, date(Datum_Zahlungsziel) ASC", ('überfällig',)).fetchall()




                colNames = cursor.execute("SELECT Name, Vorgang, Vorgangsnummer, Datum, Datum_Zahlungsziel, Tage_Zahlungsziel, Tage_Ablauf_Zahlungsziel, Betrag, Forderung, Saldo FROM kundenbewegung")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows

                #print(colNames)

                df = pd.DataFrame(tableBooks)

                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='Forderung', index=False)

                worksheet = writer.sheets['Forderung']
                writer.save()

                wb = openpyxl.load_workbook(koExc)
                sheet = wb['Forderung']

                sheet.delete_rows(1)
                sheet.insert_rows(idx=0, amount=3)
                sheet.cell(row=2, column=1).value = title
                wb.save(koExc)

                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')

                connection.close()










            if event == 'Alles' or event == sg.WIN_CLOSED:

                open_db()

                qtySm = [qtySm[0] for qtySm in cursor.execute("SELECT quantity FROM books")]
                qtySm = sum(qtySm)
                qtySm = round(qtySm, 3)


                bbtSm = [bbtSm[0] for bbtSm in cursor.execute("SELECT BBT_price FROM books")]
                bbtSm = sum(bbtSm)
                bbtSm = round(bbtSm, 3)


                tBbtSm = [tBbtSm[0] for tBbtSm in cursor.execute("SELECT total_BBT_price_quantity FROM books")]
                tBbtSm = sum(tBbtSm)
                tBbtSm = round(tBbtSm, 3)



                dtNw = date.today()
                cursor.execute("INSERT INTO books (ID, warehouse, quantity, total_BBT_price_quantity) VALUES (?, ?, ?, ?)", ("total:", "z", qtySm, tBbtSm))

                random_number = random.randint(1, 10000)
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Kommisionsbestand_des_BBT_" + dtNw + "_" + str(random_number) + ".xlsx"

                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'

                title = "Kommisionsbestand des BBT " + dtNw
                koExc = os.path.join(repDirPath, koExc)

                window1['outBericht'].update(asbKoExc)

                rows = cursor.execute("SELECT * FROM books ORDER BY warehouse, language, name_of_item").fetchall()
                colNames = cursor.execute("SELECT * FROM books")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows

                df = pd.DataFrame(tableBooks)

                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='Kommisionsbestand', index=False)

                worksheet = writer.sheets['Kommisionsbestand']



                writer.save()



                wb = openpyxl.load_workbook(koExc)
                sheet = wb['Kommisionsbestand']

                sheet.delete_rows(1)
                sheet.insert_rows(idx=0, amount=3)
                sheet.cell(row=2, column=1).value = title
                wb.save(koExc)

                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')

                connection.close()


            if event == 'Bücherbewegung' or event == sg.WIN_CLOSED:

                vonBücherCAL = values['-vBCAL-']
                bisBücherCAL = values['-bBCAL-']



                open_db()

                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Bücherbewegung_" + dtNw + "_" + str(random_number) + ".xlsx"


                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'


                title = "Bücherbewegung " + dtNw
                koExc = os.path.join(repDirPath, koExc)

                window1['outBericht'].update(asbKoExc)

                rows = cursor.execute("SELECT date, name_of_item, Vorgangsnummer, Verteiler, amount_old, Entnahme_Zuführung, amount FROM Bücherbewegung WHERE date >= ? AND date <= ? ORDER BY date(date) ASC", (vonBücherCAL, bisBücherCAL,)).fetchall()


                colNames = cursor.execute("SELECT date, name_of_item, Vorgangsnummer, Verteiler, amount_old, Entnahme_Zuführung, amount FROM Bücherbewegung")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows

                df = pd.DataFrame(tableBooks)

                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='Bücherbewegung', index=False)

                #workbook  = writer.book
                worksheet = writer.sheets['Bücherbewegung']


                #worksheet.write(4, 10, title)


                writer.save()



                wb = openpyxl.load_workbook(koExc)
                sheet = wb['Bücherbewegung']

                sheet.delete_rows(1)
                sheet.insert_rows(idx=0, amount=3)
                sheet.cell(row=2, column=1).value = title
                wb.save(koExc)





                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')

                connection.close()




            if event == "Exit" or event == sg.WIN_CLOSED:
                db_backup()
                direct = "exit"
                window1.close()




            if event == "Gauranga" or event == sg.WIN_CLOSED:
                direct = "menu"
                lineToChange = []
                window1.close()







            if event == 'Inventory' or event == sg.WIN_CLOSED:



                open_db()


                content = [distributer[0] for distributer in cursor.execute("SELECT name FROM books")]
       
                for name in content:



                    if name != "0":

                       

                        content = [books[0] for books in cursor.execute("SELECT language FROM books")]
                       
                        for langu in content:

                            qtyBe = [qtyBe[0] for qtyBe in cursor.execute("SELECT quantity FROM books WHERE name=? AND language=?", (name, langu))]
                            qtyBe = sum(qtyBe)
                            





                            cursor.execute("UPDATE books SET number=? WHERE name=? AND language=?", (qtyBe, name, langu))

                            connection.commit()



                            open_db()

                            BBT_price = cursor.execute("SELECT BBT_price FROM books WHERE name =? AND language=? LIMIT 1", (name, langu)).fetchall()

                            if BBT_price:
                                
                                BBT_price = str(BBT_price)
                                
                               
                                

                           
                            
                                    
                                for char in ['(', ')', ',', '\'', ']', '[']:
                                    if char in BBT_price:

                                        BBT_price = BBT_price.replace(char, '')

                                BBT_price = float(BBT_price)

                                total_BBT_price_number = BBT_price * qtyBe


                                cursor.execute("UPDATE books SET total_BBT_price_number=?, Info_3=? WHERE name=? AND language=?", (total_BBT_price_number, name+langu, name, langu))
                                

                                connection.commit()

    


                open_db()

                qtySm = [qtySm[0] for qtySm in cursor.execute("SELECT number FROM books WHERE name NOT IN (?) GROUP BY Info_3", ("0",))]
                qtySm = sum(qtySm)
               

                tBbtSm = [tBbtSm[0] for tBbtSm in cursor.execute("SELECT total_BBT_price_number FROM books WHERE name NOT IN (?) GROUP BY Info_3", ("0",))]
                tBbtSm = sum(tBbtSm)
                tBbtSm = round(tBbtSm, 2)


                cursor.execute("INSERT INTO books (name, language, number, total_BBT_price_number) VALUES (?, ?, ?, ?)", ("total:", "z", qtySm, tBbtSm))

                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%m_%Y")
                koExc = "inventory_of_BBT-report_for_quarter_" + dtNw + "_" + str(random_number) + ".xlsx"


                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'


                title = "inventory of BBT-report for quarter " + dtNw + " - Musterbuchhandel"
                koExc = os.path.join(repDirPath, koExc)

                window1['outBericht'].update(asbKoExc)


                rows = cursor.execute("SELECT name, language, BBT_price, number, total_BBT_price_number from books WHERE name NOT IN (?) GROUP BY Info_3 ORDER BY language, name", ("0",)).fetchall()
                colNames = cursor.execute("SELECT name, language, BBT_price, number, total_BBT_price_number FROM books")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows

                df = pd.DataFrame(tableBooks)




                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='inventory', index=False)

                #workbook  = writer.book
                worksheet = writer.sheets['inventory']


                #worksheet.write(4, 10, title)


                writer.save()



                wb = openpyxl.load_workbook(koExc)
                sheet = wb['inventory']

                sheet.delete_rows(1)
                sheet.insert_rows(idx=0, amount=3)
                sheet.cell(row=2, column=1).value = title
                wb.save(koExc)





                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')

                connection.close()




            if event == 'Preisliste Word' or event == sg.WIN_CLOSED:

                open_db()

                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Preisliste_" + dtNw + "_" + str(random_number) + ".docx"

                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'

                title = "Preisliste " + dtNw
                koExc = os.path.join(repDirPath, koExc)

                window1['outBericht'].update(asbKoExc)



                rows = cursor.execute("SELECT ID, name_of_item, type, language, warehouse, SP, P1, P2, P3_30, End_P FROM books ORDER BY language, name_of_item").fetchall()
                colNames = cursor.execute("SELECT ID, name_of_item, type, language, warehouse, SP, P1, P2, P3_30, End_P FROM books")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows

                df = pd.DataFrame(tableBooks)

                ######################################


                document = Document()



                def set_column_width(column, width):
                    for cell in column.cells:
                        cell.width = width


                tabletop = document.add_table(rows=1, cols=2, style="Table Grid")


                set_column_width(tabletop.columns[0], Cm(16))
                set_column_width(tabletop.columns[1], Cm(17))

                heading_row = tabletop.rows[0].cells

                heading_row[0].text = "Musterbuchhandel	\n Musterbücher"

                heading_row[1].text = "Mustersadresse\n Tel: 0176 57985621 \n (Godruma/Frank-M. Walz)"


                paragraph = document.add_paragraph('')

                paragraph = document.add_paragraph()
                paragraph.add_run('Preisliste - Stand: ' + str(dtNw)).font.size = Pt(12)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(22)

   


                # tableRech = document.add_table(rows=2, cols=4, style="Table Grid")


                # table = document.add_table(rows=positionsB, cols=6, style="Table Grid")



                # set_column_width(table.columns[0], Cm(1))
                # set_column_width(table.columns[1], Cm(2))
                # set_column_width(table.columns[2], Cm(2))
                # set_column_width(table.columns[3], Cm(25))
                # set_column_width(table.columns[4], Cm(6))
                # set_column_width(table.columns[5], Cm(8))
                # set_column_width(table.columns[6], Cm(2))
                # set_column_width(table.columns[7], Cm(25))
                # set_column_width(table.columns[8], Cm(6))
                # set_column_width(table.columns[9], Cm(8))


                # heading_row = table.rows[0].cells

                # heading_row[0].text = "ID"
                # heading_row[1].text = "name_of_item"
                # heading_row[2].text = "type"
                # heading_row[3].text = "language"
                # heading_row[4].text = "warehouse."
                # heading_row[5].text = "SP."
                # heading_row[6].text = "P1"
                # heading_row[7].text = "P2"
                # heading_row[8].text = "P3_30"
                # heading_row[9].text = "End_P"


                #table = document.add_table(rows=positionsB, cols=6, style="Table Grid")

                table = document.add_table(rows=(df.shape[0]), cols=df.shape[1], style="Table Grid") # First row are table headers!

                # set_column_width(table.columns[0], Cm(1))
                # set_column_width(table.columns[1], Cm(2))
                # set_column_width(table.columns[2], Cm(2))
                # set_column_width(table.columns[3], Cm(25))
                # set_column_width(table.columns[4], Cm(6))
                # set_column_width(table.columns[5], Cm(8))

            
         
                

                #table.allow_autofit = True
                #table.autofit = True
                for i, column in enumerate(df) :
                    for row in range(df.shape[0]) :
                        table.cell(row, i).text = str(df[column][row])


                paragraph = document.add_paragraph('')

                paragraph.add_run('SP: Straßenmissionspreis - für Verteiler, die die Bücher direkt im Rahmen der Straßenmission verteilen und einen gültigen DSA-Ausweis haben und Tempel, deren Verteiler aktiv und regelmäßig an der Straßenmission teilnehmen (DSA-Ausweis)')
                paragraph.paragraph_format.space_before = Pt(33)
                paragraph.paragraph_format.space_after = Pt(3)

                paragraph = document.add_paragraph(str('P1: Für Tempel, die nicht SP zahlen'))
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)

                paragraph = document.add_paragraph('P2: Devoteebuchhandlungen, verschiedene Devotees, die nicht SP zahlen, Mindestabnahme 20 Bücher')
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)

                paragraph = document.add_paragraph('P3_30%: kommerzielle Buchhandlungen')
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)

                paragraph = document.add_paragraph('End_P: Empfohlener Endverkaufspreis für ganz Deutschland, dient auch zur Orientierung für Tempelshops')
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Preisliste_" + dtNw + "_" + str(random_number) + ".docx"

                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'

                title = "Preisliste " + dtNw
                koExc = os.path.join(repDirPath, koExc)

                window1['outBericht'].update(asbKoExc)




                koExc = os.path.join(repDirPath, koExc)

                
                document.save(koExc)



                absolutePath = Path(koExc).resolve()
                os.system(f'start WINWORD.EXE "{absolutePath}"')




                ####################################




                # writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                # df.to_excel(writer, sheet_name='Preisliste', index=False)

                # #workbook  = writer.book
                # worksheet = writer.sheets['Preisliste']


                # #worksheet.write(4, 10, title)


                # writer.save()



                # wb = openpyxl.load_workbook(koExc)
                # sheet = wb['Preisliste']

                # sheet.delete_rows(1)
                # sheet.insert_rows(idx=0, amount=3)
                # sheet.cell(row=2, column=1).value = title
                # wb.save(koExc)





                # absolutePath = Path(koExc).resolve()
                # os.system(f'start excel.exe "{absolutePath}"')

                # connection.close()





            if event == 'Preisliste Excel' or event == sg.WIN_CLOSED:


                

                open_db()


                content = [distributer[0] for distributer in cursor.execute("SELECT name_of_article FROM books")]
       
                for name_of_article in content:



                    if name_of_article != "0":

                       

                        content = [books[0] for books in cursor.execute("SELECT language FROM books")]
                       
                        for langu in content:


                            cursor.execute("UPDATE books SET Info_2=? WHERE name_of_article=? AND language=?", (name_of_article+langu, name_of_article, langu))
                            

                            connection.commit()





                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Preisliste_" + dtNw + "_" + str(random_number) + ".xlsx"
                                
                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'
                
                title = "Preisliste " + dtNw
                koExc = os.path.join(repDirPath, koExc)
                
                window1['outBericht'].update(asbKoExc)
                

                
                rows = cursor.execute("SELECT name_of_article, type, language, warehouses, SP, P1, P2, P3_30, End_P FROM books WHERE name_of_article NOT IN (?) AND warehouses NOT IN (?) GROUP BY Info_2 ORDER BY language, name_of_item", ("0", "0")).fetchall()



                

                colNames = cursor.execute("SELECT name_of_article, type, language, warehouses, SP, P1, P2, P3_30, End_P FROM books")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows
                
                df = pd.DataFrame(tableBooks)

 


                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='Preisliste', index=False)
                
                #workbook  = writer.book
                worksheet = writer.sheets['Preisliste']


                #worksheet.write(4, 10, title)

                     
                writer.save()



                wb = openpyxl.load_workbook(koExc)
                sheet = wb['Preisliste']
               
                sheet.delete_rows(1)
                sheet.insert_rows(idx=0, amount=3)
                sheet.cell(row=2, column=1).value = title
                wb.save(koExc)





                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')

                connection.close()





            if event == 'Wiesenena' or event == sg.WIN_CLOSED:

                open_db()

                qtySm = [qtySm[0] for qtySm in cursor.execute("SELECT quantity FROM books WHERE warehouse = 'Wiesenena'")]
                qtySm = sum(qtySm)


                tBbtSm = [tBbtSm[0] for tBbtSm in cursor.execute("SELECT total_BBT_price_quantity FROM books WHERE warehouse = 'Wiesenena'")]
                tBbtSm = sum(tBbtSm)

                cursor.execute("INSERT INTO books (ID, quantity, total_BBT_price_quantity) VALUES (?, ?, ?)", ("total:", qtySm, tBbtSm))



                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Kommisionsbestand_des_BBT_Wiesenena_" + dtNw + "_" + str(random_number) + ".xlsx"


                asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'


                title = "Kommisionsbestand des BBT Wiesenena " + dtNw
                koExc = os.path.join(repDirPath, koExc)

                window1['outBericht'].update(asbKoExc)



                rows = cursor.execute("SELECT ID, name_of_item, language, quantity, BBT_price, SP, total_BBT_price_quantity FROM books WHERE warehouse = 'Wiesenena' OR ID = 'total:'").fetchall()
                colNames = cursor.execute("SELECT ID, name_of_item, language, quantity, BBT_price, SP, total_BBT_price_quantity FROM books")
                colNames = [cn[0] for cn in colNames.description]
                colNames = tuple(colNames)
                rows.insert(0, colNames)
                tableBooks = rows

                df = pd.DataFrame(tableBooks)




                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='Wiesenena', index=False)

                #workbook  = writer.book
                worksheet = writer.sheets['Wiesenena']


                #worksheet.write(4, 10, title)


                writer.save()



                wb = openpyxl.load_workbook(koExc)
                sheet = wb['Wiesenena']

                sheet.delete_rows(1)
                sheet.insert_rows(idx=0, amount=3)
                sheet.cell(row=2, column=1).value = title
                wb.save(koExc)





                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')

                connection.close()

#cursor.execute("DELETE FROM books WHERE ID=? ", (ID,))
#where strftime('%Y', datecolumn) = '2010'

            if event == 'Lösche alle Einträge in Verteilerbewegung' or event == sg.WIN_CLOSED:
                #WHERE Datum between date('now', 'start of day','-2 days')
                           #and date('now', 'start of day', '+1 day')
                           #SELECT date('now','start of month','+1 month','-1 day');
                open_db()
                beginn = '2023-01-01'
                ende = "2021-01-01"
                # jahre anzeigen rows = cursor.execute("SELECT * FROM kundenbewegung where strftime('%Y', Datum) = '2020'").fetchall()
                # zeitraum anzeigen dazwischen anzeigen rows = cursor.execute('SELECT * FROM kundenbewegung WHERE Datum >= "2018-01-01" and Datum <= "2021-01-01"').fetchall()
                # before Zeitpunkt anzeigen lassen rows = cursor.execute('SELECT * FROM kundenbewegung WHERE Datum <= "2015-01-01"').fetchall()

#"UPDATE kundenbewegung SET Saldo=? WHERE Verteiler_ID=?", (qtyBe, kunde))

                        #rows = cursor.execute("SELECT * FROM books WHERE type = ? AND name_of_item = ?", (typeDddd, name_of_itemDddd,)).#fetchall()
                        #



                #import datetime



                rows = cursor.execute("SELECT * FROM kundenbewegung").fetchall()
                # zeitraum anzeigen dazwischen anzeigen rows = cursor.execute('SELECT * FROM kundenbewegung WHERE Datum >= "2018-01-01" and Datum <= "2021-01-01"').fetchall()
                #print(rows)
                window1['outBericht'].update(rows)

            if event == 'Lösche alle Einträge in Bücherbewegung' or event == sg.WIN_CLOSED:

                open_db()
                connection.execute('DELETE FROM bücherbewegung;',);
                connection.commit()
                connection.close()
                window1['outBericht'].update("Im System wurden alle Zeilen aus Bücherbewegung gelöscht. Die Berichte im Berichte Ordner wurden nicht gelöscht.")

            if event == 'Kalpa Taru' or event == sg.WIN_CLOSED:

                direct = "kalpa"
                window1.close()


        while direct == "kalpa":

            open_db()




            warehouseDdd = [books[0] for books in cursor.execute("SELECT warehouse FROM books")]
            warehouseDdd = list(set(warehouseDdd))
            warehouseDdd = sorted(warehouseDdd)


            languageDdd = [books[0] for books in cursor.execute("SELECT language FROM books")]
            languageDdd = list(set(languageDdd))
            languageDdd = sorted(languageDdd)

            typeDdd = [books[0] for books in cursor.execute("SELECT type FROM books")]
            typeDdd = list(set(typeDdd))
            typeDdd = sorted(typeDdd)

            name_of_itemDdd = [books[0] for books in cursor.execute("SELECT name_of_item FROM books")]
            name_of_itemDdd = list(set(name_of_itemDdd))
            name_of_itemDdd = sorted(name_of_itemDdd)


            typeDddd = ""
            name_of_itemDddd = ""
            warehouseDddd = ""
            languageDddd = ""



            layout108 = [
                [sg.Button('Clear', size=(9,1)), sg.Text('', size=(24, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text('')],
                [sg.Text('name_of_item', size=(36, 1)), sg.Text('type', size=(9, 1)), sg.Text('language', size=(10, 1)), sg.Text('warehouse', size=(9, 1))],
                [sg.Combo(name_of_itemDdd, size=(39, 1), readonly=True, key="name_of_itemDddd"), sg.Combo(typeDdd, size=(9, 1), readonly=True, key="typeDddd"), sg.Combo(languageDdd, size=(9, 1), readonly=True, key="languageDddd"), sg.Combo(warehouseDdd, size=(9, 1), readonly=True, key="warehouseDddd")],

                [sg.Button('Ok', size=(9,1)), sg.Text('Ein, mehrere oder alle Suchkritärien auswählen und dann auf OK klicken.')],
                [sg.Text(key="outKalpa")],
            ]

            window108 = sg.Window("Suchkriterien auswählen", layout108)

            while direct == "kalpa":

                event, values = window108.read()

                if event == "Clear" or event == sg.WIN_CLOSED:

                    window108['name_of_itemDddd'].update('')
                    window108['warehouseDddd'].update('')
                    window108['languageDddd'].update('')
                    window108['typeDddd'].update('')
                    window108['outKalpa'].update('')

                if event == "Ok" or event == sg.WIN_CLOSED:

                    typeDddd = values['typeDddd']
                    name_of_itemDddd = values['name_of_itemDddd']
                    warehouseDddd = values['warehouseDddd']
                    languageDddd = values['languageDddd']


                    if typeDddd != "" and name_of_itemDddd != "" and warehouseDddd != "" and languageDddd != "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"

                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'

                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE type = ? AND name_of_item = ? AND language = ? AND warehouse = ?", (typeDddd, name_of_itemDddd, languageDddd, warehouseDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""



                    if typeDddd != "" and name_of_itemDddd != "" and warehouseDddd != "" and languageDddd == "":




                        open_db()



                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"

                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'

                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)
                        rows = cursor.execute("SELECT * FROM books WHERE type = ? AND name_of_item = ? AND warehouse = ?", (typeDddd, name_of_itemDddd, warehouseDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()


                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""


                    if typeDddd != "" and name_of_itemDddd != "" and warehouseDddd == "" and languageDddd == "":




                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"



                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'


                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE type = ? AND name_of_item = ?", (typeDddd, name_of_itemDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""




                    if typeDddd != "" and name_of_itemDddd == "" and warehouseDddd == "" and languageDddd == "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"



                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'



                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE type = ?", (typeDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""





                    if typeDddd != "" and name_of_itemDddd == "" and warehouseDddd == "" and languageDddd != "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"


                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'



                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE type = ? AND language = ?", (typeDddd, languageDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""


                    if typeDddd == "" and name_of_itemDddd == "" and warehouseDddd == "" and languageDddd != "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"



                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'


                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE language = ?", (languageDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""






                    if typeDddd == "" and name_of_itemDddd == "" and warehouseDddd != "" and languageDddd != "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"


                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'




                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE language = ? AND warehouse = ?", (languageDddd, warehouseDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""


                    if typeDddd == "" and name_of_itemDddd != "" and warehouseDddd == "" and languageDddd == "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"


                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'



                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE name_of_item = ?", (name_of_itemDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""




                    if typeDddd == "" and name_of_itemDddd != "" and warehouseDddd != "" and languageDddd != "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"


                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'


                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE name_of_item = ? AND language = ? AND warehouse = ?", (name_of_itemDddd, languageDddd, warehouseDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""


                    if typeDddd == "" and name_of_itemDddd != "" and warehouseDddd != "" and languageDddd == "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"


                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'



                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE name_of_item = ? AND warehouse = ?", (name_of_itemDddd, warehouseDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""






                    if typeDddd == "" and name_of_itemDddd == "" and warehouseDddd != "" and languageDddd == "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"



                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'



                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE warehouse = ?", (warehouseDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""


                    if typeDddd == "" and name_of_itemDddd != "" and warehouseDddd == "" and languageDddd != "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"


                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'


                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE name_of_item = ? AND language = ?", (name_of_itemDddd, languageDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""



                    if typeDddd != "" and name_of_itemDddd == "" and warehouseDddd != "" and languageDddd == "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"


                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'



                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE type = ? AND warehouse = ?", (typeDddd, warehouseDddd,)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""



                    if typeDddd != "" and name_of_itemDddd != "" and warehouseDddd == "" and languageDddd != "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"


                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'


                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE type = ? AND language = ? AND name_of_item  = ?", (typeDddd, name_of_itemDddd, languageDddd)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""



                    if typeDddd != "" and name_of_itemDddd == "" and warehouseDddd != "" and languageDddd != "":


                        open_db()


                        random_number = random.randint(1, 10000)
                        date.today()
                        dtNw = date.today()
                        dtNw = dtNw.strftime("%d_%m_%Y")
                        koExc = "Kalpa_Taru_" + dtNw + "_" + str(random_number) + ".xlsx"


                        asbKoExc = koExc + ' wird gelöscht da alle Dateien im tmp Order automatisiert gelöscht werden.'


                        title = "Kalpa Taru " + dtNw + "_" + str(random_number)
                        koExc = os.path.join(repDirPath, koExc)

                        window108['outKalpa'].update(asbKoExc)

                        rows = cursor.execute("SELECT * FROM books WHERE language = ? AND type = ? AND warehouse = ?",  (languageDddd, typeDddd, warehouseDddd)).fetchall()



                        colNames = cursor.execute("SELECT * FROM books")
                        colNames = [cn[0] for cn in colNames.description]
                        colNames = tuple(colNames)
                        rows.insert(0, colNames)
                        tableBooks = rows

                        df = pd.DataFrame(tableBooks)




                        writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                        df.to_excel(writer, sheet_name='Kalpa_Taru', index=False)

                        #workbook  = writer.book
                        worksheet = writer.sheets['Kalpa_Taru']


                        #worksheet.write(4, 10, title)


                        writer.save()



                        wb = openpyxl.load_workbook(koExc)
                        sheet = wb['Kalpa_Taru']

                        sheet.delete_rows(1)
                        sheet.insert_rows(idx=0, amount=3)
                        sheet.cell(row=2, column=1).value = title
                        wb.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start excel.exe "{absolutePath}"')

                        connection.close()

                        typeDddd = ""
                        name_of_itemDddd = ""
                        warehouseDddd = ""
                        languageDddd = ""



                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"
                    lineToChange = []
                    connection.close()
                    window108.close()

                if event == "Back" or event == sg.WIN_CLOSED:
                    direct = "bbtReports"
                    connection.close()
                    window108.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window108.close()



    while direct == "laksmi":



            dtNew = datetime.datetime.now()
            dtNew = dtNew.strftime('%Y-%m-%d')


            open_db()


            try:
                inRech_NrDistri = nameLa
                # inMatchcodeDistri = custo[2]
                # inVerteiler_IDDistri = custo[0]
                # inVerteiler_IDDistri = custo[3]


            except:
                inVerteiler_IDDistri = ""
                inRech_NrDistri = ""
                inMatchcodeDistri = ''
                inVerteiler_IDDistri = ""


            # cursor.execute("CREATE TABLE kundenbewegung (number INTEGER PRIMARY KEY, Datum TEXT, Verteiler_ID TEXT, Rech_Nr TEXT, Schulden REAL, Bezahlt REAL)")
            dtNew = date.today()
            dtNew = dtNew.strftime('%Y-%m-%d')



            layout1 = [
                [sg.Text("", size=(20, 1))],
                [sg.Button('1. Verteiler', size=(9, 1)), sg.Text(size=(9, 1), key="outRech"), sg.CalendarButton('2. Datum', target='-CAL-', size=(9, 1), pad=None, key='_CALENDAR_', format=('%Y-%m-%d')), sg.In(dtNew, key='-CAL-', readonly=True, size=(10, 1)), sg.InputText("3. Betrag", size=(9, 1), key="bezahlt"), sg.Button('4. Zahlung übernehmen', size=(18, 1)), sg.Text("", size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text(key="outKun")],

                [sg.Text(key="outBuch")],
            ]


            window1 = sg.Window("Verteiler Eingangszahlungen im System eintragen", layout1, finalize=True)
            window1['outRech'].update(inRech_NrDistri)
            #window1['outKun'].update(inVerteiler_IDDistri + " " + inMatchcodeDistri + "" + inVerteiler_IDDistri)

            # Datensatz hinzüfgen

            while direct == "laksmi":

                event, values = window1.read()




                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"
                    window1['bezahlt'].update('')
                    window1['-CAL-'].update('')
                    window1['outRech'].update("")
                    window1['outKun'].update("")
                    nameLa = ""


                    custo = []
                    rechNumb = ""
                    connection.close()
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:

                    window1['bezahlt'].update('')
                    window1['-CAL-'].update('')
                    window1['outRech'].update("")
                    window1['outKun'].update("")

                    custo = []
                    rechNumb = ""

                    direct = "Kunden"
                    window1.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()


                if event == "4. Zahlung übernehmen" or event == sg.WIN_CLOSED:


                    try:
                        userRech = inRech_NrDistri
                        if userRech:
                            rechAuWa = "OK"
                        else:
                            rechAuWa = "NOK"
                            #print("userRechNOK")
                    except:
                        rechAuWa = "NOK"


                    try:
                        userDA = values['-CAL-']
                        if userDA:
                            userAuDA = "OK"
                        else:
                            userAuDA = "NOK"
                    except:
                        userAuDA = "NOK"



                    try:
                        userbeza = values['bezahlt']
                        userbeza = float(userbeza)
                        userbeza = round(userbeza, 2)
                        if float(userbeza):
                            userAubeza = "OK"
                            #print("userbeza ok")
                        else:
                            userAubeza = "NOK"

                    except:
                        userAubeza = "NOK"



                    if rechAuWa == "OK" and userAuDA == "OK" and userAubeza == "OK":
                        open_db()

                        #cursor.execute("INSERT INTO kundenbewegung (Datum, Verteiler_ID, Betrag, Vorgang) VALUES (?, ?, ?, ?)", (userDA, userRech, userbeza, "Ü"))


                        nameVer = [nameVer[0] for nameVer in cursor.execute("SELECT Name FROM distributer Where Verteiler_ID=? ",(userRech,))]

                        if len(nameVer) == 1:

                            for i in nameVer:
                                nameVer = i



                        cursor.execute("INSERT INTO kundenbewegung (Verteiler_ID,  Name, Vorgang, Datum, Betrag) VALUES (?, ?, ?, ?, ?)", (str(userRech), str(nameVer), 'Ü', str(userDA), str(userbeza)))

                        connection.commit()

                        window1['outBuch'].update("Für Verteiler " + str(nameLa) + " wurde für " + userDA + " folgeder Zahlungseingang in der " + dbName + " gespeichert: "  + str(userbeza) + " Euro.", text_color='Green')
                        window1['bezahlt'].update('')
                        window1['-CAL-'].update('')
                        window1['outRech'].update("")
                        window1['outKun'].update("")


                        content = [kundenbewegung[0] for kundenbewegung in cursor.execute("SELECT Vorgangsnummer FROM kundenbewegung Where Verteiler_ID=?  ORDER BY date(Datum_Zahlungsziel) ASC",(userRech,))]
                        #print('ooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooo' + str(content))
                        for vorgangsnummer in content:

                            if vorgangsnummer is not None:
                                #print('eeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeee')
                                #print(vorgangsnummer)

                                betrag = cursor.execute("SELECT Betrag FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()


                                for t in betrag:
                                        for x in t:
                                            betrag = x


                                betrag = float(betrag)
                                #print('betrag')
                                #print(betrag)
                                #print('userbeza')
                                #print(userbeza)


 



                                bezahlung = cursor.execute("SELECT Forderung FROM kundenbewegung WHERE Vorgangsnummer =?",(vorgangsnummer,)).fetchall()

                                for t in bezahlung:
                                    for x in t:
                                        bezahlung = x
                                #print(bezahlung)



                                if not 'bezahlt' in str(bezahlung):


                                    rech = userbeza + betrag


                                    if rech >= 0:

                                        cursor.execute("UPDATE kundenbewegung SET Forderung=? WHERE Vorgangsnummer=?", ('bezahlt', vorgangsnummer))

                                        connection.commit()

                                        userbeza = userbeza + betrag

                                        #print(vorgangsnummer)
                                        #print('bezelt')


                        connection.close()







                    if userAubeza == "NOK":

                        window1['outBuch'].update("Der Betrag muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88..", text_color='Red')


                    if userAuDA == "NOK":

                        window1['outBuch'].update("Das Datum an dem bezahlt wurde muss ausgewählt werden.", text_color='Red')

                    if rechAuWa == "NOK":

                        window1['outBuch'].update("Der Verteiler muss ausgewählt werden.", text_color='Red')



                if event == "1. Verteiler" or event == sg.WIN_CLOSED:
                    direct = "findRechNumber"
                    window1.close()


            while direct == "findRechNumber":
                open_db()
                content = [distributer[0] for distributer in cursor.execute("SELECT Verteiler_ID FROM distributer")]
                content = list(set(content))
                content = sorted(content)
                layout1 = [
                    [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                    [sg.Text("Verteiler auswählen:")],
                    [sg.Combo(content, readonly=True, size=(79, 1), key='nameKey')],
                    [sg.Button('Ok', size=(9,1)), sg.Text(key="outPostleitzahl_OrtChooseBuch2")],
                ]

                window1 = sg.Window("Verteiler finden", layout1)

                while direct == "findRechNumber":

                    event, values = window1.read()

                    if event == "Ok" or event == sg.WIN_CLOSED:

                        chooseContent = ""

                        try:
                            nameLa = values['nameKey']
                            #print("yesiii " + nameLa)
                            #print(type(nameLa))
                            direct = "laksmi"
                            window1.close()


                        except:
                            pass

                    else:
                        direct = "findRechNumber"
                        window1['outPostleitzahl_OrtChooseBuch2'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')



                    if event == "Gauranga" or event == sg.WIN_CLOSED:
                        direct = "menu"
                        lineToChange = []
                        connection.close()
                        window1.close()

                    if event == "Back" or event == sg.WIN_CLOSED:
                        lineToChange = []
                        direct = "eingang"
                        window1.close()

                    if event == "Exit" or event == sg.WIN_CLOSED:
                        db_backup()
                        direct = "exit"
                        connection.close()
                        window1.close()




            if event == "Exit" or event == sg.WIN_CLOSED:
                db_backup()
                direct = "exit"
                window1.close()




            if event == "Gauranga" or event == sg.WIN_CLOSED:
                direct = "menu"
                lineToChange = []
                window1.close()



    while direct == "versand":


        dtNew = date.today()
        dtNew = dtNew.strftime('%Y-%m-%d')



        open_db()

        try:
            inbookToLieferDict = bookToLieferDict
        except:
            inbookToLieferDict = ""

        try:
            inVerteiler = verteiler
        except:
            inVerteiler = ''


        try:
            inZahlZielKey = outZahlZielKey
        except:
            inZahlZielKey = ''


        try:
            inLager = lager
        except:
            inLager = ''

        try:
            inbüchLi = büchLi
            #print(büchLi)
        except:
            inbüchLi = ''


        try:
            inImport = imporText
            #print(inImport)
        except:
            inImport = ''



        try:

            ininfoStartBA = infoStartBA
            inrechGrunddatensatz = rechGrunddatensatz


        except:

            ininfoStartBA = ''
            inrechGrunddatensatz = ''



        preise = ['BBT_price', 'SP', 'P1', 'P2', 'P3_40', 'P3_30', 'End_P']




        layout1 = [
            [sg.Text("")],
            [sg.Button('Clear', size=(12,1)), sg.Text(size=(100, 1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
            [sg.Text("")],
            [sg.Button('1. Verteiler', size=(12, 1)), sg.Text(size=(14, 1), key="outVerteiler"), sg.Button('2. Lager', size=(12, 1)), sg.Text(size=(14, 1), key="outLager"), sg.Text('3. Preiskategorie'), sg.Combo(preise, size=(12, 1), readonly=True, key='choosePreis') ,sg.CalendarButton('4. Datum', target='-CAL-', size=(12, 1), pad=None, key='_CALENDAR_', format=('%Y-%m-%d')), sg.In(dtNew, key='-CAL-', readonly=True, size=(12, 1)), sg.InputText("5. Porto ", size=(9, 1), key="outPorto"), sg.Button('6. Übernehmen', size=(12, 1))],
            [sg.Text(key="outBuch")],
            [sg.Text(key="outBuch1")],
            [sg.Button('EXCEL Bestellung Vorlage download', size=(28,1)), sg.FileBrowse('1. EXCEL Bestellung auswählen', size=(28,1), key="impKey"), sg.Button('2. EXCEL Bestellung anzeigen', size=(30,1))],
            [sg.Text(size=(138, 3), key="outBuch123")],
            [sg.Button('1. Buch', size=(12, 1)), sg.Text(key="outBuchRE"), sg.InputText("2. Anzahl", size=(11, 1), key="anzahl"), sg.Text("", size=(0, 1)),
            sg.Button('3. Hinzufügen', size=(12, 1)), sg.Button('Vorschau', size=(12, 1))],
            [sg.Text(size=(138, 8), key="outBuch2")],
            [sg.Button('Lieferschein und Rechnung erstellen', size=(28, 1)), sg.Button('Liefers. und Rech. entfernen', size=(28, 1)), sg.InputText("Vorgangsnummer eingeben", size=(25, 1), key="vrgngNummer")],
            [sg.Text(size=(137, 1), key="outBuch3")],
        ]



        window1 = sg.Window("Versand", layout1, finalize=True)
        window1['outVerteiler'].update(inVerteiler)
        window1['outLager'].update(inLager)
        window1['outBuch'].update(inrechGrunddatensatz, text_color='Blue')
        window1['outBuch1'].update(ininfoStartBA)
        window1['outBuchRE'].update(inbüchLi)
        window1['outBuch2'].update(inbookToLieferDict, text_color='Blue')
        window1['outBuch123'].update(inImport)

        # Datensatz hinzüfgen

        while direct == "versand":

            event, values = window1.read()

            if event == "1. Buch" or event == sg.WIN_CLOSED:
                if ersteBestellZeileOK != "OK":
                    window1['outBuch'].update("Zuerst Verteiler, Lager, Preiskategorie und Datum auswäheln.", text_color='Red')
                else:
                    direct = "bücherZuLieferschein"
                    window1.close()



            if event == "Clear" or event == sg.WIN_CLOSED:


                window1['-CAL-'].update('')
                verteiler = ""
                lager = ""
                window1['choosePreis'].update("")
                ersteBestellZeileOK = "NOK"
                zweiteBestellZeileOK = "NOK"
                rechGrunddatensatz = ""
                infoStartBA = ""
                bookToLieferDict = {}
                window1['outVerteiler'].update("")
                window1['outLager'].update("")
                window1['outBuch'].update("")
                window1['outBuch1'].update("")
                window1['outBuch3'].update("")
                window1['outBuchRE'].update("")
                window1['outBuch2'].update("")
                window1['outPorto'].update("")
                window1['anzahl'].update("")
                window1['vrgngNummer'].update("")
                window1['outBuch123'].update("")
                ininfoStartBA = ''
                inrechGrunddatensatz = ''
                inbüchLi = ''
                inLager = ''
                inVerteiler = ''
                ininfoStartBA = ''
                inrechGrunddatensatz = ''
                inbookToLieferDict = ""
                inverteiler = ''
                inlager = ''
                inrechGrunddatensatz = ''
                verteiler = ''
                lager = ''
                rechGrunddatensatz = ''
                büchLi = ''






            if event == "Gauranga" or event == sg.WIN_CLOSED:
                direct = "menu"

                window1['-CAL-'].update('')
                verteiler = ""
                lager = ""
                window1['choosePreis'].update("")
                ersteBestellZeileOK = "NOK"
                zweiteBestellZeileOK ="NOK"
                rechGrunddatensatz = ""
                infoStartBA = ""
                bookToLieferDict = {}
                window1['outPorto'].update("")
                window1['outVerteiler'].update("")
                window1['outLager'].update("")
                window1['outBuch'].update("")
                window1['outBuch1'].update("")
                window1['outBuchRE'].update("")
                window1['outBuch2'].update("")
                window1['anzahl'].update("")
                window1['outBuch123'].update("")
                imporText = ''

                ininfoStartBA = ''
                inrechGrunddatensatz = ''
                inbüchLi = ''
                inLager = ''
                inVerteiler = ''
                ininfoStartBA = ''
                inrechGrunddatensatz = ''
                inbookToLieferDict = ""
                inverteiler = ''
                inlager = ''
                inrechGrunddatensatz = ''
                verteiler = ''
                lager = ''
                rechGrunddatensatz = ''
                büchLi = ''
                connection.close()
                window1.close()

            if event == "Back" or event == sg.WIN_CLOSED:

                window1['-CAL-'].update('')
                window1['outVerteiler'].update("")
                window1['outLager'].update("")
                window1['choosePreis'].update("")
                window1['outPorto'].update("")

                direct = "versand"
                window1.close()

            if event == "Exit" or event == sg.WIN_CLOSED:
                db_backup()
                direct = "exit"
                connection.close()
                window1.close()

            if event == "EXCEL Bestellung Vorlage download" or event == sg.WIN_CLOSED:
                #print("lololololoololoooo")




                direct = "vorlageBestell"

                window1.close()


            if event == "2. EXCEL Bestellung anzeigen" or event == sg.WIN_CLOSED:
                if ersteBestellZeileOK != "OK":
                    window1['outBuch'].update("Zuerst Verteiler, Lager, Preiskategorie und Datum auswäheln.", text_color='Red')
                else:
                    direct = "bestImport"
                    window1.close()


            if event == "3. EXCEL Bestellung nicht mehr anzeigen" or event == sg.WIN_CLOSED:
                if ersteBestellZeileOK != "OK":
                    window1['outBuch'].update("Zuerst Verteiler, Lager, Preiskategorie und Datum auswäheln.", text_color='Red')
                else:
                    direct = "xxxx"
                    window1.close()



            if event == "Liefers. und Rech. entfernen" or event == sg.WIN_CLOSED:

                try:

                    #vorgangsnummer, buchdatensatz, verteilerdatensatz, summe
                    open_db()
                    vrgngNummer = values['vrgngNummer']

                    checkvorgangsnummer = cursor.execute("SELECT vorgangsnummer FROM vorgang WHERE vorgangsnummer =?",(vrgngNummer,)).fetchall()
                    #print(checkvorgangsnummer)
                    for t in checkvorgangsnummer:
                            for x in t:
                                checkvorgangsnummer = x

                    if checkvorgangsnummer == vrgngNummer:

                        bookRestore = cursor.execute("SELECT buchdatensatz FROM vorgang WHERE vorgangsnummer =?",(vrgngNummer,)).fetchall()
                        for t in bookRestore:
                                for x in t:
                                    bookRestore = x


                        #print("bookRestore")
                        #print(bookRestore)


                        grndDtnstzRestore = cursor.execute("SELECT verteilerdatensatz FROM vorgang WHERE vorgangsnummer =?",(vrgngNummer,)).fetchall()
                        for t in grndDtnstzRestore:
                                for x in t:
                                    grndDtnstzRestore = x


                        #print("grndDtnstzRestore")
                        #print(grndDtnstzRestore)


                        summeRestore = cursor.execute("SELECT summe FROM vorgang WHERE vorgangsnummer =?",(vrgngNummer,)).fetchall()

                        for t in summeRestore:
                                for x in t:
                                    summeRestore = x


                        #print("summeRestore")
                        #print(summeRestore)


                        grndDtnstzRestore = eval(grndDtnstzRestore)

                        bookRestore = eval(bookRestore)


                        for k,v in bookRestore.items():
                            #print(k + str(v[0]) + str(v[1]) + str(v[2]))
                            anzahlJB = int(v[0])
                            vID = str(v[3])


                            open_db()

                            anzahlDB = cursor.execute("SELECT quantity FROM books WHERE ID = ?", (k,)).fetchall()
                            anzahlDB = str(anzahlDB)
                            for char in ['(', ')', ',', '\'', ']', '[']:

                                if char in anzahlDB:

                                    anzahlDB = anzahlDB.replace(char, '')

                            #print(anzahlDB)
                            anzahlDB = int(anzahlDB)

                            anazahlNew = anzahlDB + anzahlJB


                            BBT_price = cursor.execute("SELECT BBT_price FROM books WHERE ID = ?", (k,)).fetchall()
                            #print(BBT_price)
                            BBT_price = str(BBT_price)

                            for char in ['(', ')', ',', '\'', ']', '[']:

                                if char in BBT_price:

                                    BBT_price = BBT_price.replace(char, '')

                            BBT_price = float(BBT_price)
                            #print(BBT_price)
                            total_BBT_price_quantity = BBT_price * anazahlNew
                            #print(total_BBT_price_quantity)

                            cursor.execute("UPDATE books SET quantity=?, total_BBT_price_quantity=? WHERE id=? ", (anazahlNew, total_BBT_price_quantity, k))

                            connection.commit()

                            a1 = cursor.execute("SELECT * FROM bücherbewegung WHERE bbID=?", (vID,)).fetchall()

                            a2 = cursor.execute("SELECT * FROM kundenbewegung WHERE vbID=?", (vID,)).fetchall()


                            cursor.execute("DELETE FROM bücherbewegung WHERE bbID=?", (vID,))
                            connection.commit()


                            cursor.execute("DELETE FROM kundenbewegung WHERE vbID=?", (vID,))
                            connection.commit()




                            #print("DELETE")
                            #print(a1)
                            #print(a2)




                            connection.commit()


                        open_db()



                        cursor.execute("DELETE FROM vorgang WHERE vorgangsnummer=?", (vrgngNummer,))

                        connection.commit()
                        connection.close()

                        try:
                            for p in Path(rechLiDir).glob("*"+vrgngNummer+"*"):
                                p.unlink()
                            window1['outBuch3'].update("Bücherentnahmen und Saldenberechnungen beim Vorgang mit der Nummer " + str(vrgngNummer) + " wurden rückgängig gemacht. Lieferschein und Rechnung wurden aus rechnungen_lieferscheine gelöscht.", text_color='Green')

                        except:
                            window1['outBuch3'].update("Bücherentnahmen und Saldenberechnungen beim Vorgang mit der Nummer " + str(vrgngNummer) + " wurden rückgängig gemacht. Lieferschein und Rechnung bitte aus rechnungen_lieferscheine löschen.", text_color='Red')


                        # document.save(koExc)

                        # absolutePath = Path(koExc).resolve()
                        # os.system(f'start WINWORD.EXE "{absolutePath}"')





                        vorgangsnummer = ""
                        sumBi = 0
                        window1['-CAL-'].update('')
                        verteiler = ""
                        lager = ""
                        window1['choosePreis'].update("")
                        ersteBestellZeileOK = "NOK"
                        zweiteBestellZeileOK = "NOK"
                        rechGrunddatensatz = {}
                        infoStartBA = ""
                        bookToLieferDict = {}
                        window1['outVerteiler'].update("")
                        window1['outLager'].update("")
                        window1['outBuch'].update("")
                        window1['outBuch1'].update("")
                        window1['outBuchRE'].update("")
                        window1['outBuch2'].update("")
                        window1['anzahl'].update("")
                        window1['outPorto'].update("")
                        ininfoStartBA = ''
                        inrechGrunddatensatz = ''
                        inbüchLi = ''
                        inLager = ''
                        inVerteiler = ''
                        ininfoStartBA = ''
                        inrechGrunddatensatz = ''
                        inbookToLieferDict = ""
                        inverteiler = ''
                        inlager = ''
                        inrechGrunddatensatz = ''
                        verteiler = ''
                        lager = ''
                        rechGrunddatensatz = ''
                        büchLi = ''



                    else:
                        window1['outBuch3'].update('Entweder wurde keine gültige Vorgnagsnummer eingegeben oder der Vorgang wurde schon aus dem System gelöscht.', text_color='Red')

                        #print(checkvorgangsnummer)
                        #print(vrgngNummer)

                except:
                    window1['outBuch3'].update('Es wurde keine gültige Rechnungs- oder Lieferscheinnummer eingegeben.', text_color='Red')



            if event == "Lieferschein und Rechnung erstellen" or event == sg.WIN_CLOSED:

                if ersteBestellZeileOK != "OK":
                    window1['outBuch'].update("Zuerst Verteiler, Lager, Preiskategorie und Datum auswäheln und dann auf 6. Übernehmen klicken.", text_color='Red')
                    window1['outBuch123'].update("")
                else:
                    if zweiteBestellZeileOK != "ok":
                        window1['outBuch2'].update("Es muss mindestens ein Buch mit dazugehöriger Anzahl hinzugefügt werden.", text_color='Red')
                        window1['outBuch123'].update("")

                    else:


#cursor.execute("CREATE TABLE vorgang (vorgangsnummer TEXT PRIMARY KEY, buchdatensatz TEXT, verteilerdatensatz TEXT, summe INTEGER)")

                    # grndDtnstzRestore = cursor.execute("SELECT vorgangsnummer FROM vorgang WHERE vorgangsnummer =?",(vrgngNummer,)).fetchall()
                    # for t in grndDtnstzRestore:
                    #         for x in t:
                    #             grndDtnstzRestore = x



                        dtNw = date.today()
                        dtNw = dtNw.strftime("%#d%#m%y")

                        lager = rechGrunddatensatz.get('Lager')
                        lager = lager.upper()
                        lager = lager[0]


                        open_db()
                        #vorgaengeVomTag = cursor.execute("SELECT vorgangsnummer FROM vorgang WHERE vorgangsnummer LIKE =?",(dtNw,)).fetchall()
                        vorgaengeVomTag = [vorgaengeVomTag[0] for vorgaengeVomTag in cursor.execute("SELECT vorgangsnummer FROM vorgang WHERE vorgangsnummer LIKE ?",('%'+dtNw+'%',))]
                        #print("mmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmm")
                        #print(vorgaengeVomTag)

                        last = []

                        for strip in vorgaengeVomTag:

                            sep = dtNw

                            stripped = strip.split(sep, 1)[1]
                            #print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
                            #print(stripped)
                            last.append(int(stripped))

                        try:
                            fNumber = int(max(last)) + 1
                        except:
                            fNumber = 1

                        vrgngsnmmr = 'G' + lager + dtNw + str(fNumber)



                        #print('2808 ++++++++++++++++++++++++++++++++++++++++++++++++')
                        for k,v in bookToLieferDict.items():
                            #print(k + str(v[0]) + str(v[1]) + str(v[2]))
                            anzahlJB = int(v[0])
                            sumBu = float(v[2])
                            sumBu = round(sumBu, 2)

                            vID = str(v[3])

                            sumBi += sumBu

                            sumBi = round(sumBi, 2)

                            open_db()
                            kN = cursor.execute("SELECT name_of_item FROM books WHERE ID = ?", (k,)).fetchall()
                            kN = str(kN)
                            for char in ['(', ')', ',', '\'', ']', '[']:

                                if char in kN:

                                       kN = kN.replace(char, '')

                            ##print(name_of_item)
                            anzahlDB = cursor.execute("SELECT quantity FROM books WHERE ID = ?", (k,)).fetchall()
                            anzahlDB = str(anzahlDB)
                            for char in ['(', ')', ',', '\'', ']', '[']:

                                if char in anzahlDB:

                                       anzahlDB = anzahlDB.replace(char, '')

                            #print(anzahlDB)
                            anzahlDB = int(anzahlDB)

                            anazahlNew = anzahlDB - anzahlJB


                            BBT_price = cursor.execute("SELECT BBT_price FROM books WHERE ID = ?", (k,)).fetchall()
                            #print(BBT_price)
                            BBT_price = str(BBT_price)

                            for char in ['(', ')', ',', '\'', ']', '[']:

                                if char in BBT_price:

                                       BBT_price = BBT_price.replace(char, '')

                            BBT_price = float(BBT_price)
                            #print(BBT_price)
                            total_BBT_price_quantity = BBT_price * anazahlNew
                            #print(total_BBT_price_quantity)

                            cursor.execute("UPDATE books SET quantity=?, total_BBT_price_quantity=? WHERE id=? ", (anazahlNew, total_BBT_price_quantity, k))

                            connection.commit()



                            datum = rechGrunddatensatz.get('Datum')

                            verteilerG = rechGrunddatensatz.get('Verteiler')

                            cursor.execute("INSERT INTO bücherbewegung (bbID, book, name_of_item, date, amount, amount_old, Entnahme_Zuführung, Verteiler, Vorgangsnummer) VALUES(?,?,?,?,?,?,?,?,?)", (vID, k, kN, datum, anazahlNew, anzahlDB, anzahlJB, verteilerG, vrgngsnmmr))


                            connection.commit()

                        portopauschaleAD = rechGrunddatensatz.get('Portopauschale')

                        tageAblaufZahlungsziel = rechGrunddatensatz.get('Zahlungsziel')
                        tageAblaufZahlungsziel = int(tageAblaufZahlungsziel)

                        sumBiNoPorto = sumBi

                        sumBi = sumBi + portopauschaleAD
                        sumBi = round(sumBi, 2)


                        sumBiM = -abs(sumBi)
                        sumBiM = round(sumBiM, 2)






                        datumErstellung = datetime.datetime.strptime(datum, '%Y-%m-%d')

                        #print("datumErstellung")
                        #print(datumErstellung)
                        #print(type(datumErstellung))

                        ablaufZahlungsziel = datetime.timedelta(tageAblaufZahlungsziel)



                        datumZahlungsziel = datumErstellung + ablaufZahlungsziel

                        #print("datumZahlungsziel")
                        #print(datumZahlungsziel)
                        #print(type(datumZahlungsziel))

                        WarnungAblaufZahlungsziel = ""

                        if datumZahlungsziel <= datumErstellung:

                            WarnungAblaufZahlungsziel = "Austehende Zahlung ist überfällig!"

                            #print('WarnungAblaufZahlungsziel')

                        #currentDate = date.today()

                        #currentDate = datetime.datetime.now().date()

                        currentDate =  datetime.datetime.now()

                        taTageAblaufZahlungsziel = datumZahlungsziel - currentDate

                        #print("taTageAblaufZahlungsziel")
                        #print(taTageAblaufZahlungsziel)
                        #print(type(taTageAblaufZahlungsziel))

                        open_db()

                        verteiler = rechGrunddatensatz.get('Verteiler')

                        nameVer = [nameVer[0] for nameVer in cursor.execute("SELECT Name FROM distributer Where Verteiler_ID=? ",(verteiler,))]

                        if len(nameVer) == 1:

                            for i in nameVer:
                                nameVer = i


                        vrgngsnmmr = str(vrgngsnmmr)
                        bookToLieferDict = str(bookToLieferDict)
                        rechGrunddatensatz = str(rechGrunddatensatz)
                        sumbiM = str(sumBiM)
                        #print(vrgngsnmmr)
                        cursor.execute("INSERT INTO kundenbewegung (vbID, Verteiler_ID,  Name, Vorgang, Vorgangsnummer, Tage_Zahlungsziel, Datum, Datum_Zahlungsziel, Betrag) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)", (vID, verteiler, nameVer, "L", vrgngsnmmr, str(tageAblaufZahlungsziel), str(datum), str(datumZahlungsziel), sumBiM))


                        cursor.execute("INSERT INTO vorgang (vorgangsnummer, buchdatensatz, verteilerdatensatz, summe) VALUES (?, ?, ?, ?)", (vrgngsnmmr,  bookToLieferDict, rechGrunddatensatz, sumbiM))



                        connection.commit()
                        connection.close()


                        mehrWertS = sumBi / 107 * 7
                        mehrWertS = round(mehrWertS, 2)

                        rechGrunddatensatz = eval(rechGrunddatensatz)

                        bookToLieferDict = eval(bookToLieferDict)

                        positionsB = (len(bookToLieferDict))

                        positionsB += 1







                        verteilerG = rechGrunddatensatz.get('Verteiler')
                        datumG = rechGrunddatensatz.get('Datum')
                        lagerG = rechGrunddatensatz.get('Lager')




                        open_db()
                        nameVer = [nameVer[0] for nameVer in cursor.execute("SELECT Name FROM distributer Where Verteiler_ID=? ",(verteilerG,))]

                        if len(nameVer) == 1:

                            for i in nameVer:
                                nameVer = i

                        straVer = [straVer[0] for straVer in cursor.execute("SELECT Straße FROM distributer Where Verteiler_ID=? ",(verteilerG,))]

                        if len(straVer) == 1:

                            for i in straVer:
                                straVer = i



                        ortV = [ortV[0] for ortV in cursor.execute("SELECT Postleitzahl_Ort FROM distributer Where Verteiler_ID=? ",(verteilerG,))]
                        #print(ortV)
                        if len(ortV) == 1:

                            for i in ortV:
                                ortV = i

                        #print("uuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuu")
                        #print(ortV)
                        #print("uuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuu")



                        UST_ID = [UST_ID[0] for UST_ID in cursor.execute("SELECT UST_ID FROM distributer Where Verteiler_ID=? ",(verteilerG,))]
                        #print(UST_ID)
                        if len(UST_ID) == 1:

                            for i in UST_ID:
                                UST_ID = i

                        #print("uuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuu")
                        #print(UST_ID)
                        #print("uuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuuu")






                        #print("tttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttt")

                        kindOfInvoice = rechGrunddatensatz.get('kindOfInvoice')
                        #print(kindOfInvoice)

                        if kindOfInvoice == "Standard Rechnung":

                            # Rechnung

                            document = Document()




                            def set_column_width(column, width):
                                for cell in column.cells:
                                    cell.width = width

                            paragraph = document.add_paragraph()
                            paragraph.add_run('Musterbuchhandel').font.size = Pt(16)
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = document.add_paragraph()
                            paragraph.add_run('Musterbücher').font.size = Pt(12)
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(22)



                            tabletop = document.add_table(rows=1, cols=2, style="Table Grid")



                            set_column_width(tabletop.columns[0], Cm(15))
                            set_column_width(tabletop.columns[1], Cm(15))


                            heading_row = tabletop.rows[0].cells

                            heading_row[0].text = "Mustersadresse\nTel: 0000000\nFax: 000000\nUSt.-ID 0000000\nSteuer-Nr. 000000"

                            heading_row[1].text = "Bankverbindung: \nMusterbank\nMusterbankname\nIBAN: 000000\nSWIFT-BIC: 00000"


                            paragraph = document.add_paragraph()
                            paragraph.add_run('Empfänger:').underline = True
                            paragraph.paragraph_format.space_before = Pt(22)
                            paragraph.paragraph_format.space_after = Pt(4)

                            paragraph = document.add_paragraph(str(nameVer))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = document.add_paragraph(str(straVer))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = document.add_paragraph(str(ortV))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(22)

                            if UST_ID != None:

                                paragraph = document.add_paragraph("USt.-ID: " + str(UST_ID))
                                paragraph.paragraph_format.space_before = Pt(0)
                                paragraph.paragraph_format.space_after = Pt(22)
                            

                            paragraph = document.add_paragraph()
                            paragraph.add_run('R e c h n u n g').bold = True
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(4)



                            tableRech = document.add_table(rows=2, cols=4, style="Table Grid")


                            set_column_width(tableRech.columns[0], Cm(5))
                            set_column_width(tableRech.columns[1], Cm(5))
                            set_column_width(tableRech.columns[2], Cm(5))
                            set_column_width(tableRech.columns[3], Cm(5))


                            heading_row = tableRech.rows[0].cells

                            heading_row[0].text = "Nummer:"

                            heading_row[1].text = str(vrgngsnmmr)


                            heading_row[2].text = "Datum:"

                            heading_row[3].text = str(datumG)


                            data_rowre = tableRech.rows[1].cells

                            data_rowre[0].text = "Bestellung am:"
                            data_rowre[1].text = str(datumG)


                            data_rowre[2].text = "Lieferung am:"
                            data_rowre[3].text = str(datumG)





                            paragraph = document.add_paragraph('Zahlungsbedingungen: 10 Tage')
                            paragraph.paragraph_format.space_before = Pt(22)
                            paragraph.paragraph_format.space_after = Pt(22)

                            table = document.add_table(rows=positionsB, cols=6, style="Table Grid")



                            set_column_width(table.columns[0], Cm(1))
                            set_column_width(table.columns[1], Cm(2))
                            set_column_width(table.columns[2], Cm(2))
                            set_column_width(table.columns[3], Cm(25))
                            set_column_width(table.columns[4], Cm(6))
                            set_column_width(table.columns[5], Cm(8))

                            heading_row = table.rows[0].cells

                            heading_row[0].text = "Pos"
                            heading_row[1].text = "Menge"
                            heading_row[2].text = "Typ"
                            heading_row[3].text = "Text"
                            heading_row[4].text = "Einzelpr."
                            heading_row[5].text = "Gesamtpr."

                            count = 0

                            for k,v in bookToLieferDict.items():
                                #print(k + str(v[0]) + str(v[1]) + str(v[2]))
                                anzahlJB = int(v[0])
                                preisBue = float(v[1])
                                sumBu = float(v[2])
                                sumBu = round(sumBu, 2)
                                vID = str(v[3])


                                typeRe = cursor.execute("SELECT type FROM books WHERE ID = ?", (k,)).fetchall()

                                for t in typeRe:
                                    for x in t:
                                        typeRe = x


                                #print(count)

                                count += 1

                                data_row = table.rows[count].cells



                                data_row[0].text = str(count)
                                data_row[1].text = str(anzahlJB)
                                data_row[2].text = str(typeRe)


                                open_db()
                                kN = cursor.execute("SELECT name_of_item FROM books WHERE ID = ?", (k,)).fetchall()
                                kN = str(kN)
                                for char in ['(', ')', ',', '\'', ']', '[']:

                                    if char in kN:

                                        kN = kN.replace(char, '')


                                data_row[3].text = str(kN)

                                preisBue = format(preisBue, '.2f')
                                preisBue = str(preisBue)

                                preisBue = preisBue.replace('.',',')
                                preisBue = preisBue + " Euro"



                                data_row[4].text = str(preisBue)
                                data_row[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT



                                sumBu = format(sumBu, '.2f')
                                sumBu = str(sumBu)

                                sumBu = sumBu.replace('.',',')

                                sumBu = (re.sub(r'(?<!^)(?=(\d{3})+,)', r'.', sumBu))

                                sumBu = sumBu + " Euro"




                                data_row[5].text = str(sumBu)
                                data_row[5].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            if porto != 0:

                                paragraph = document.add_paragraph('')

                                table21 = document.add_table(rows=1, cols=2, style="Table Grid")

                                set_column_width(table21.columns[0], Cm(2.5))
                                set_column_width(table21.columns[1], Cm(2.5))

                                heading_row2 = table21.rows[0].cells

                                heading_row2[0].text = "Portopauschale:"


                                portoAD = rechGrunddatensatz.get('Portopauschale')
                                #print('portoAD')
                                #print(portoAD)
                                portoAD = format(portoAD, '.2f')
                                portoAD = str(portoAD)

                                portoAD = portoAD.replace('.',',')
                                portoAD = portoAD + " Euro"

                                heading_row2[1].text = str(portoAD)
                                heading_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            paragraph = document.add_paragraph('')

                            table2 = document.add_table(rows=2, cols=2, style="Table Grid")

                            set_column_width(table2.columns[0], Cm(5.2))
                            set_column_width(table2.columns[1], Cm(4))

                            heading_row2 = table2.rows[0].cells

                            heading_row2[0].text = "Rechnungsbetrag brutto:"



                            sumBi = format(sumBi, '.2f')


                            sumBi = str(sumBi)



                            sumBi = sumBi.replace('.',',')



                            sumBi = (re.sub(r'(?<!^)(?=(\d{3})+,)', r'.', sumBi))



                            sumBi = sumBi + " Euro"


                            heading_row2[1].text = str(sumBi)
                            heading_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT



                            data_row2 = table2.rows[1].cells

                            data_row2[0].text = "enthält 7 % Mehrwertsteuer:"

                            mehrWertS = format(mehrWertS, '.2f')
                            mehrWertS = str(mehrWertS)
                            mehrWertS = mehrWertS.replace('.',',')

                            mehrWertS = (re.sub(r'(?<!^)(?=(\d{3})+,)', r'.', mehrWertS))

                            mehrWertS = mehrWertS + " Euro"


                            data_row2[1].text = str(mehrWertS)
                            data_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            paragraph = document.add_paragraph()
                            paragraph.add_run('Die gelieferte Ware bleibt bis zur vollständigen Bezahlung Eigentum des Lieferanten.').underline = True
                            paragraph.paragraph_format.space_before = Pt(22)
                            paragraph.paragraph_format.space_after = Pt(4)


                            sumBitDaNa = sumBi.replace(',','_')

                            sumBitDaNa = sumBitDaNa.replace(' ','_')

                            verteilerhDaNa = verteilerG.replace(' ','_')

                            koExc = str(vrgngsnmmr) + "_R_" + sumBitDaNa + "_" + verteilerhDaNa + ".docx"

                            koExc = os.path.join(rechLiDir, koExc)

                            try:
                                document.save(koExc)

                            except:

                                random_number = random.randint(1, 10000)

                                vrgngsnmmr = vrgngsnmmr + "_" + str(random_number)

                                koExc = str(vrgngsnmmr) + "_R_" + sumBitDaNa + "_" + verteilerhDaNa + ".docx"

                                koExc = os.path.join(rechLiDir, koExc)

                                document.save(koExc)



                            absolutePath = Path(koExc).resolve()
                            os.system(f'start WINWORD.EXE "{absolutePath}"')


    # Lieferschein

                            documentl = Document()




                            def set_column_width(column, width):
                                for cell in column.cells:
                                    cell.width = width

                            paragraph = documentl.add_paragraph()
                            paragraph.add_run('Musterbuchhandel').font.size = Pt(16)
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = documentl.add_paragraph()
                            paragraph.add_run('Musterbücher').font.size = Pt(12)
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(22)



                            tabletop = documentl.add_table(rows=1, cols=2, style="Table Grid")



                            set_column_width(tabletop.columns[0], Cm(15))
                            set_column_width(tabletop.columns[1], Cm(15))


                            heading_row = tabletop.rows[0].cells

                            heading_row[0].text = "Mustersadresse\nTel: 0000000\nFax: 000000\nUSt.-ID 0000000\nSteuer-Nr. 000000"

                            heading_row[1].text = "Bankverbindung: \nMusterbank\nMusterbankname\nIBAN: 000000\nSWIFT-BIC: 00000"


                            paragraph = documentl.add_paragraph()
                            paragraph.add_run('Empfänger:').underline = True
                            paragraph.paragraph_format.space_before = Pt(22)
                            paragraph.paragraph_format.space_after = Pt(4)

                            paragraph = documentl.add_paragraph(str(nameVer))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = documentl.add_paragraph(str(straVer))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = documentl.add_paragraph(str(ortV))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(22)

                            paragraph = documentl.add_paragraph()
                            paragraph.add_run('L i e f e r s c h e i n').bold = True
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(4)



                            tableRech = documentl.add_table(rows=1, cols=6, style="Table Grid")


                            set_column_width(tableRech.columns[0], Cm(2))
                            set_column_width(tableRech.columns[1], Cm(4))
                            set_column_width(tableRech.columns[2], Cm(2))
                            set_column_width(tableRech.columns[3], Cm(4))
                            set_column_width(tableRech.columns[4], Cm(2))
                            set_column_width(tableRech.columns[5], Cm(4))


                            heading_row = tableRech.rows[0].cells



                            lager = rechGrunddatensatz.get('Lager')


                            heading_row[0].text = "Lager:"

                            heading_row[1].text = str(lager)

                            heading_row[2].text = "Nummer:"

                            heading_row[3].text = str(vrgngsnmmr)


                            heading_row[4].text = "Datum:"

                            heading_row[5].text = str(datumG)



                            paragraph = documentl.add_paragraph('')

                            table = documentl.add_table(rows=positionsB, cols=6, style="Table Grid")




                            set_column_width(table.columns[0], Cm(1))
                            set_column_width(table.columns[1], Cm(2))
                            set_column_width(table.columns[2], Cm(2))
                            set_column_width(table.columns[3], Cm(25))
                            set_column_width(table.columns[4], Cm(6))
                            set_column_width(table.columns[5], Cm(8))

                            heading_row = table.rows[0].cells

                            heading_row[0].text = "Pos"
                            heading_row[1].text = "Menge"
                            heading_row[2].text = "Typ"
                            heading_row[3].text = "Text"
                            heading_row[4].text = "Einzelpr."
                            heading_row[5].text = "Gesamtpr."

                            count = 0

                            for k,v in bookToLieferDict.items():
                                #print(k + str(v[0]) + str(v[1]) + str(v[2]))
                                anzahlJB = int(v[0])
                                preisBue = float(v[1])
                                sumBu = float(v[2])
                                sumBu = round(sumBu, 2)
                                vID = str(v[3])


                                typeRe = cursor.execute("SELECT type FROM books WHERE ID = ?", (k,)).fetchall()

                                for t in typeRe:
                                    for x in t:
                                        typeRe = x


                                #print(count)

                                count += 1

                                data_row = table.rows[count].cells



                                data_row[0].text = str(count)
                                data_row[1].text = str(anzahlJB)
                                data_row[2].text = str(typeRe)

                                open_db()
                                kN = cursor.execute("SELECT name_of_item FROM books WHERE ID = ?", (k,)).fetchall()
                                kN = str(kN)
                                for char in ['(', ')', ',', '\'', ']', '[']:

                                    if char in kN:

                                        kN = kN.replace(char, '')

                                data_row[3].text = str(kN)



                                preisBue = f'{preisBue:.2f}'
                                preisBue = str(preisBue)
                                preisBue = preisBue.replace('.',',')
                                preisBue = preisBue + " Euro"

                                data_row[4].text = str(preisBue)
                                data_row[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                                sumBu = f'{sumBu:.2f}'
                                sumBu = str(sumBu)
                                sumBu = sumBu.replace('.',',')

                                sumBu = (re.sub(r'(?<!^)(?=(\d{3})+,)', r'.', sumBu))

                                sumBu = sumBu + " Euro"


                                data_row[5].text = str(sumBu)
                                data_row[5].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            if porto != 0:

                                paragraph = documentl.add_paragraph('')

                                table21 = documentl.add_table(rows=1, cols=2, style="Table Grid")

                                set_column_width(table21.columns[0], Cm(2.5))
                                set_column_width(table21.columns[1], Cm(2.5))

                                heading_row2 = table21.rows[0].cells

                                heading_row2[0].text = "Portopauschale:"


                                portoAD = rechGrunddatensatz.get('Portopauschale')
                                #print('portoAD')
                                #print(portoAD)
                                portoAD = format(portoAD, '.2f')
                                portoAD = str(portoAD)

                                portoAD = portoAD.replace('.',',')
                                portoAD = portoAD + " Euro"

                                heading_row2[1].text = str(portoAD)
                                heading_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            paragraph = documentl.add_paragraph('')

                            table2 = documentl.add_table(rows=1, cols=2, style="Table Grid")

                            set_column_width(table2.columns[0], Cm(2.5))
                            set_column_width(table2.columns[1], Cm(4))

                            heading_row2 = table2.rows[0].cells

                            heading_row2[0].text = "Gesamtbetrag:"





                            heading_row2[1].text = str(sumBi)
                            heading_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


                            sumBitDaNa = sumBi.replace(',','_')

                            sumBitDaNa = sumBitDaNa.replace(' ','_')

                            verteilerhDaNa = verteilerG.replace(' ','_')

                            koExc = str(vrgngsnmmr) + "_L_" + sumBitDaNa + "_" + verteilerhDaNa + ".docx"

                            koExc = os.path.join(rechLiDir, koExc)

                            try:
                                documentl.save(koExc)

                            except:

                                random_number = random.randint(1, 10000)

                                vrgngsnmmr = vrgngsnmmr + "_" + str(random_number)

                                koExc = str(vrgngsnmmr) + "_L_" + sumBitDaNa + "_" + verteilerhDaNa + ".docx"

                                koExc = os.path.join(rechLiDir, koExc)

                                documentl.save(koExc)



                            absolutePath = Path(koExc).resolve()
                            os.system(f'start WINWORD.EXE "{absolutePath}"')

                            window1['outBuch3'].update("Der Vorgang " + str(vrgngsnmmr) + " wurde erfolgreich erfasst und in der Datenbank gespeichert. Lieferschein und Rechnung wurden erstellt und unter " + str(rechLiDir) + " gespeichert.", text_color='Green')

                            vrgngsnmmr = ""
                            vorgangsnummer = ""
                            sumBi = 0
                            window1['-CAL-'].update('')
                            verteiler = ""
                            lager = ""
                            window1['choosePreis'].update("")
                            ersteBestellZeileOK = "NOK"
                            zweiteBestellZeileOK = "NOK"
                            rechGrunddatensatz = {}
                            infoStartBA = ""
                            bookToLieferDict = {}
                            window1['outVerteiler'].update("")
                            window1['outLager'].update("")
                            window1['outBuch'].update("")
                            window1['outBuch1'].update("")
                            window1['outBuchRE'].update("")
                            window1['outBuch2'].update("")
                            window1['anzahl'].update("")
                            window1['outBuch123'].update("")
                            ininfoStartBA = ''
                            inrechGrunddatensatz = ''
                            inbüchLi = ''
                            inLager = ''
                            inVerteiler = ''
                            ininfoStartBA = ''
                            inrechGrunddatensatz = ''
                            inbookToLieferDict = ""
                            inverteiler = ''
                            inlager = ''
                            inrechGrunddatensatz = ''
                            verteiler = ''
                            lager = ''
                            rechGrunddatensatz = ''
                            büchLi = ''


                            ersteBestellZeileOK = "NOK"
                            zweiteBestellZeileOK = "NOK"


                        elif kindOfInvoice == "Innergemeinschaftliche Lieferung": 

                            # Rechnung

                            document = Document()




                            def set_column_width(column, width):
                                for cell in column.cells:
                                    cell.width = width

                            paragraph = document.add_paragraph()
                            paragraph.add_run('Musterbuchhandel').font.size = Pt(16)
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = document.add_paragraph()
                            paragraph.add_run('Musterbücher').font.size = Pt(12)
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(22)



                            tabletop = document.add_table(rows=1, cols=2, style="Table Grid")



                            set_column_width(tabletop.columns[0], Cm(15))
                            set_column_width(tabletop.columns[1], Cm(15))


                            heading_row = tabletop.rows[0].cells

                            heading_row[0].text = "Mustersadresse\nTel: 0000000\nFax: 000000\nUSt.-ID 0000000\nSteuer-Nr. 000000"

                            heading_row[1].text = "Bankverbindung: \nMusterbank\nMusterbankname\nIBAN: 000000\nSWIFT-BIC: 00000"


                            paragraph = document.add_paragraph()
                            paragraph.add_run('Empfänger:').underline = True
                            paragraph.paragraph_format.space_before = Pt(22)
                            paragraph.paragraph_format.space_after = Pt(4)

                            paragraph = document.add_paragraph(str(nameVer))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = document.add_paragraph(str(straVer))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = document.add_paragraph(str(ortV))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(22)

                            if UST_ID != None:

                                paragraph = document.add_paragraph("USt.-ID: " + str(UST_ID))
                                paragraph.paragraph_format.space_before = Pt(0)
                                paragraph.paragraph_format.space_after = Pt(22)

                            paragraph = document.add_paragraph()
                            paragraph.add_run('R e c h n u n g').bold = True
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(4)



                            tableRech = document.add_table(rows=2, cols=4, style="Table Grid")


                            set_column_width(tableRech.columns[0], Cm(5))
                            set_column_width(tableRech.columns[1], Cm(5))
                            set_column_width(tableRech.columns[2], Cm(5))
                            set_column_width(tableRech.columns[3], Cm(5))


                            heading_row = tableRech.rows[0].cells

                            heading_row[0].text = "Nummer:"

                            heading_row[1].text = str(vrgngsnmmr)


                            heading_row[2].text = "Datum:"

                            heading_row[3].text = str(datumG)


                            data_rowre = tableRech.rows[1].cells

                            data_rowre[0].text = "Bestellung am:"
                            data_rowre[1].text = str(datumG)


                            data_rowre[2].text = "Lieferung am:"
                            data_rowre[3].text = str(datumG)





                            paragraph = document.add_paragraph('Zahlungsbedingungen: 10 Tage')
                            paragraph.paragraph_format.space_before = Pt(22)
                            paragraph.paragraph_format.space_after = Pt(22)

                            table = document.add_table(rows=positionsB, cols=6, style="Table Grid")



                            set_column_width(table.columns[0], Cm(1))
                            set_column_width(table.columns[1], Cm(2))
                            set_column_width(table.columns[2], Cm(2))
                            set_column_width(table.columns[3], Cm(25))
                            set_column_width(table.columns[4], Cm(6))
                            set_column_width(table.columns[5], Cm(8.5))

                            heading_row = table.rows[0].cells

                            heading_row[0].text = "Pos"
                            heading_row[1].text = "Menge"
                            heading_row[2].text = "Typ"
                            heading_row[3].text = "Text"
                            heading_row[4].text = "Einzelpr."
                            heading_row[5].text = "gesamt brutto"

                            count = 0

                            for k,v in bookToLieferDict.items():
                                #print(k + str(v[0]) + str(v[1]) + str(v[2]))
                                anzahlJB = int(v[0])
                                preisBue = float(v[1])
                                sumBu = float(v[2])
                                sumBu = round(sumBu, 2)
                                vID = str(v[3])



                                typeRe = cursor.execute("SELECT type FROM books WHERE ID = ?", (k,)).fetchall()

                                for t in typeRe:
                                    for x in t:
                                        typeRe = x


                                #print(count)

                                count += 1

                                data_row = table.rows[count].cells



                                data_row[0].text = str(count)
                                data_row[1].text = str(anzahlJB)
                                data_row[2].text = str(typeRe)


                                open_db()
                                kN = cursor.execute("SELECT name_of_item FROM books WHERE ID = ?", (k,)).fetchall()
                                kN = str(kN)
                                for char in ['(', ')', ',', '\'', ']', '[']:

                                    if char in kN:

                                        kN = kN.replace(char, '')


                                data_row[3].text = str(kN)

                                preisBue = format(preisBue, '.2f')
                                preisBue = str(preisBue)

                                preisBue = preisBue.replace('.',',')
                                preisBue = preisBue + " Euro"



                                data_row[4].text = str(preisBue)
                                data_row[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT



                                sumBu = format(sumBu, '.2f')
                                sumBu = str(sumBu)

                                sumBu = sumBu.replace('.',',')

                                sumBu = (re.sub(r'(?<!^)(?=(\d{3})+,)', r'.', sumBu))

                                sumBu = sumBu + " Euro"




                                data_row[5].text = str(sumBu)
                                data_row[5].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            if porto != 0:

                                paragraph = document.add_paragraph('')

                                table21 = document.add_table(rows=1, cols=2, style="Table Grid")

                                set_column_width(table21.columns[0], Cm(2.5))
                                set_column_width(table21.columns[1], Cm(2.5))

                                heading_row2 = table21.rows[0].cells

                                heading_row2[0].text = "Portopauschale:"


                                portoAD = rechGrunddatensatz.get('Portopauschale')
                                #print('portoAD')
                                #print(portoAD)
                                portoAD = format(portoAD, '.2f')
                                portoAD = str(portoAD)

                                portoAD = portoAD.replace('.',',')
                                portoAD = portoAD + " Euro"

                                heading_row2[1].text = str(portoAD)
                                heading_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            paragraph = document.add_paragraph('')

                            table2 = document.add_table(rows=1, cols=2, style="Table Grid")

                            set_column_width(table2.columns[0], Cm(2.5))
                            set_column_width(table2.columns[1], Cm(4))


                            heading_row2 = table2.rows[0].cells

                            heading_row2[0].text = "Rechnungsbetrag:"






                            sumBi = format(sumBi, '.2f')


                            sumBi = str(sumBi)



                            sumBi = sumBi.replace('.',',')



                            sumBi = (re.sub(r'(?<!^)(?=(\d{3})+,)', r'.', sumBi))



                            sumBi = sumBi + " Euro"


                            heading_row2[1].text = str(sumBi)
                            heading_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


                            paragraph = document.add_paragraph()
                            paragraph.add_run('Innergemeinschaftliche Lieferung').underline = True
                            paragraph.paragraph_format.space_before = Pt(22)
                            paragraph.paragraph_format.space_after = Pt(4)                            


                            paragraph = document.add_paragraph()
                            paragraph.add_run('Die gelieferte Ware bleibt bis zur vollständigen Bezahlung Eigentum des Lieferanten.').underline = True
                            paragraph.paragraph_format.space_before = Pt(22)
                            paragraph.paragraph_format.space_after = Pt(4)


                            sumBitDaNa = sumBi.replace(',','_')

                            sumBitDaNa = sumBitDaNa.replace(' ','_')

                            verteilerhDaNa = verteilerG.replace(' ','_')

                            koExc = str(vrgngsnmmr) + "_R_" + sumBitDaNa + "_" + verteilerhDaNa + ".docx"

                            koExc = os.path.join(rechLiDir, koExc)

                            try:
                                document.save(koExc)

                            except:

                                random_number = random.randint(1, 10000)

                                vrgngsnmmr = vrgngsnmmr + "_" + str(random_number)

                                koExc = str(vrgngsnmmr) + "_R_" + sumBitDaNa + "_" + verteilerhDaNa + ".docx"


                                koExc = os.path.join(rechLiDir, koExc)

                                document.save(koExc)



                            absolutePath = Path(koExc).resolve()
                            os.system(f'start WINWORD.EXE "{absolutePath}"')


    # Lieferschein

                            documentl = Document()




                            def set_column_width(column, width):
                                for cell in column.cells:
                                    cell.width = width

                            paragraph = documentl.add_paragraph()
                            paragraph.add_run('Musterbuchhandel').font.size = Pt(16)
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = documentl.add_paragraph()
                            paragraph.add_run('Musterbücher').font.size = Pt(12)
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(22)



                            tabletop = documentl.add_table(rows=1, cols=2, style="Table Grid")



                            set_column_width(tabletop.columns[0], Cm(15))
                            set_column_width(tabletop.columns[1], Cm(15))


                            heading_row = tabletop.rows[0].cells

                            heading_row[0].text = "Mustersadresse\nTel: 0000000\nFax: 000000\nUSt.-ID 0000000\nSteuer-Nr. 000000"

                            heading_row[1].text = "Bankverbindung: \nMusterbank\nMusterbankname\nIBAN: 000000\nSWIFT-BIC: 00000"


                            paragraph = documentl.add_paragraph()
                            paragraph.add_run('Empfänger:').underline = True
                            paragraph.paragraph_format.space_before = Pt(22)
                            paragraph.paragraph_format.space_after = Pt(4)

                            paragraph = documentl.add_paragraph(str(nameVer))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = documentl.add_paragraph(str(straVer))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)

                            paragraph = documentl.add_paragraph(str(ortV))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(22)

                            paragraph = documentl.add_paragraph()
                            paragraph.add_run('L i e f e r s c h e i n').bold = True
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(4)



                            tableRech = documentl.add_table(rows=1, cols=6, style="Table Grid")


                            set_column_width(tableRech.columns[0], Cm(2))
                            set_column_width(tableRech.columns[1], Cm(4))
                            set_column_width(tableRech.columns[2], Cm(2))
                            set_column_width(tableRech.columns[3], Cm(4))
                            set_column_width(tableRech.columns[4], Cm(2))
                            set_column_width(tableRech.columns[5], Cm(4))


                            heading_row = tableRech.rows[0].cells



                            lager = rechGrunddatensatz.get('Lager')


                            heading_row[0].text = "Lager:"

                            heading_row[1].text = str(lager)

                            heading_row[2].text = "Nummer:"

                            heading_row[3].text = str(vrgngsnmmr)


                            heading_row[4].text = "Datum:"

                            heading_row[5].text = str(datumG)



                            paragraph = documentl.add_paragraph('')

                            table = documentl.add_table(rows=positionsB, cols=6, style="Table Grid")




                            set_column_width(table.columns[0], Cm(1))
                            set_column_width(table.columns[1], Cm(2))
                            set_column_width(table.columns[2], Cm(2))
                            set_column_width(table.columns[3], Cm(25))
                            set_column_width(table.columns[4], Cm(6))
                            set_column_width(table.columns[5], Cm(8))

                            heading_row = table.rows[0].cells

                            heading_row[0].text = "Pos"
                            heading_row[1].text = "Menge"
                            heading_row[2].text = "Typ"
                            heading_row[3].text = "Text"
                            heading_row[4].text = "Einzelpr."
                            heading_row[5].text = "Gesamtpr."

                            count = 0

                            for k,v in bookToLieferDict.items():
                                #print(k + str(v[0]) + str(v[1]) + str(v[2]))
                                anzahlJB = int(v[0])
                                preisBue = float(v[1])
                                sumBu = float(v[2])
                                sumBu = round(sumBu, 2)
                                vID = str(v[3])


                                typeRe = cursor.execute("SELECT type FROM books WHERE ID = ?", (k,)).fetchall()

                                for t in typeRe:
                                    for x in t:
                                        typeRe = x


                                #print(count)

                                count += 1

                                data_row = table.rows[count].cells



                                data_row[0].text = str(count)
                                data_row[1].text = str(anzahlJB)
                                data_row[2].text = str(typeRe)

                                open_db()
                                kN = cursor.execute("SELECT name_of_item FROM books WHERE ID = ?", (k,)).fetchall()
                                kN = str(kN)
                                for char in ['(', ')', ',', '\'', ']', '[']:

                                    if char in kN:

                                        kN = kN.replace(char, '')

                                data_row[3].text = str(kN)



                                preisBue = f'{preisBue:.2f}'
                                preisBue = str(preisBue)
                                preisBue = preisBue.replace('.',',')
                                preisBue = preisBue + " Euro"

                                data_row[4].text = str(preisBue)
                                data_row[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                                sumBu = f'{sumBu:.2f}'
                                sumBu = str(sumBu)
                                sumBu = sumBu.replace('.',',')

                                sumBu = (re.sub(r'(?<!^)(?=(\d{3})+,)', r'.', sumBu))

                                sumBu = sumBu + " Euro"


                                data_row[5].text = str(sumBu)
                                data_row[5].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            if porto != 0:

                                paragraph = documentl.add_paragraph('')

                                table21 = documentl.add_table(rows=1, cols=2, style="Table Grid")

                                set_column_width(table21.columns[0], Cm(2.5))
                                set_column_width(table21.columns[1], Cm(2.5))

                                heading_row2 = table21.rows[0].cells

                                heading_row2[0].text = "Portopauschale:"


                                portoAD = rechGrunddatensatz.get('Portopauschale')
                                #print('portoAD')
                                #print(portoAD)
                                portoAD = format(portoAD, '.2f')
                                portoAD = str(portoAD)

                                portoAD = portoAD.replace('.',',')
                                portoAD = portoAD + " Euro"

                                heading_row2[1].text = str(portoAD)
                                heading_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            paragraph = documentl.add_paragraph('')

                            table2 = documentl.add_table(rows=1, cols=2, style="Table Grid")

                            set_column_width(table2.columns[0], Cm(2.5))
                            set_column_width(table2.columns[1], Cm(4))

                            heading_row2 = table2.rows[0].cells

                            heading_row2[0].text = "Gesamtbetrag:"





                            heading_row2[1].text = str(sumBi)
                            heading_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT




                            sumBitDaNa = sumBi.replace(',','_')

                            sumBitDaNa = sumBitDaNa.replace(' ','_')

                            verteilerhDaNa = verteilerG.replace(' ','_')

                            koExc = str(vrgngsnmmr) + "_L_" + sumBitDaNa + "_" + verteilerhDaNa + ".docx"

                            koExc = os.path.join(rechLiDir, koExc)

                            try:
                                documentl.save(koExc)

                            except:

                                random_number = random.randint(1, 10000)

                                vrgngsnmmr = vrgngsnmmr + "_" + str(random_number)

                                koExc = str(vrgngsnmmr) + "_L_" + sumBitDaNa + "_" + verteilerhDaNa + ".docx"

                                koExc = os.path.join(rechLiDir, koExc)

                                documentl.save(koExc)






                            absolutePath = Path(koExc).resolve()
                            os.system(f'start WINWORD.EXE "{absolutePath}"')

                            window1['outBuch3'].update("Der Vorgang " + str(vrgngsnmmr) + " wurde erfolgreich erfasst und in der Datenbank gespeichert. Lieferschein und Rechnung wurden erstellt und unter " + str(rechLiDir) + " gespeichert.", text_color='Green')

                            vrgngsnmmr = ""
                            vorgangsnummer = ""
                            sumBi = 0
                            window1['-CAL-'].update('')
                            verteiler = ""
                            lager = ""
                            window1['choosePreis'].update("")
                            ersteBestellZeileOK = "NOK"
                            zweiteBestellZeileOK = "NOK"
                            rechGrunddatensatz = {}
                            infoStartBA = ""
                            bookToLieferDict = {}
                            window1['outVerteiler'].update("")
                            window1['outLager'].update("")
                            window1['outBuch'].update("")
                            window1['outBuch1'].update("")
                            window1['outBuchRE'].update("")
                            window1['outBuch2'].update("")
                            window1['anzahl'].update("")
                            window1['outBuch123'].update("")
                            ininfoStartBA = ''
                            inrechGrunddatensatz = ''
                            inbüchLi = ''
                            inLager = ''
                            inVerteiler = ''
                            ininfoStartBA = ''
                            inrechGrunddatensatz = ''
                            inbookToLieferDict = ""
                            inverteiler = ''
                            inlager = ''
                            inrechGrunddatensatz = ''
                            verteiler = ''
                            lager = ''
                            rechGrunddatensatz = ''
                            büchLi = ''


                            ersteBestellZeileOK = "NOK"
                            zweiteBestellZeileOK = "NOK"










            if event == "Vorschau" or event == sg.WIN_CLOSED:
                random_number = random.randint(1, 10000)
                date.today()
                dtNw = date.today()
                dtNw = dtNw.strftime("%d_%m_%Y")
                koExc = "Vorschau Bestellung_" + str(random_number) + ".xlsx"
                title = "Vorschau Bestellung"
                koExc = os.path.join(repDirPath, koExc)



                df = pd.DataFrame(bookToLieferDict)




                writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                df.to_excel(writer, sheet_name='Vorschau Bestellung', index=False)

                #workbook  = writer.book
                worksheet = writer.sheets['Vorschau Bestellung']


                #worksheet.write(4, 10, title)


                writer.save()



                wb = openpyxl.load_workbook(koExc)
                sheet = wb['Vorschau Bestellung']

                sheet.delete_rows(5)
                # sheet.insert_rows(idx=0, amount=3)
                # sheet.cell(row=2, column=1).value = title
                # sheet.cell(row=3, column=1).value = lagExc
                # sheet.cell(row=4, column=5).value = "Anzahl:"
                wb.save(koExc)





                absolutePath = Path(koExc).resolve()
                os.system(f'start excel.exe "{absolutePath}"')






            if event == "3. Hinzufügen" or event == sg.WIN_CLOSED:
                if ersteBestellZeileOK != "OK":
                    window1['outBuch'].update("Zuerst Verteiler, Lager, Preiskategorie und Datum auswählen.", text_color='Red')
                else:

                    try:
                        anzahlBücher = values['anzahl']
                        anzahlBücher = int(anzahlBücher)
                        anzahlBüCheck = "OK"

                    except:
                        anzahlBüCheck = "NOK"

                    try:
                        büchLiefer = büchLi
                        if büchLiefer:
                            büchLieferCheck = "OK"
                        else:
                            büchLieferCheck = "NOK"
                            #print("userRechNOK")
                    except:
                        büchLieferCheck = "NOK"


                    if anzahlBüCheck == "OK" and büchLieferCheck == "OK":

                        random_number = random.randint(1, 100000000000)
                        vID = str(random_number)

                        open_db()
                        #print('3374')
                        anzahlDbB = [anzahlDbB[0] for anzahlDbB in cursor.execute("SELECT quantity FROM books Where ID=? ",(büchLi,))]

                        #print(anzahlDbB)
                        if len(anzahlDbB) == 1:

                            for i in anzahlDbB:
                                anzahlDbB = i

                        #a.setdefault("somekey",[]).append("bob")

                        anzahlDbB = int(anzahlDbB)
                        anzahlBücher = int(anzahlBücher)

                        dif = anzahlDbB - anzahlBücher

                        if dif >= 0:

                            preisK = rechGrunddatensatz.get('Preiskategorie')

                            preisER = cursor.execute("SELECT " + preisK + " FROM books Where ID =?",(büchLiefer,)).fetchall()
                            #print(preisER)
                            for t in preisER:
                                for x in t:
                                    preisER = x


                            if kindOfInvoice == "Standard Rechnung":

                                summe = anzahlBücher * preisER
                                summe = round(summe, 2)

                            elif kindOfInvoice == "Innergemeinschaftliche Lieferung":

                                abzug = preisER / 100 * 7
                                preisER = preisER - abzug
                                summe = anzahlBücher * preisER
                                summe = round(summe, 2)


                            bookToLieferDict[büchLiefer] = [anzahlBücher, preisER, summe, vID]

                            #print(bookToLieferDict)

                            for k,v in bookToLieferDict.items():

                                if v[0] == 0:
                                    keyDel = k

                            try:
                                del bookToLieferDict[keyDel]
                                keyDel = ""
                            except:
                                pass

                            if not bookToLieferDict:

                                zweiteBestellZeileOK = "nok"
                            else:
                                zweiteBestellZeileOK = "ok"



                            window1['outBuch2'].update(bookToLieferDict, text_color='Blue')


                        else:

                            window1['outBuch2'].update('Es sind nur ' + str(anzahlDbB) + ' ' +  büchLi + ' Bücher im Lager ' + lager + ' vorhanden. Falls das nicht stimmt muss die Anzahl der Bücher bei Bücher - Lagerbestände geändert werden.', text_color='Red')




                    if büchLieferCheck == "NOK":

                        window1['outBuch2'].update("Ein Buch muss ausgewählt werden.", text_color='Red')


                    elif anzahlBüCheck == "NOK":

                        window1['outBuch2'].update("Bei 2.Anzahl muss eine Zahl eingegeben werden z.B. 108.", text_color='Red')



                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"

                    window1['-CAL-'].update('')
                    verteiler = ""
                    lager = ""
                    window1['choosePreis'].update("")
                    ersteBestellZeileOK = "NOK"
                    zweiteBestellZeileOK = "NOK"
                    rechGrunddatensatz = ""
                    infoStartBA = ""

                    bookToLieferDict = {}
                    custo = []
                    rechNumb = ""
                    connection.close()
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:

                    window1['-CAL-'].update('')
                    window1['outVerteiler'].update("")
                    window1['outLager'].update("")
                    window1['choosePreis'].update("")

                    direct = "versand"
                    window1.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()


            if event == "6. Übernehmen" or event == sg.WIN_CLOSED:


                ersteBestellZeileOK = "NOK"
                zweiteBestellZeileOK = "NOK"
                rechGrunddatensatz = ""
                infoStartBA = ""
                bookToLieferDict = {}
                window1['outBuch'].update("")
                window1['outBuch1'].update("")
                window1['outBuchRE'].update("")
                window1['outBuch2'].update("")
                window1['outPorto'].update("5. Porto")
                window1['outBuch123'].update("")
                ininfoStartBA = ''
                inrechGrunddatensatz = ''
                inbüchLi = ''
                ininfoStartBA = ''
                inrechGrunddatensatz = ''
                rechGrunddatensatz = ''
                window1['anzahl'].update("")
                anzahlBücher = ""
                inbookToLieferDict = ''
                bookToLieferDict = {}


                try:
                    choosePreis = values['choosePreis']
                    if choosePreis != "":
                        choosePreisCheck = "OK"
                        #print(choosePreis)
                    else:
                        choosePreisCheck = "NOK"
                        #print("userPreis1")

                except:
                    choosePreisCheck = "NOK"
                    #print("noprice2")

                try:
                    verteiler = inVerteiler
                    if verteiler:
                        verteilerCheck = "OK"
                    else:
                        verteilerCheck = "NOK"
                        #print("userRechNOK")
                except:
                    verteilerCheck = "NOK"



                try:
                    userzahlZiel = inZahlZielKey
                    userzahlZiel = int(userzahlZiel)
                    if int(userzahlZiel):
                        #print(userzahlZiel)
                        userAUzahlZiel = "OK"
                    else:
                        userAUzahlZiel = "NOK"
                        #print("userAUzahlZiel")
                except:
                    userAUzahlZiel = "NOK"
                    # #print('userAUzahlZiel = "NOK"')




                try:
                    lDate = values['-CAL-']
                    if lDate:
                        lDateCheck = "OK"
                    else:
                        lDateCheck = "NOK"
                except:
                    lDateCheck = "NOK"

                try:
                    lager = inLager
                    if lager:
                        lagerCheck = "OK"
                    else:
                        lagerCheck = "NOK"
                        #print("userbeza")
                except:
                    lagerCheck = "NOK"


                try:
                    porto = values['outPorto']
                    porto = float(porto)
                    if float(porto):
                        userAuPorto = "OK"
                    elif porto == 0:
                        userAuPorto = "OK"

                    else:
                        userAuPorto = "NOK"

                except:
                    userAuPorto = "NOK"




                if verteilerCheck == "OK" and lagerCheck == "OK" and choosePreisCheck == "OK" and lDateCheck == "OK" and userAuPorto == "OK" and userAUzahlZiel == "OK":

                    bookToLieferDict = {}
                    window1['outBuch2'].update(inbookToLieferDict, text_color='Blue')

                    rechGrunddatensatz = {"Verteiler": verteiler, "Lager": lager, "Datum": lDate, "Preiskategorie": choosePreis, "Portopauschale": porto, "Zahlungsziel": userzahlZiel, "kindOfInvoice": kindOfInvoice}

                    window1['outBuch'].update(rechGrunddatensatz, text_color='Blue')

                    infoStartBA = 'Nacheinander beliebig viele Bücher hinzufügen oder deren Anzahl (0 zum löschen) korrigieren oder die EXCEL Bestellung verwenden die dann auch angepasst werden kann.'
                    window1['outBuch1'].update('Nacheinander beliebig viele Bücher hinzufügen oder deren Anzahl (0 zum löschen) korrigieren oder die EXCEL Bestellung verwenden die dann auch angepasst werden kann.')

                    ersteBestellZeileOK = "OK"

                    window1['-CAL-'].update('')
                    verteiler = ""
                    window1['choosePreis'].update("")


                if verteilerCheck == "NOK":

                    window1['outBuch'].update("Der Verteiler muss ausgewählt werden.", text_color='Red')


                elif lagerCheck == "NOK":

                    window1['outBuch'].update("Das Lager muss ausgewählt werden.", text_color='Red')

                elif lagerCheck == "NOK":

                    window1['outBuch'].update("Das Lager muss ausgewählt werden.", text_color='Red')

                elif choosePreisCheck == "NOK":

                    window1['outBuch'].update("Die Preiskategorie muss ausgewählt werden.", text_color='Red')


                elif lDateCheck == "NOK":

                    window1['outBuch'].update("Das Bestelldatum muss ausgewählt werden.", text_color='Red')

                elif userAuPorto == "NOK":

                    window1['outBuch'].update("Die Portopauschale muss eingetragen sein und es muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88. Wenn kein Porto dann 0 eintragen.", text_color='Red')




            if event == "1. Verteiler" or event == sg.WIN_CLOSED:



                direct = "findVerteiler1"


                window1.close()

            if event == "2. Lager" or event == sg.WIN_CLOSED:
                #print("laaager")
                direct = "findLager1"
                window1.close()


            if event == "3. Preis" or event == sg.WIN_CLOSED:
                direct = "findPreis"
                window1.close()


        while direct == "findVerteiler1":
            rechStaIn = ['Standard Rechnung', 'Innergemeinschaftliche Lieferung', 'Gutschrift']
            open_db()
            content = [distributer[0] for distributer in cursor.execute("SELECT Verteiler_ID FROM distributer")]
            content = list(set(content))
            content = sorted(content)

            layout1 = [
                [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text("Verteiler auswählen:")],
                [sg.Combo(content, readonly=True, size=(79, 1), key='nameKey')],
                [sg.Text('')],
                [sg.Text("Zahlungsziel ab Rechnungs- bzw. Lieferscheindatum in Tagen:")],
                [sg.InputText("104", size=(4, 1), key="outZahlZielKey")],
                [sg.Text('')],
                [sg.Combo(rechStaIn, readonly=True, size=(39, 1), default_value='Standard Rechnung', key='kindOfInvoice')],
                [sg.Text('')],
                [sg.Button('Ok', size=(9,1)), sg.Text(key="outPostleitzahl_OrtChooseBuch2")]
        
                
            ]

            window1 = sg.Window("Verteiler finden, Zahlungsziel und Rechnungsart auswählen", layout1)

            while direct == "findVerteiler1":

                event, values = window1.read()

                if event == "Ok" or event == sg.WIN_CLOSED:

                    chooseContent = ""

                    try:
                        verteiler = values['nameKey']
                        outZahlZielKey = values['outZahlZielKey']
                        outZahlZielKey = int(outZahlZielKey)
                        kindOfInvoice = values['kindOfInvoice']

                        if kindOfInvoice == 'Gutschrift':
                            direct = "Gutschrift"
                            #print('Gutschrift1')
                        else:
                            direct = "versand"

                        #print(direct)    
                        window1.close()


                    except:
                        pass

                else:
                    direct = "findRechNumber"
                    window1['outPostleitzahl_OrtChooseBuch2'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')

                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"
                    lineToChange = []
                    connection.close()
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:

                    direct = "versand"
                    window1.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()


        while direct == "Gutschrift":

            
            dtNew = date.today()
            dtNew = dtNew.strftime('%Y-%m-%d')



            open_db()

            
            try:
                inVerteiler = verteiler
            except:
                inVerteiler = ''






            layout1 = [
                [sg.Text("")],
                [sg.Button('Clear', size=(12,1)), sg.Text(size=(9, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text(key="outVerteiler", text_color='Blue')],
                [sg.CalendarButton('1. Datum', target='-CAL-', size=(12, 1), pad=None, key='_CALENDAR_', format=('%Y-%m-%d')), sg.In(dtNew, key='-CAL-', readonly=True, size=(12, 1)), sg.InputText("2. Gutschrift ", size=(12, 1), key="outGutschrit"), sg.Button('3. Gutschrift erstellen', size=(18, 1))],
                [sg.Text(key="outBuch3")]
            ]



            window1 = sg.Window("Gutschrift", layout1, finalize=True)
            window1['outVerteiler'].update(inVerteiler)
           
            window1['outBuch3'].update(inrechGrunddatensatz, text_color='Blue')
        
            

            # Datensatz hinzüfgen

            while direct == "Gutschrift":

                event, values = window1.read()


                if event == "1. Datum" or event == sg.WIN_CLOSED:

                    pass

                    #print("Datum")



                if event == "3. Gutschrift erstellen" or event == sg.WIN_CLOSED:

                    try:
                        gutschrift = values['outGutschrit']
                        gutschrift = float(gutschrift)
                        if float(gutschrift):
                            userAugutschrift = "OK"
                        

                        else:
                            userAugutschrift = "NOK"

                    except:
                        userAugutschrift = "NOK"



                    try:
                        verteilerh = inVerteiler
                        if verteilerh:
                            verteilerCheck = "OK"
                        else:
                            verteilerCheck = "NOK"

                    except:
                        verteilerCheck = "NOK"


                    try:
                        lDate = values['-CAL-']
                        if lDate:
                            lDateCheck = "OK"
                        else:
                            lDateCheck = "NOK"

                    except:
                        lDateCheck = "NOK"




                    if verteilerCheck == "OK" and userAugutschrift == "OK" and lDateCheck == "OK":



                        dtNw = date.today()
                        dtNw = dtNw.strftime("%#d%#m%y")



                        open_db()
                        #vorgaengeVomTag = cursor.execute("SELECT vorgangsnummer FROM vorgang WHERE vorgangsnummer LIKE =?",(dtNw,)).fetchall()
                        vorgaengeVomTag = [vorgaengeVomTag[0] for vorgaengeVomTag in cursor.execute("SELECT vorgangsnummer FROM vorgang WHERE vorgangsnummer LIKE ?",('%'+dtNw+'%',))]
                       #print("mmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmm")
                       #print(vorgaengeVomTag)

                        last = []

                        for strip in vorgaengeVomTag:

                            sep = dtNw

                            stripped = strip.split(sep, 1)[1]
                           #print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
                           #print(stripped)
                            last.append(int(stripped))

                        try:
                            fNumber = int(max(last)) + 1
                        except:
                            fNumber = 1

                        vrgngsnmmr = 'G' + dtNw + str(fNumber)


                        open_db()
                        nameVer = [nameVer[0] for nameVer in cursor.execute("SELECT Name FROM distributer Where Verteiler_ID=? ",(verteilerh,))]

                        if len(nameVer) == 1:

                            for i in nameVer:
                                nameVer = i

                        straVer = [straVer[0] for straVer in cursor.execute("SELECT Straße FROM distributer Where Verteiler_ID=? ",(verteilerh,))]

                        if len(straVer) == 1:

                            for i in straVer:
                                straVer = i



                        ortV = [ortV[0] for ortV in cursor.execute("SELECT Postleitzahl_Ort FROM distributer Where Verteiler_ID=? ",(verteilerh,))]
                          
                        if len(ortV) == 1:

                            for i in ortV:
                                ortV = i      


                        UST_ID = [UST_ID[0] for UST_ID in cursor.execute("SELECT UST_ID FROM distributer Where Verteiler_ID=? ",(verteilerh,))]
                        if len(UST_ID) == 1:

                            for i in UST_ID:
                                UST_ID = i


                        






                        # Gutschrict

                        document = Document()




                        def set_column_width(column, width):
                            for cell in column.cells:
                                cell.width = width

                        paragraph = document.add_paragraph()
                        paragraph.add_run('Musterbuchhandel').font.size = Pt(16)
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)

                        paragraph = document.add_paragraph()
                        paragraph.add_run('Musterbücher').font.size = Pt(12)
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(22)



                        tabletop = document.add_table(rows=1, cols=2, style="Table Grid")



                        set_column_width(tabletop.columns[0], Cm(15))
                        set_column_width(tabletop.columns[1], Cm(15))


                        heading_row = tabletop.rows[0].cells

                        heading_row[0].text = "Mustersadresse\nTel: 0000000\nFax: 000000\nUSt.-ID 0000000\nSteuer-Nr. 000000"

                        heading_row[1].text = "Bankverbindung: \nMusterbank\nMusterbankname\nIBAN: 000000\nSWIFT-BIC: 00000"


                        paragraph = document.add_paragraph()
                        paragraph.add_run('Empfänger:').underline = True
                        paragraph.paragraph_format.space_before = Pt(22)
                        paragraph.paragraph_format.space_after = Pt(4)

                        paragraph = document.add_paragraph(str(verteiler))
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)

                        paragraph = document.add_paragraph(str(straVer))
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)

                        paragraph = document.add_paragraph(str(ortV))
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(22)

                        if UST_ID != None:

                            paragraph = document.add_paragraph("USt.-ID: " + str(UST_ID))
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(22)
                        

                        paragraph = document.add_paragraph()
                        paragraph.add_run('G u t s c h r i f t').bold = True
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(4)



                        tableRech = document.add_table(rows=1, cols=4, style="Table Grid")


                        set_column_width(tableRech.columns[0], Cm(5))
                        set_column_width(tableRech.columns[1], Cm(5))
                        set_column_width(tableRech.columns[2], Cm(5))
                        set_column_width(tableRech.columns[3], Cm(5))


                        heading_row = tableRech.rows[0].cells

                        heading_row[0].text = "Nummer:"

                        heading_row[1].text = str(vrgngsnmmr)


                        heading_row[2].text = "Datum:"

                        heading_row[3].text = str(lDate)


                        paragraph = document.add_paragraph('')

                        table21 = document.add_table(rows=1, cols=2, style="Table Grid")

                        set_column_width(table21.columns[0], Cm(3))
                        set_column_width(table21.columns[1], Cm(4))

                        heading_row2 = table21.rows[0].cells

                        heading_row2[0].text = "Gutschriftsbetrag:"

                        nurBeGutschrift = gutschrift

                       #print(gutschrift)
                        gutschrift = format(gutschrift, '.2f')
                        gutschrift = str(gutschrift)

                        gutschrift = gutschrift.replace('.',',')
                        gutschrift = gutschrift + " Euro"

                        heading_row2[1].text = str(gutschrift)
                        heading_row2[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                        paragraph = document.add_paragraph('')

                        gutschriftDaNa = gutschrift.replace(',','_')

                        gutschriftDaNa = gutschriftDaNa.replace(' ','_')

                        verteilerhDaNa = verteilerh.replace(' ','_')

                        koExc = str(vrgngsnmmr) + "_G_" + gutschriftDaNa +"_" + verteilerhDaNa + ".docx"

                        koExc = os.path.join(rechLiDir, koExc)

                        try:
                            document.save(koExc)

                        except:

                            random_number = random.randint(1, 10000)

                            vrgngsnmmr = vrgngsnmmr + "_" + str(random_number)

                            koExc = str(vrgngsnmmr) + "_G_" + gutschriftDaNa +"_" + verteilerhDaNa + ".docx"
                            

                            koExc = os.path.join(rechLiDir, koExc)

                            document.save(koExc)





                        absolutePath = Path(koExc).resolve()
                        os.system(f'start WINWORD.EXE "{absolutePath}"')












                        window1['outBuch3'].update(vrgngsnmmr + " wurde erfasst und im Rechnungsordner gespeichert.", text_color='Green')


                        window1['-CAL-'].update('')
                        window1['outVerteiler'].update('')
                        window1['outGutschrit'].update('')


                        inVerteiler = ''
                        



                        cursor.execute("INSERT INTO kundenbewegung (Verteiler_ID,  Name, Vorgang, Vorgangsnummer, Datum, Betrag) VALUES (?, ?, ?, ?, ?, ?)", (verteilerh, nameVer, "G", vrgngsnmmr, lDate, nurBeGutschrift))


                        cursor.execute("INSERT INTO vorgang (vorgangsnummer, buchdatensatz, verteilerdatensatz, summe) VALUES (?, ?, ?, ?)", (vrgngsnmmr, verteilerh, "Gutschrift", nurBeGutschrift))



                        connection.commit()
                        connection.close()

               






                    elif verteilerCheck == "NOK":

                        window1['outBuch3'].update("Der Verteiler muss ausgewählt werden.", text_color='Red')


                    elif userAugutschrift == "NOK":

                        window1['outBuch3'].update("Der Gutschriftsbetrag muss eingetragen werden.", text_color='Red')


                    elif lDateCheck == "NOK":

                        window1['outBuch3'].update("Das Gutschriftdatum muss ausgewählt werden.", text_color='Red')





                if event == "Clear" or event == sg.WIN_CLOSED:


                    window1['-CAL-'].update('')

                    window1['outBuch3'].update("")

                    window1['outGutschrit'].update("")

                    
     

               






                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"

                    window1['-CAL-'].update('')
                    verteiler = ""
                    window1['outVerteiler'].update("")
                    window1['outBuch3'].update("")
                    window1.close()

     
                    inVerteiler = ''
               
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:

                    direct = "versand"
                    window1['-CAL-'].update('')
                    verteiler = ""
                    window1['outVerteiler'].update("")
                    window1['outBuch3'].update("")
                    window1.close()

     
                    inVerteiler = ''

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()










        while direct == "bestImport":

            try:

                anzky = values['impKey']

                lagerFrFile = openpyxl.load_workbook(anzky)

                lagerFrFile = lagerFrFile.active

                lagerFrFile = lagerFrFile['A3']

                lagerFrFile = lagerFrFile.value

                #print("warehouse")
                #print(lagerFrFile)



                if lagerFrFile == lager:

                    dictIm = pd.read_excel(anzky, index_col=0, skiprows=4, squeeze=True, header=None, usecols=(0,8))
                    dictIm = dictIm.to_dict()
                    #print(type(dictIm))
                    #print(dictIm)


                    for k,v in dictIm.items():
                        #print(k)


                        if not math.isnan(v):


                            anzahlBücher = int(v)

                            random_number = random.randint(1, 100000000000)
                            vID = str(random_number)

                            # open_db()
                            # anzahlDbB = [anzahlDbB[0] for anzahlDbB in cursor.execute("SELECT quantity FROM books Where ID=? ",(k,))]

                            # #print(anzahlDbB)
                            # if len(anzahlDbB) == 1:

                            #     for i in anzahlDbB:
                            #         anzahlDbB = i



                            # anzahlDbB = int(anzahlDbB)
                            # anzahlBücher = int(anzahlBücher)

                            # #print(anzahlDbB)
                            # #print(anzahlBücher)

                            #         dif = anzahlDbB - anzahlBücher

                            #         if dif >= 0:

                            preisK = rechGrunddatensatz.get('Preiskategorie')

                            preisER = cursor.execute("SELECT " + preisK + " FROM books Where ID =?",(k,)).fetchall()
                            #print(preisER)
                            for t in preisER:
                                for x in t:
                                    preisER = x

                                summe = anzahlBücher * preisER


                            bookToLieferDict[k] = [anzahlBücher, preisER, summe, vID]

                        büchLieferCheck = "OK"
                        zweiteBestellZeileOK = "ok"

                        imporText = 'Die Bestellung wurde in das System importiert und kann unterhalb bearbeitet werden oder es kann gleich "Lieferschein und Rechnung erstellen" ausgewählt werden.'

                else:
                    imporText = 'Das Lager das oberhalb in Gauranga ausgewählt wurde (' + lager + ') stimmt nicht mit dem Lager das im Importfile steht (' + lagerFrFile + ') überein. Bitte in Gauranga oder im Importfile das Lager anpassen und den Vorgang wiederholen.'
            except:
                imporText = 'Entweder es wurde keine Bestellung in das System importiert. Bitte zuerst "1. EXCEL Bestellung auswählen" oder ggf. den aktiven Import fertig bearbeiten.'





            direct = "versand"
            window1.close()







             # maha man
             #
        while direct == "vorlageBestell":


            #print("lololololoololoooo44444444444444444444444444444444444")
            open_db()
            content = [books[0] for books in cursor.execute("SELECT warehouse FROM books")]
            #print("go")
            content = list(set(content))
            #print(content)
            content = sorted(content)
            layout1 = [
                [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text("Lager auswählen:")],
                [sg.Combo(content, readonly=True, size=(79, 1), key='nameLKey')],
                [sg.Button('Ok', size=(9,1)), sg.Text(key="outLagerCSV")],
            ]

            window1 = sg.Window("Lager finden", layout1)

            while direct == "vorlageBestell":

                event, values = window1.read()

                if event == "Ok" or event == sg.WIN_CLOSED:

                    chooseContent = ""

                    try:
                        nameLa = values['nameLKey']
                        if nameLa != "":



                            random_number = random.randint(1, 10000)
                            date.today()
                            dtNw = date.today()
                            dtNw = dtNw.strftime("%d_%m_%Y")
                            koExc = "Bestellung_Bücher_" + nameLa + "_" + str(random_number) + ".xlsx"
                            title = "Bestellung_Bücher"
                            lagExc = nameLa
                            koExc = os.path.join(repDirPath, koExc)

                            imporText = "Die Datei " + koExc + " wurde erstellt und wird automatisiert aus dem Verzeichnis gelöscht und kann dem entsprechenen Verteiler zur Verfügung gestellt werden. Das Verteiler muss in die Spalte Anzahl die Anzahl der jweiligen Bücher schreiben und darf ansonsten nichts verändern. Dann kann die Datei in das System importiert werden."




                                                # rows = cursor.execute("SELECT * FROM books WHERE type = ? AND name_of_item = ? AND warehouse = ?", ##(typeDddd, name_of_itemDddd, warehouseDddd,)).fetchall()

                            # cursor.execute("UPDATE kunden_saldo SET kunden_saldo=? WHERE warehouse=? ", (nameLa))
                            open_db()
                            #content = [distributer[0] for distributer in cursor.execute("SELECT Name FROM distributer")]
                            rows = cursor.execute("SELECT ID, name_of_item, type, language, SP, P1, P2, P3_30 FROM books WHERE warehouse = ? ORDER BY language, name_of_item", (nameLa,)).fetchall()
                            colNames = cursor.execute("SELECT ID, name_of_item, type, language, SP, P1, P2, P3_30 FROM books")
                            colNames = [cn[0] for cn in colNames.description]
                            colNames = tuple(colNames)
                            rows.insert(0, colNames)
                            tableBooks = rows

                            df = pd.DataFrame(tableBooks)




                            writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                            df.to_excel(writer, sheet_name='Bestellung Bücher', index=False)

                            #workbook  = writer.book
                            worksheet = writer.sheets['Bestellung Bücher']


                            #worksheet.write(4, 10, title)


                            writer.save()



                            wb = openpyxl.load_workbook(koExc)
                            sheet = wb['Bestellung Bücher']

                            sheet.delete_rows(1)
                            sheet.insert_rows(idx=0, amount=3)
                            sheet.cell(row=2, column=1).value = title
                            sheet.cell(row=3, column=1).value = lagExc
                            sheet.cell(row=4, column=9).value = "Anzahl:"
                            wb.save(koExc)





                            absolutePath = Path(koExc).resolve()
                            os.system(f'start excel.exe "{absolutePath}"')

                            connection.close()

                            direct = "versand"
                            window1.close()


                    except:
                        pass

                else:
                    direct = "findLager"
                    window1['outLagerCSV'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')


                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"
                    lineToChange = []
                    connection.close()
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:

                    direct = "versand"
                    window1.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()




        while direct == "bücherZuLieferschein":
            open_db()
            content = [books[0] for books in cursor.execute("SELECT name_of_item FROM books Where warehouse =?",(lager,))]


            #print(content)
            content = sorted(content)


            #print("2557")
            #print(lager)
            layout1 = [
                [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text("Buch auswählen:")],
                [sg.Combo(content, readonly=True, size=(79, 1), key='nameKey')],





                [sg.Button('Ok', size=(9,1)), sg.Text(key="outPostleitzahl_OrtChooseBuch2")],
            ]



            window1 = sg.Window("Buch finden", layout1)

            while direct == "bücherZuLieferschein":

                event, values = window1.read()

                if event == "Ok" or event == sg.WIN_CLOSED:

                    chooseContent = ""

                    try:
                        büchLi = values['nameKey']

                        büchLi = cursor.execute("SELECT ID FROM books WHERE name_of_item = ?", (büchLi,)).fetchall()

                        büchLi = str(büchLi)

                        for char in ['(', ')', ',', '\'', ']', '[']:
                            if char in büchLi:

                                büchLi = büchLi.replace(char, '')


                        direct = "versand"
                        window1.close()


                    except:
                        pass

                else:
                    direct = "findRechNumber"
                    window1['outPostleitzahl_OrtChooseBuch2'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')


                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"
                    lineToChange = []
                    connection.close()
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:

                    direct = "versand"
                    window1.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()



        if event == "Exit" or event == sg.WIN_CLOSED:
            db_backup()
            direct = "exit"
            window1.close()




        if event == "Gauranga" or event == sg.WIN_CLOSED:
            direct = "menu"
            lineToChange = []
            window1.close()









        while direct == "findLager1":
            open_db()
            content = [books[0] for books in cursor.execute("SELECT warehouse FROM books")]
            content = list(set(content))
            content = sorted(content)

            layout1 = [
                [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text("Lager auswählen:")],
                [sg.Combo(content, readonly=True, size=(79, 1), key='nameKey')],
                [sg.Button('Ok', size=(9,1)), sg.Text(key="outPostleitzahl_OrtChooseBuch2")],
            ]

            window1 = sg.Window("Lager finden", layout1)

            while direct == "findLager1":

                event, values = window1.read()

                if event == "Ok" or event == sg.WIN_CLOSED:

                    chooseContent = ""

                    try:
                        lager = values['nameKey']
                        direct = "versand"
                        window1.close()


                    except:
                        pass

                else:
                    direct = "findRechNumber"
                    window1['outPostleitzahl_OrtChooseBuch2'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')


                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"
                    lineToChange = []
                    connection.close()
                    window1.close()
                    window1['outBuch123'].update("")

                if event == "Back" or event == sg.WIN_CLOSED:

                    direct = "versand"
                    window1.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()



        if event == "Exit" or event == sg.WIN_CLOSED:
            db_backup()
            direct = "exit"
            window1.close()




        if event == "Gauranga" or event == sg.WIN_CLOSED:
            direct = "menu"
            lineToChange = []
            window1.close()




    if event == "Kunden" or event == sg.WIN_CLOSED:

        direct = "Kunden"
        window0.close()

    if event == "change" or event == sg.WIN_CLOSED:

        direct = "change"
        window0.close()

    if event == "Exit" or event == sg.WIN_CLOSED:
        db_backup()
        direct = "exit"
        window0.close()
        break




    while direct == "Kunden":

        open_db()


        try:
            inVerteiler_IDDistri = lineToChange[0]
            inEmailDistri = lineToChange[5]
            inLandDistri = lineToChange[4]
            inNameDistri = lineToChange[1]
            inStraßeDistri = lineToChange[2]
            inPostleitzahl_OrtDistri = lineToChange[3]
            #inOrtDistri = lineToChange[7]
            inTelfonnummerDistri = lineToChange[6]
            inUST_ID = lineToChange[7]

        except:

            inVerteiler_IDDistri = None
            inEmailDistri = None
            inNameDistri = None
            inStraßeDistri = None
            inLandDistri = None
            inPostleitzahl_OrtDistri = None
            inOrtDistri = None
            inTelfonnummerDistri = None
            inUST_ID = None



#Create distributor  table
#cursor.execute("CREATE TABLE distributer (Verteiler_ID  Name  Straße Postleitzahl_Ort Email Telefonnummer


        layout1 = [
            [sg.Text("", size=(154, 1))],
            [sg.Button('Datensatz finden', size=(13,1)), sg.Button('Clear', size=(13,1)), sg.Text('', size=(102, 1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
            [sg.Text("")],
            [sg.Text('Verteiler_ID:', size=(13, 1)), sg.InputText(inVerteiler_IDDistri, size=(41, 1), key='Verteiler_ID'), sg.Text('Name:', size=(13, 1)), sg.InputText(inNameDistri, size=(41, 1), key="Name"), sg.Text('Straße:', size=(13, 1)), sg.InputText(inStraßeDistri, size=(41, 1),key="Straße")],
            [sg.Text(size=(50, 1), key="outVerteiler_IDDistri"), sg.Text(size=(51, 1), key="outNameDistri"), sg.Text(size=(51, 1), key="outStraßeDistri")],
            [sg.Text("")],
            [sg.Text('Postleitzahl_Ort:', size=(13, 1)), sg.InputText(inPostleitzahl_OrtDistri, size=(41, 1), key="Postleitzahl_Ort"), sg.Text('Land:', size=(13, 1)), sg.InputText(inLandDistri, size=(41, 1), key="Land")],
            [sg.Text(size=(51, 1), key="outPostleitzahl_OrtDistri"), sg.Text(size=(50, 1), key="outLandDistri")],
            [sg.Text("")],
            [sg.Text('Email:', size=(13, 1)), sg.InputText(inEmailDistri, size=(41, 1), key="Email"), sg.Text('Telefonnummer:', size=(13, 1)), sg.InputText(inTelfonnummerDistri, size=(41, 1), key="Telefonnummer")],
            [sg.Text(size=(50, 1), key="outEmailDistri"), sg.Text(size=(51, 1), key="outTelefonnummerDistri")],
            [sg.Text("")],
            [sg.Text('USt.-ID:', size=(13, 1)), sg.InputText(inUST_ID, size=(41, 1), key="UST_ID")],
            [sg.Text(size=(50, 1), key="outUST_ID")],
            [sg.Text("")],
            [sg.Button('Neuen Datensatz anlegen', size=(37,1)), sg.Button('Bestehenden Datensatz über Verteiler_ID ändern', size=(37,1)), sg.Button('Lösche Datensatz über Verteiler_ID', size=(37,1))],
            [sg.Text(key="outBuch")],
        ]


        window1 = sg.Window("Verteiler - Verteilerdatensatz ändern, hinzufügen oder löschen", layout1)

        # Datensatz hinzüfgen

        while direct == "Kunden":

            event, values = window1.read()


            if event == "Lösche Datensatz über Verteiler_ID" or event == sg.WIN_CLOSED:

                try:
                    window1['outVerteiler_IDDistri'].update('')
                    window1['outEmailDistri'].update('')
                    window1['outLandDistri'].update('')
                    window1['outNameDistri'].update('')
                    window1['outStraßeDistri'].update('')
                    window1['outPostleitzahl_OrtDistri'].update('')
                    window1['outTelefonnummerDistri'].update('')
                    window1['outUST_ID'].update('')


                except:
                    pass


                try:
                    Verteiler_ID = values['Verteiler_ID']
                    Verteiler_IDs = [distributer[0] for distributer in cursor.execute("SELECT Verteiler_ID FROM distributer")]
                    if Verteiler_ID in Verteiler_IDs:
                        Verteiler_ID = str(Verteiler_ID)
                        cursor.execute("DELETE FROM distributer WHERE Verteiler_ID=? ", (Verteiler_ID,))
                        connection.commit()

                        window1['outBuch'].update('Der Datensatz mit der Verteiler_ID "' + Verteiler_ID + '" wurde in der Datenbank "' + dbName + '" gelöscht.', text_color='Green')
                        clear_Ditri()
                        connection.close()

                    else:
                        window1['outVerteiler_IDDistri'].update("Verteiler_ID muss eine vorhandene Verteiler_ID sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht gelöscht, bitte Fehlermeldungen beachten.", text_color='Red')
                except:
                        window1['outVerteiler_IDDistri'].update("Verteiler_ID muss eine vorhandende Verteiler_ID sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht gelöscht, bitte Fehlermeldungen beachten.", text_color='Red')




            if event == "Neuen Datensatz anlegen" or event == sg.WIN_CLOSED:
                #window1['test'].update("P3_30 test.")
                countBuch = 0



                try:
                    Verteiler_ID = values['Verteiler_ID']
                    #print(Verteiler_ID)
                    open_db()
                    Verteiler_IDs = [distributer[0] for distributer in cursor.execute("SELECT Verteiler_ID FROM distributer")]
                    if Verteiler_ID and Verteiler_ID not in Verteiler_IDs:
                        window1['outVerteiler_IDDistri'].update("Verteiler_ID ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outVerteiler_IDDistri'].update("Verteiler_ID darf nicht leer sein und darf nicht vorhanden sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')
                except:
                    window1['outVerteiler_IDDistri'].update("Verteiler_ID darf nicht leer sein und darf nicht vorhanden sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')





                try:
                    Email = values['Email']
                except:
                    pass

                try:
                    Land = values['Land']
                except:
                    pass

                try:
                    Name = values['Name']
                    if Name:
                        window1['outNameDistri'].update("Name ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outNameDistri'].update("Name darf nicht leer sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                except:
                    window1['outNameDistri'].update("Name darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    Straße = values['Straße']

                except:
                    pass

                try:
                    Postleitzahl_Ort = values['Postleitzahl_Ort']

                except:
                    pass




                try:
                    Telefonnummer = values['Telefonnummer']
                except:
                    pass


                try:
                    UST_ID = values['UST_ID']
                except:
                    pass







                if countBuch == 2:

                    #print("Es wird in die DB geschrieben")
                    try:

#cursor.execute("CREATE TABLE distributer (Verteiler_ID TEXT PRIMARY KEY, Name TEXT, Straße TEXT, Postleitzahl_Ort TEXT, Land TEXT, Email TEXT, Telefonnummer TEXT)")


                        cursor.execute("INSERT INTO distributer (Verteiler_ID, Name, Straße, Postleitzahl_Ort, Land, Email, Telefonnummer, UST_ID) VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (Verteiler_ID, Name, Straße, Postleitzahl_Ort, Land, Email, Telefonnummer, UST_ID))

                        connection.commit()
                        connection.close()


                        window1['outBuch'].update('Der Datensatz mit der Verteiler_ID "' + Verteiler_ID + '" wurde in der Datenbank "' + dbName + '" gespeichert.', text_color='Green')


                        clear_Ditri()

                    except:

                        window1['outVerteiler_IDDistri'].update("Diese Verteiler_ID existiert schon.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')







            if event == 'Clear' or event == sg.WIN_CLOSED:


                window1['outBuch'].update('')



                clear_Ditri()












            if event == 'Bestehenden Datensatz über Verteiler_ID ändern' or event == sg.WIN_CLOSED:

                countBuch = 0

                try:
                    Verteiler_ID = values['Verteiler_ID']
                    Verteiler_IDs = [distributer[0] for distributer in cursor.execute("SELECT Verteiler_ID FROM distributer")]
                    if Verteiler_ID and Verteiler_ID in Verteiler_IDs:
                        window1['outVerteiler_IDDistri'].update("Verteiler_ID ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outVerteiler_IDDistri'].update("Verteiler_ID darf nicht leer sein und muss eine vorhandene Verteiler_ID sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')
                except:
                    window1['outVerteiler_IDDistri'].update("Verteiler_ID darf nicht leer sein und muss eine vorhandene Verteiler_ID sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')




                try:
                    Email = values['Email']
                except:
                    pass


                try:
                    Land = values['Land']

                except:
                    pass

                try:
                    Name = values['Name']
                    if Name:
                        window1['outNameDistri'].update("Name ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outNameDistri'].update("Name darf nicht leer sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                except:
                    window1['outNameDistri'].update("Name darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    Straße = values['Straße']

                except:
                    pass

                try:
                    Postleitzahl_Ort = values['Postleitzahl_Ort']


                except:
                    pass

                try:
                    Telefonnummer = values['Telefonnummer']
                except:
                    pass


                try:
                    UST_ID = values['UST_ID']
                except:
                    pass


                if countBuch == 2:

                   # try:

                    cursor.execute("UPDATE distributer SET Postleitzahl_Ort=?, Land=?, Name=?, Straße=?, Telefonnummer=?, Email=?, UST_ID=? WHERE Verteiler_ID=? ", (Postleitzahl_Ort, Land, Name, Straße, Telefonnummer, Email, UST_ID, Verteiler_ID))

                    connection.commit()
                    connection.close()


                    window1['outBuch'].update('Der Datensatz mit der Verteiler_ID "' + Verteiler_ID + '" wurde in der Datenbank "' + dbName + '" gespeichert.', text_color='Green')


                    clear_Ditri()

                   # except:

                  #      window1['outVerteiler_IDDistri'].update("Diese ID existiert nicht.", text_color='Red')
                  #      window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')














            if event == "Exit" or event == sg.WIN_CLOSED:
                db_backup()
                direct = "exit"
                connection.close()
                window1.close()




            if event == "Gauranga" or event == sg.WIN_CLOSED:
                direct = "menu"
                connection.close()
                lineToChange = []
                window1.close()

            if event == "Datensatz finden" or event == sg.WIN_CLOSED:
                direct = "change"
                window1.close()

        while direct == "change":
            open_db()
            # create dynamic drop down for table culums
            colNames = cursor.execute("SELECT * FROM distributer")
            colNames = [cn[0] for cn in colNames.description]


            layout1 = [
                [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text("Spalte auswählen:")],
                [sg.Combo(colNames, readonly=True, size=(79, 1), key='chooseColumn')],
                [sg.Button('Ok', size=(9,1)), sg.Text(key="outPostleitzahl_OrtChooseBuch1")],


            ]

            window1 = sg.Window("Verteilerdatensatz auswählen", layout1)

            while direct == "change":

                event, values = window1.read()

                if event == "Ok" or event == sg.WIN_CLOSED:

                    chooseColumn = ""

                    try:
                        chooseColumn = values['chooseColumn']

                    except:
                        pass

                    if chooseColumn != "":

                        direct = "content"
                        window1.close()

                    else:
                        window1['outPostleitzahl_OrtChooseBuch1'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')
                        direct = "change"
                        #window1.close()







                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"
                    connection.close()
                    lineToChange = []
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:
                    pass

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()


            while direct == "content":

                content = [distributer[0] for distributer in cursor.execute("SELECT " + chooseColumn + " FROM distributer")]
                content = sorted(content)

                layout1 = [
                    [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                    [sg.Text("Inhalt auswählen:")],
                    [sg.Combo(content, readonly=True, size=(79, 1), key='chooseContent')],
                    [sg.Button('Ok', size=(9,1)), sg.Text(key="outPostleitzahl_OrtChooseBuch2")],
                ]

                window1 = sg.Window("Verteilerdatensatz auswählen", layout1)

                while direct == "content":

                    event, values = window1.read()

                    if event == "Ok" or event == sg.WIN_CLOSED:

                        chooseContent = ""

                        try:
                            chooseContent = values['chooseContent']

                        except:
                            pass

                        if chooseContent != "":
                            lines = [distributer[0] for distributer in cursor.execute("SELECT * FROM distributer Where " + chooseColumn + "=?",(chooseContent,))]
                            #lines = cursor.execute("SELECT * FROM distributer Where " + chooseColumn + "=?",(chooseContent,))
                            #lines = lines.fetchall()

                            if len(lines) == 1:

                                for i in lines:
                                    Verteiler_ID = i

                                linVerteiler_ID = cursor.execute("SELECT * FROM distributer Where Verteiler_ID=?",(Verteiler_ID,))
                                linVerteiler_ID = linVerteiler_ID.fetchall()

                                result = []

                                for t in linVerteiler_ID:
                                    for x in t:
                                        result.append(x)

                                lineToChange = result
                                window1.close()
                                direct = "Kunden"
                                break

                            else:
                                direct = "line"
                                window1.close()


                        else:
                            direct = "content"
                            window1['outPostleitzahl_OrtChooseBuch2'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')


                    if event == "Gauranga" or event == sg.WIN_CLOSED:
                        direct = "menu"
                        lineToChange = []
                        connection.close()
                        window1.close()

                    if event == "Back" or event == sg.WIN_CLOSED:

                        direct = "Kunden"
                        window1.close()

                    if event == "Exit" or event == sg.WIN_CLOSED:
                        db_backup()
                        direct = "exit"
                        connection.close()
                        window1.close()





            while direct == "line":

                lines = cursor.execute("SELECT * FROM distributer Where " + chooseColumn + "=?",(chooseContent,))
                lines = lines.fetchall()
                lines = list(set(lines))
                ##print(lines)

                layout1 = [
                    [sg.Text('', size=(90, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                    [sg.Text("Zeile auswählen:")],
                    [sg.Combo(lines, readonly=True, size=(140, 1), key='chooseLines')],
                    [sg.Button('Ok', size=(9,1)), sg.Text(key="outPostleitzahl_OrtChooseBuch3")],
                ]

                window1 = sg.Window("Buchdatensatz auswählen", layout1)

                while direct == "line":

                    event, values = window1.read()

                    if event == "Ok" or event == sg.WIN_CLOSED:

                        chooseLines = ""

                        try:
                            chooseLines = values['chooseLines']

                        except:
                            pass

                        if chooseLines != "":

                            chooseline = list(chooseLines)
                            lineToChange = chooseline

                            window1.close()
                            direct = "Kunden"
                            break

                        else:
                            direct = "line"
                            window1['outPostleitzahl_OrtChooseBuch3'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')


                    if event == "Gauranga" or event == sg.WIN_CLOSED:
                        direct = "menu"
                        lineToChange = []
                        connection.close()
                        window1.close()

                    if event == "Back" or event == sg.WIN_CLOSED:

                        direct = "Kunden"
                        window1.close()

                    if event == "Exit" or event == sg.WIN_CLOSED:
                        db_backup()
                        direct = "exit"
                        connection.close()
                        window1.close()










    while direct == "add":

        open_db()

        try:
            inIDBuch = lineToChange[0]
            inquantityBuch = lineToChange[6]
            inBestandBuch = lineToChange[6]
            inNameOfItemBuch = lineToChange[1]
            inLanguageBuch = lineToChange[4]
            inwarehouseBuch = lineToChange[5]
            intypeBuch = lineToChange[3]
            inSPBuch = lineToChange[9]
            inBBT_priceBuch = lineToChange[8]
            inP1Buch = lineToChange[10]
            inP2Buch = lineToChange[11]
            inP3_40Buch = lineToChange[12]
            inP3_30Buch = lineToChange[13]
            inEnd_PBuch = lineToChange[14]
            inInventoName = lineToChange[17]
            InventoNameBuch = lineToChange[17]
            DEinLaPreis = lineToChange[21]
            La_Preis = lineToChange[21]
            inPreiName = lineToChange[20]
            PreiNameBuch = lineToChange[20]






        except:

            inIDBuch = ""
            inNameOfItemBuch = "neu eintragen nur wenn nicht im dropdown"
            inLanguageBuch = "neu eintragen nur wenn nicht im dropdown"
            inwarehouseBuch = "neu eintragen nur wenn nicht im dropdown"
            intypeBuch = "neu eintragen nur wenn nicht im dropdown"
            InventoNameBuch = "neu eintragen nur wenn nicht im dropdown"
            DEinLaPreis = "neu eintragen nur wenn nicht im dropdown"
            PreiNameBuch = "neu eintragen nur wenn nicht im dropdown"
            inBestandBuch = None
            inPreiName = None
            inquantityBuch = None
            inSPBuch = None
            inBBT_priceBuch = None
            inP1Buch = None
            inP2Buch = None
            inP3_40Buch = None
            inP3_30Buch = None
            inEnd_PBuch = None
            inInventoName = None
            inPreiName = None
  
            



        warehouseDd = [books[0] for books in cursor.execute("SELECT warehouse FROM books")]
        warehouseDd = list(set(warehouseDd))
        warehouseDd = sorted(warehouseDd)

        languageDd = [books[0] for books in cursor.execute("SELECT language FROM books")]
        languageDd = list(set(languageDd))
        languageDd = sorted(languageDd)

        typeDd = [books[0] for books in cursor.execute("SELECT type FROM books")]
        typeDd = list(set(typeDd))
        typeDd = sorted(typeDd)

        name_of_itemDd = [books[0] for books in cursor.execute("SELECT name_of_item FROM books")]
        name_of_itemDd = list(set(name_of_itemDd))
        name_of_itemDd = sorted(name_of_itemDd)

        inInventoName = [books[0] for books in cursor.execute("SELECT name FROM books")]
        inInventoName= list(set(inInventoName))
        inInventoName = sorted(inInventoName)

        inLaPreis = [books[0] for books in cursor.execute("SELECT warehouses FROM books")]
        inLaPreis = list(set(inLaPreis))
        inLaPreis = sorted(inLaPreis)

        inPreiName = [books[0] for books in cursor.execute("SELECT name_of_article FROM books")]
        inPreiName = list(set(inPreiName))
        inPreiName = sorted(inPreiName)



        layout1 = [
            [sg.Text("", size=(154, 1))],
            [sg.Button('Datensatz finden', size=(13,1)), sg.Button('Clear', size=(13,1)), sg.Text('', size=(102, 1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
            [sg.Text("")],
            [sg.Text('ID: ', size=(13, 1)), sg.InputText(inIDBuch, size=(41, 1), key="ID"), sg.Text('Name:', size=(13, 1)), sg.Combo(name_of_itemDd, size=(39, 1), default_value=inNameOfItemBuch, key="name_of_item"), sg.Text('Sprache:', size=(13, 1)), sg.Combo(languageDd, size=(39, 1), default_value=inLanguageBuch, key='language')],
            [sg.Text(size=(50, 1), key="outIDBuch"), sg.Text(size=(51, 1), key="outNameOfItemBuch"),sg.Text(size=(51, 1), key="outLanguageBuch")],
            [sg.Text("")],
            [sg.Text('Lager:', size=(13, 1)), sg.Combo(warehouseDd, size=(39, 1), default_value=inwarehouseBuch, key='warehouse'), sg.Text('Typ:', size=(13, 1)), sg.Combo(typeDd, size=(39, 1), default_value=intypeBuch, key="type"), sg.Text('BBT_price:', size=(13, 1)), sg.InputText(inBBT_priceBuch, size=(41, 1), key="BBT_price")],
            [sg.Text(size=(50, 1), key="outwarehouseBuch"), sg.Text(size=(50, 1), key="outtypeBuch"), sg.Text(size=(50, 1),key="outBBT_priceBuch")],
            [sg.Text("")],
            [sg.Text('SP:', size=(13, 1)), sg.InputText(inSPBuch, size=(41, 1), key="SP"), sg.Text('P1:', size=(13, 1)), sg.InputText(inP1Buch, size=(41, 1), key="P1"), sg.Text('P2:', size=(13, 1)), sg.InputText(inP2Buch, size=(41, 1), key="P2")],
            [sg.Text(size=(50, 1), key="outSPBuch"), sg.Text(size=(51, 1), key="outP1Buch"), sg.Text(size=(51, 1), key="outP2Buch")],
            [sg.Text("")],
            [sg.Text('P3_40:', size=(13, 1)), sg.InputText(inP3_40Buch, size=(41, 1), key="P3_40"), sg.Text('P3_30:', size=(13, 1)), sg.InputText(inP3_30Buch, size=(41, 1), key="P3_30"), sg.Text('End_P:', size=(13, 1)), sg.InputText(inEnd_PBuch, size=(41, 1), key="End_P")],
            [sg.Text(size=(50, 1), key="outP3_40Buch"), sg.Text(size=(51, 1), key="outP3_30Buch"), sg.Text(size=(51, 1), key="outEnd_PBuch")],
            [sg.Text("")],
            [sg.Text('Inventory Name:', size=(13, 1)), sg.Combo(inInventoName, size=(39,1), default_value=InventoNameBuch, key="Inventory_Name"), sg.Text('Lager Preisliste:', size=(13, 1)), sg.Combo(inLaPreis, size=(39,1), default_value=DEinLaPreis, key="La_Preis"), sg.Text('Name Preisliste:', size=(13, 1)), sg.Combo(inPreiName, size=(39,1), default_value=PreiNameBuch, key="Preis_Liste")],
            [sg.Text(size=(50, 1), key="outInventory_Name"), sg.Text(size=(50, 1), key="outLa_Preis"), sg.Text(size=(50, 1), key="outName_Preis")],
            [sg.Text("")],
            [sg.Text('Bestand:', size=(13, 1)), sg.InputText(inBestandBuch, size=(41, 1), key="Bestand"), sg.Text("", size=(15, 1)), sg.Text('Wenn Inventory Name, Lager Preisliste oder Name Preisliste nicht aufgeführt werden sollen, bitte jeweils 0 eintragen.', text_color='Blue')],
            [sg.Text(size=(50, 1), key="outBestandBuch")],
            [sg.Button('Neuen Datensatz anlegen', size=(29,1)), sg.Button('Bestehenden Datensatz über ID ändern', size=(29,1)), sg.Button('Lösche Datensatz über ID', size=(29,1)), sg.Button('Lagerbestände ändern', size=(29,1)), sg.Button('Alle Preise von allen Büchern updaten', size=(29,1))],
            [sg.Text(key="outBuch")],
        ]




        window1 = sg.Window("Bücher - Buchdatensatz ändern, hinzufügen oder löschen", layout1, finalize=True)
        #if inIDBuch:
            #window1['anzahlBu'].update("Anzahl " + inIDBuch + ": " + str(inquantityBuch), text_color='Blue')
            # window1['outPostleitzahl_OrtChooseBuch3'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')

        try:
            if anzahlB == 0:

                window1['outBuch'].update("Es existriert kein Buchdatensatz von daher kann die Bücheranzahl nicht geändert werden. Bitte zuerst neuen Buchdatensatz anlegen und dann die Anzahl ändern.", text_color='Red')
                anzahlB = 1
        except:
            pass
        # Datensatz hinzüfgen

        while direct == "add":

            event, values = window1.read()


            if event == "Lösche Datensatz über ID" or event == sg.WIN_CLOSED:

                try:
                    window1['outLanguageBuch'].update('')
                    window1['outwarehouseBuch'].update('')
                    window1['outNameOfItemBuch'].update('')
                    window1['outtypeBuch'].update('')
                    window1['outBBT_priceBuch'].update('')
                    window1['outSPBuch'].update('')
                    window1['outP1Buch'].update('')
                    window1['outP2Buch'].update('')
                    window1['outP3_40Buch'].update('')
                    window1['outP3_30Buch'].update('')
                    window1['outEnd_PBuch'].update('')
                    window1['outInventory_Name'].update('')
                    window1['outBestandBuch'].update('')
                    window1['outLa_Preis'].update('')
                    window1['outName_Preis'].update('')


                except:
                    pass


                try:
                    open_db()
                    ID = values['ID']
                    ids = [books[0] for books in cursor.execute("SELECT ID FROM books")]
                    if ID in ids:
                        ID = str(ID)
                        cursor.execute("DELETE FROM books WHERE ID=? ", (ID,))
                        connection.commit()

                        window1['outBuch'].update('Der Datensatz mit der ID "' + ID + '" wurde in der Datenbank "' + dbName + '" gelöscht.', text_color='Green')
                        clear_Buch()
                        connection.close()

                    else:
                        window1['outIDBuch'].update("ID muss eine vorhandene ID sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht gelöscht, bitte Fehlermeldungen beachten.", text_color='Red')
                except:
                        window1['outIDBuch'].update("ID muss eine vorhandende ID sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht gelöscht, bitte Fehlermeldungen beachten.", text_color='Red')

            if event == "Alle Preise von allen Büchern updaten" or event == sg.WIN_CLOSED:
                #window1['test'].update("P3_30 test.")
                countBuch = 0
                clear_Buch()
                lineToChange = []
                direct = "preiChan"
                window1.close()




            if event == "Lagerbestände ändern" or event == sg.WIN_CLOSED:
                #window1['test'].update("P3_30 test.")
                countBuch = 0
                clear_Buch()
                lineToChange = []
                direct = "anzahl"
                window1.close()




            if event == "Neuen Datensatz anlegen" or event == sg.WIN_CLOSED:
                #window1['test'].update("P3_30 test.")
                countBuch = 0



                try:
                    open_db()
                    ID = values['ID']
                    ids = [books[0] for books in cursor.execute("SELECT ID FROM books")]
                    if ID and ID not in ids:
                        window1['outIDBuch'].update("ID ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outIDBuch'].update("ID darf nicht leer sein und darf nicht vorhanden sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')
                except:
                    window1['outIDBuch'].update("ID darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')



                try:
                    language = values['language']
                    if language:
                        window1['outLanguageBuch'].update("language ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outLanguageBuch'].update("language darf nicht leer sein. Bitte mit dropdown auswälen.", text_color='Red')
                except:
                    window1['outLanguageBuch'].update("language darf nicht leer sein. Bitte mit dropdown auswälen.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    warehouse = values['warehouse']
                    if warehouse:
                        window1['outwarehouseBuch'].update("warehouse ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outwarehouseBuch'].update("warehouse darf nicht leer sein. Bitte mit dropdown auswälen.", text_color='Red')
                except:
                    window1['outwarehouseBuch'].update("warehouse darf nicht leer sein. Bitte mit dropdown auswälen.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')




                try:
                    open_db()
                    name_of_item = values['name_of_item']
                    name_of_items = [name_of_items[0] for name_of_items in cursor.execute("SELECT name_of_item FROM books")]
                    if name_of_item and name_of_item not in name_of_items:
                        window1['outIDBuch'].update("Name ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outNameOfItemBuch'].update("Name darf nicht leer sein und darf nicht vorhanden sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')
                except:
                    window1['outNameOfItemBuch'].update("Name darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    type = values['type']
                    if type:
                        window1['outtypeBuch'].update("type ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outtypeBuch'].update("type darf nicht leer sein.", text_color='Red')
                except:
                    window1['outtypeBuch'].update("type darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')
                try:
                    BBT_price = values['BBT_price']
                    BBT_price = float(BBT_price)
                    window1['outBBT_priceBuch'].update("BBT_price ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outBBT_priceBuch'].update("BBT_price muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    SP = values['SP']
                    SP = float(SP)
                    window1['outSPBuch'].update("SP ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outSPBuch'].update("SP muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    P1 = values['P1']
                    P1 = float(P1)
                    window1['outP1Buch'].update("P1 ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outP1Buch'].update("P1 muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    P2 = values['P2']
                    P2 = float(P2)
                    window1['outP2Buch'].update("P2 ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outP2Buch'].update("P2 muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    P3_40 = values['P3_40']
                    P3_40 = float(P3_40)
                    window1['outP3_40Buch'].update("P3_40 ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outP3_40Buch'].update("P3_40 muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    P3_30 = values['P3_30']
                    P3_30 = float(P3_30)
                    window1['outP3_30Buch'].update("P3_30 ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outP3_30Buch'].update("P3_30 muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    End_P = values['End_P']
                    End_P = float(End_P)
                    window1['outEnd_PBuch'].update("End_P ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outEnd_PBuch'].update("End_P muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    Bestand = values['Bestand']
                    Bestand = int(Bestand)
                    window1['outBestandBuch'].update("Bestand ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outBestandBuch'].update("Bestand muss eine Zahl sein z.B.: 3", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')



                try:
                    Inventory_Name = values['Inventory_Name']
                    if Inventory_Name:
                        window1['outInventory_Name'].update("Inventory Name ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outInventory_Name'].update("Inventory Name darf nicht leer sein.", text_color='Red')
                except:
                    window1['outInventory_Name'].update("Inventory Name darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    La_Preis = values['La_Preis']
                    if La_Preis:
                        window1['outLa_Preis'].update("Lager Preisliste ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outLa_Preis'].update("Lager Preisliste darf nicht leer sein.", text_color='Red')
                except:
                    window1['outLa_Preis'].update("Lager Preisliste darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    Preis_Liste = values['Preis_Liste']
                    if Preis_Liste:
                        window1['outName_Preis'].update("Name Preisliste ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outName_Preis'].update("Name Preisliste darf nicht leer sein.", text_color='Red')
                except:
                    window1['outName_Preis'].update("Name Preisliste darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                if countBuch == 16:


                   

                    

                    total_BBT_price_quantity = BBT_price * Bestand



                    #try:


                    




                    open_db()



                    quantitiyB = cursor.execute("SELECT quantity FROM books WHERE ID = ?", (ID,)).fetchall()

                    quantitiyB = str(quantitiyB)
                    for char in ['(', ')', ',', '\'', ']', '[']:
                        if char in quantitiyB:

                            quantitiyB = quantitiyB.replace(char, '')






                    today1 = date.today()        


                    try:

                        Entnahme_Zuführung = int(quantitiyB) - int(Bestand)


                    except:

                        Entnahme_Zuführung = Bestand

                   
                    try:

                        cursor.execute("INSERT INTO bücherbewegung (book, date, amount, name_of_item, amount_old, Entnahme_Zuführung, Vorgangsnummer) VALUES(?,?,?,?,?,?,?)", (ID, today1, Bestand, name_of_item, quantitiyB, Entnahme_Zuführung, "Bestand_Neu"))                        

                        connection.commit()





                        cursor.execute("INSERT INTO books (ID, name_of_item, type, language, warehouse, quantity, BBT_price, SP, P1, P2, P3_40, P3_30, End_P, name, name_of_article, warehouses, total_BBT_price_quantity) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)", (ID, name_of_item, type, language, warehouse, Bestand, BBT_price, SP, P1, P2, P3_40, P3_30, End_P, Inventory_Name, Preis_Liste, La_Preis, total_BBT_price_quantity))

                        connection.commit()
                        connection.close()


                        window1['outBuch'].update('Der Datensatz mit der ID "' + ID + '" wurde in der Datenbank "' + dbName + '" gespeichert.', text_color='Green')


                        clear_Buch()

                    except:

                        window1['outIDBuch'].update("Diese ID existiert schon.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')







            if event == 'Clear' or event == sg.WIN_CLOSED:


                window1['outBuch'].update('')



                clear_Buch()












            if event == 'Bestehenden Datensatz über ID ändern' or event == sg.WIN_CLOSED:



                countBuch = 0

                try:
                    open_db()
                    ID = values['ID']
                    ids = [books[0] for books in cursor.execute("SELECT ID FROM books")]
                    if ID in ids:
                        window1['outIDBuch'].update("ID ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outIDBuch'].update("ID muss eine vorhandene ID sein.", text_color='Red')
                        window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')
                except:
                    window1['outIDBuch'].update("ID darf nicht leer sein und sie muss vorhanden sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')



                try:
                    language = values['language']
                    if language:
                        window1['outLanguageBuch'].update("language ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outLanguageBuch'].update("language darf nicht leer sein. Bitte mit dropdown auswälen.", text_color='Red')
                except:
                    window1['outLanguageBuch'].update("language darf nicht leer sein. Bitte mit dropdown auswälen.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    warehouse = values['warehouse']
                    if warehouse:
                        window1['outwarehouseBuch'].update("warehouse ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outwarehouseBuch'].update("warehouse darf nicht leer sein. Bitte mit dropdown auswälen.", text_color='Red')
                except:
                    window1['outwarehouseBuch'].update("warehouse darf nicht leer sein. Bitte mit dropdown auswälen.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    name_of_item = values['name_of_item']
                    if name_of_item:
                        window1['outNameOfItemBuch'].update("name_of_item ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outNameOfItemBuch'].update("name_of_item darf nicht leer sein.", text_color='Red')
                except:
                    window1['outNameOfItemBuch'].update("name_of_item darf nicht leer sein., text_color='Red'")
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    type = values['type']
                    if type:
                        window1['outtypeBuch'].update("type ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outtypeBuch'].update("type darf nicht leer sein.", text_color='Red')
                except:
                    window1['outtypeBuch'].update("type darf nicht leer sein.")
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    BBT_price = values['BBT_price']
                    BBT_price = float(BBT_price)
                    window1['outBBT_priceBuch'].update("BBT_price ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outBBT_priceBuch'].update("BBT_price muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    SP = values['SP']
                    SP = float(SP)
                    window1['outSPBuch'].update("SP ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outSPBuch'].update("SP muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    P1 = values['P1']
                    P1 = float(P1)
                    window1['outP1Buch'].update("P1 ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outP1Buch'].update("P1 muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    P2 = values['P2']
                    P2 = float(P2)
                    window1['outP2Buch'].update("P2 ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outP2Buch'].update("P2 muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')


                try:
                    P3_40 = values['P3_40']
                    P3_40 = float(P3_40)
                    window1['outP3_40Buch'].update("P3_40 ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outP3_40Buch'].update("P3_40 muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    P3_30 = values['P3_30']
                    P3_30 = float(P3_30)
                    window1['outP3_30Buch'].update("P3_30 ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outP3_30Buch'].update("P3_30 muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    End_P = values['End_P']
                    End_P = float(End_P)
                    window1['outEnd_PBuch'].update("End_P ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outEnd_PBuch'].update("End_P muss eine Zahl sein die mit Punkt getrennt wird z.B.: 3.88.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht geändert, bitte Fehlermeldungen beachten.", text_color='Red')

                try:
                    Bestand = values['Bestand']
                    Bestand = int(Bestand)
                    window1['outBestandBuch'].update("Bestand ist OK.", text_color='Green')
                    countBuch += 1
                except:
                    window1['outBestandBuch'].update("Bestand muss eine Zahl sein z.B.: 3", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    Inventory_Name = values['Inventory_Name']
                    if Inventory_Name:
                        window1['outInventory_Name'].update("Inventory Name ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outInventory_Name'].update("Inventory Name darf nicht leer sein.", text_color='Red')
                except:
                    window1['outInventory_Name'].update("Inventory Name darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')


                try:
                    La_Preis = values['La_Preis']
                    if La_Preis:
                        window1['outLa_Preis'].update("Lager Preisliste ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outLa_Preis'].update("Lager Preisliste darf nicht leer sein.", text_color='Red')
                except:
                    window1['outLa_Preis'].update("Lager Preisliste darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')



                try:
                    Preis_Liste = values['Preis_Liste']
                    if Preis_Liste:
                        window1['outName_Preis'].update("Name Preisliste ist OK.", text_color='Green')
                        countBuch += 1
                    else:
                        window1['outName_Preis'].update("Name Preisliste darf nicht leer sein.", text_color='Red')
                except:
                    window1['outName_Preis'].update("Name Preisliste darf nicht leer sein.", text_color='Red')
                    window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.", text_color='Red')



                if countBuch == 16:

                                        

                    total_BBT_price_quantity = BBT_price * Bestand



                    #try:


                    




                    open_db()



                    quantitiyB = cursor.execute("SELECT quantity FROM books WHERE ID = ?", (ID,)).fetchall()

                    quantitiyB = str(quantitiyB)
                    for char in ['(', ')', ',', '\'', ']', '[']:
                        if char in quantitiyB:

                            quantitiyB = quantitiyB.replace(char, '')






                    today1 = date.today()        


                    

                   


                    

                   
                    if int(Bestand) != int(quantitiyB):

                        Entnahme_Zuführung = int(quantitiyB) - int(Bestand)
            

                        cursor.execute("INSERT INTO bücherbewegung (book, date, amount, name_of_item, amount_old, Entnahme_Zuführung, Vorgangsnummer) VALUES(?,?,?,?,?,?,?)", (ID, today1, Bestand, name_of_item, quantitiyB, Entnahme_Zuführung, "Bestand_Änderung"))                        

                        connection.commit()




                    

                    total_BBT_price_quantity = BBT_price * Bestand
                    
                    
                    #print("total_BBT_price_quantity")
                    #print(total_BBT_price_quantity)


                    cursor.execute("UPDATE books SET name_of_item=?, type=?, language=?, warehouse=?, quantity=?, BBT_price=?, SP=?, P1=?, P2=?, P3_40=?, P3_30=?, End_P=?, total_BBT_price_quantity=?, name=?, name_of_article=?, warehouses=? WHERE id=? ", (name_of_item, type, language, warehouse, Bestand, BBT_price, SP, P1, P2, P3_40, P3_30, End_P, total_BBT_price_quantity, Inventory_Name, Preis_Liste, La_Preis, ID))


                    connection.commit()
                    connection.close()


                    window1['outBuch'].update('Der Datensatz mit der ID "' + ID + '" wurde in der Datenbank "' + dbName + '" gespeichert.', text_color='Green')


                    clear_Buch()



                        #window1['outIDBuch'].update("Diese ID existiert nicht.")
                        #window1['outBuch'].update("Datensatz wurde nicht übernommen, bitte Fehlermeldungen beachten.")








            if event == "Exit" or event == sg.WIN_CLOSED:
                db_backup()
                direct = "exit"
                connection.close()
                window1.close()




            if event == "Gauranga" or event == sg.WIN_CLOSED:
                direct = "menu"
                connection.close()
                lineToChange = []
                window1.close()

            if event == "Datensatz finden" or event == sg.WIN_CLOSED:
                direct = "change"
                window1.close()

        while direct == "change":
            open_db()
            # create dynamic drop down for table culums
            colNames = cursor.execute("SELECT * FROM books")
            colNames = [cn[0] for cn in colNames.description]


            layout1 = [
                [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text("Spalte auswählen:")],
                [sg.Combo(colNames, readonly=True, size=(79, 1), key='chooseColumn')],
                [sg.Button('Ok', size=(9,1)), sg.Text(key="outInfoChooseBuch1")],


            ]

            window1 = sg.Window("Buchdatensatz auswählen", layout1)

            while direct == "change":

                event, values = window1.read()

                if event == "Ok" or event == sg.WIN_CLOSED:

                    chooseColumn = ""

                    try:
                        chooseColumn = values['chooseColumn']

                    except:
                        pass

                    if chooseColumn != "":

                        direct = "content"
                        window1.close()

                    else:
                        window1['outInfoChooseBuch1'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')
                        direct = "change"
                        #window1.close()







                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"
                    connection.close()
                    lineToChange = []
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:

                    direct = "add"
                    window1.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()


            while direct == "content":

                content = [books[0] for books in cursor.execute("SELECT " + chooseColumn + " FROM books")]
                content = list(set(content))
                content = sorted(content)

                layout1 = [
                    [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                    [sg.Text("Inhalt auswählen:")],
                    [sg.Combo(content, readonly=True, size=(79, 1), key='chooseContent')],
                    [sg.Button('Ok', size=(9,1)), sg.Text(key="outInfoChooseBuch2")],
                ]

                window1 = sg.Window("Buchdatensatz auswählen", layout1)

                while direct == "content":

                    event, values = window1.read()

                    if event == "Ok" or event == sg.WIN_CLOSED:

                        chooseContent = ""

                        try:
                            chooseContent = values['chooseContent']

                        except:
                            pass

                        if chooseContent != "":
                            lines = [books[0] for books in cursor.execute("SELECT * FROM books Where " + chooseColumn + "=?",(chooseContent,))]
                            #lines = cursor.execute("SELECT * FROM books Where " + chooseColumn + "=?",(chooseContent,))
                            #lines = lines.fetchall()

                            if len(lines) == 1:

                                for i in lines:
                                    id = i

                                linID = cursor.execute("SELECT * FROM books Where ID=?",(id,))
                                linID = linID.fetchall()

                                result = []

                                for t in linID:
                                    for x in t:
                                        result.append(x)

                                lineToChange = result
                                window1.close()
                                direct = "add"
                                break

                            else:
                                direct = "line"
                                window1.close()


                        else:
                            direct = "content"
                            window1['outInfoChooseBuch2'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')


                    if event == "Gauranga" or event == sg.WIN_CLOSED:
                        direct = "menu"
                        lineToChange = []
                        connection.close()
                        window1.close()

                    if event == "Back" or event == sg.WIN_CLOSED:

                        direct = "add"
                        window1.close()

                    if event == "Exit" or event == sg.WIN_CLOSED:
                        db_backup()
                        direct = "exit"
                        connection.close()
                        window1.close()





            while direct == "line":

                lines = cursor.execute("SELECT * FROM books Where " + chooseColumn + "=?",(chooseContent,))
                lines = lines.fetchall()
                lines = list(set(lines))
                ##print(lines)

                layout1 = [
                    [sg.Text('', size=(90, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                    [sg.Text("Zeile auswählen:")],
                    [sg.Combo(lines, readonly=True, size=(140, 1), key='chooseLines')],
                    [sg.Button('Ok', size=(9,1)), sg.Text(key="outInfoChooseBuch3")],
                ]

                window1 = sg.Window("Buchdatensatz auswählen", layout1)

                while direct == "line":

                    event, values = window1.read()

                    if event == "Ok" or event == sg.WIN_CLOSED:

                        chooseLines = ""

                        try:
                            chooseLines = values['chooseLines']

                        except:
                            pass

                        if chooseLines != "":

                            chooseline = list(chooseLines)
                            lineToChange = chooseline

                            window1.close()
                            direct = "add"
                            break

                        else:
                            direct = "line"
                            window1['outInfoChooseBuch3'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')


                    if event == "Gauranga" or event == sg.WIN_CLOSED:
                        direct = "menu"
                        lineToChange = []
                        connection.close()
                        window1.close()

                    if event == "Back" or event == sg.WIN_CLOSED:

                        direct = "add"
                        window1.close()

                    if event == "Exit" or event == sg.WIN_CLOSED:
                        db_backup()
                        direct = "exit"
                        connection.close()
                        window1.close()


        while direct == "preiChan":




            layout1 = [
                [sg.Text('')],
                [sg.Text('Alle Preise ausser BBT von allen Büchern mit Prozentsatz:'), sg.Text('', size=(19, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Button('ERHÖHEN', size=(12, 1)), sg.Text('', size=(0, 1)), sg.InputText("Zahl eingeben", size=(12, 1), key="proPreiAnpa1"), sg.Text('%'), sg.Text('', size=(10, 1))],
                [sg.Button('REDUZIEREN', size=(12, 1)), sg.Text('', size=(0, 1)), sg.InputText("Zahl eingeben", size=(12, 1), key="proPreiAnpa0"), sg.Text('%')],
                [sg.Text(key="proInfoOpen")],
                [sg.Text(key="proInfo")]
            ]

            window1 = sg.Window("Alle Preise von allen Büchern ändern", layout1)

            while direct == "preiChan":

                event, values = window1.read()

                if event == "ERHÖHEN" or event == sg.WIN_CLOSED:

                    try:
                        proPreiAnpa1 = values['proPreiAnpa1']
                        proPreiAnpa1 = int(proPreiAnpa1)

                        if int(proPreiAnpa1):


                            content = [books[0] for books in cursor.execute("SELECT ID FROM books")]
                            #print(content)

                            # for book in content:
                            #     BBT_price = cursor.execute("SELECT BBT_price FROM books WHERE ID =?",(book,)).fetchall()


                            #     for t in BBT_price:
                            #         for x in t:
                            #             BBT_price = x

                            #     neuerPreis1 = BBT_price * proPreiAnpa1 / 100
                            #     neuerPreis1 = BBT_price + neuerPreis1
                            #     #print('alt:')
                            #     #print(BBT_price)
                            #     #print("neu")
                            #     neuerPreis1 = round(neuerPreis1, 2)
                            #     cursor.execute("UPDATE books SET BBT_price=? WHERE ID=?", (neuerPreis1, book))
                            #     connection.commit()


                            #     quantity = cursor.execute("SELECT quantity FROM books WHERE ID =?",(book,)).fetchall()


                            #     for t in quantity:
                            #         for x in t:
                            #             quantity = x


                            #     quantity = int(quantity)

                            #     BBT_price = float(BBT_price)

                            #     total_BBT_price_quantity = neuerPreis1 * quantity

                            #     #print('&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&')
                            #     #print(total_BBT_price_quantity)
                            #     cursor.execute("UPDATE books SET total_BBT_price_quantity=? WHERE id=? ", (total_BBT_price_quantity, book))

                            #     connection.commit()



                            for book in content:
                                SP = cursor.execute("SELECT SP FROM books WHERE ID =?",(book,)).fetchall()


                                for t in SP:
                                    for x in t:
                                        SP = x

                                neuerPreis1 = SP * proPreiAnpa1 / 100
                                neuerPreis1 = SP + neuerPreis1
                                #print('alt:')
                                #print(SP)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET SP=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()



                            for book in content:
                                P1 = cursor.execute("SELECT P1 FROM books WHERE ID =?",(book,)).fetchall()


                                for t in P1:
                                    for x in t:
                                        P1 = x

                                neuerPreis1 = P1 * proPreiAnpa1 / 100
                                neuerPreis1 = P1 + neuerPreis1
                                #print('alt:')
                                #print(P1)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET P1=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()


                            for book in content:
                                P2 = cursor.execute("SELECT P2 FROM books WHERE ID =?",(book,)).fetchall()


                                for t in P2:
                                    for x in t:
                                        P2 = x

                                neuerPreis1 = P2 * proPreiAnpa1 / 100
                                neuerPreis1 = P2 + neuerPreis1
                                #print('alt:')
                                #print(P2)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET P2=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()


                            for book in content:
                                P3_40 = cursor.execute("SELECT P3_40 FROM books WHERE ID =?",(book,)).fetchall()


                                for t in P3_40:
                                    for x in t:
                                        P3_40 = x

                                neuerPreis1 = P3_40 * proPreiAnpa1 / 100
                                neuerPreis1 = P3_40 + neuerPreis1
                                #print('alt:')
                                #print(P3_40)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET P3_40=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()



                            for book in content:
                                P3_30 = cursor.execute("SELECT P3_30 FROM books WHERE ID =?",(book,)).fetchall()


                                for t in P3_30:
                                    for x in t:
                                        P3_30 = x

                                neuerPreis1 = P3_30 * proPreiAnpa1 / 100
                                neuerPreis1 = P3_30 + neuerPreis1
                                #print('alt:')
                                #print(P3_30)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET P3_30=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()


                            for book in content:
                                End_P = cursor.execute("SELECT End_P FROM books WHERE ID =?",(book,)).fetchall()


                                for t in End_P:
                                    for x in t:
                                        End_P = x

                                neuerPreis1 = End_P * proPreiAnpa1 / 100
                                neuerPreis1 = End_P + neuerPreis1
                                #print('alt:')
                                #print(End_P)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET End_P=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()


                            direct = "preiChan"
                            window1['proInfo'].update('Alle Buchpreise wurden um ' + str(proPreiAnpa1) + '% erhöht.', text_color='Green')
                            window1['proPreiAnpa1'].update('')
                        else:

                            direct = "preiChan"
                            window1['proInfo'].update('Der Prozentsatz mus eine Zahl sein z.B.: 10 ohne Prozent Zeichen.', text_color='Red')

                    except:
                            direct = "preiChan"
                            window1['proInfo'].update('Der Prozentsatz mus eine Zahl sein z.B.: 10 ohne Prozent Zeichen.', text_color='Red')


                if event == "REDUZIEREN" or event == sg.WIN_CLOSED:

                    try:
                        proPreiAnpa1 = values['proPreiAnpa0']
                        proPreiAnpa1 = int(proPreiAnpa1)
                        if int(proPreiAnpa1):



                            content = [books[0] for books in cursor.execute("SELECT ID FROM books")]
                            #print(content)

                            # for book in content:
                            #     BBT_price = cursor.execute("SELECT BBT_price FROM books WHERE ID =?",(book,)).fetchall()


                            #     for t in BBT_price:
                            #         for x in t:
                            #             BBT_price = x

                            #     neuerPreis1 = BBT_price * proPreiAnpa1 / 100
                            #     neuerPreis1 = BBT_price - neuerPreis1
                            #     #print('alt:')
                            #     #print(BBT_price)
                            #     #print("neu")
                            #     neuerPreis1 = round(neuerPreis1, 2)
                            #     cursor.execute("UPDATE books SET BBT_price=? WHERE ID=?", (neuerPreis1, book))
                            #     connection.commit()


                            #     quantity = cursor.execute("SELECT quantity FROM books WHERE ID =?",(book,)).fetchall()


                            #     for t in quantity:
                            #         for x in t:
                            #             quantity = x


                            #     quantity = int(quantity)


                            #     total_BBT_price_quantity = neuerPreis1 * quantity

                            #     #print('&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&')
                            #     #print(total_BBT_price_quantity)
                            #     cursor.execute("UPDATE books SET total_BBT_price_quantity=? WHERE id=? ", (total_BBT_price_quantity, book))

                            #     connection.commit()




                            for book in content:
                                SP = cursor.execute("SELECT SP FROM books WHERE ID =?",(book,)).fetchall()


                                for t in SP:
                                    for x in t:
                                        SP = x

                                neuerPreis1 = SP * proPreiAnpa1 / 100
                                neuerPreis1 = SP - neuerPreis1
                                #print('alt:')
                                #print(SP)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET SP=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()



                            for book in content:
                                P1 = cursor.execute("SELECT P1 FROM books WHERE ID =?",(book,)).fetchall()


                                for t in P1:
                                    for x in t:
                                        P1 = x

                                neuerPreis1 = P1 * proPreiAnpa1 / 100
                                neuerPreis1 = P1 - neuerPreis1
                                #print('alt:')
                                #print(P1)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET P1=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()


                            for book in content:
                                P2 = cursor.execute("SELECT P2 FROM books WHERE ID =?",(book,)).fetchall()


                                for t in P2:
                                    for x in t:
                                        P2 = x

                                neuerPreis1 = P2 * proPreiAnpa1 / 100
                                neuerPreis1 = P2 - neuerPreis1
                                #print('alt:')
                                #print(P2)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET P2=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()


                            for book in content:
                                P3_40 = cursor.execute("SELECT P3_40 FROM books WHERE ID =?",(book,)).fetchall()


                                for t in P3_40:
                                    for x in t:
                                        P3_40 = x

                                neuerPreis1 = P3_40 * proPreiAnpa1 / 100
                                neuerPreis1 = P3_40 - neuerPreis1
                                #print('alt:')
                                #print(P3_40)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET P3_40=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()



                            for book in content:
                                P3_30 = cursor.execute("SELECT P3_30 FROM books WHERE ID =?",(book,)).fetchall()


                                for t in P3_30:
                                    for x in t:
                                        P3_30 = x

                                neuerPreis1 = P3_30 * proPreiAnpa1 / 100
                                neuerPreis1 = P3_30 - neuerPreis1
                                #print('alt:')
                                #print(P3_30)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET P3_30=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()


                            for book in content:
                                End_P = cursor.execute("SELECT End_P FROM books WHERE ID =?",(book,)).fetchall()


                                for t in End_P:
                                    for x in t:
                                        End_P = x

                                neuerPreis1 = End_P * proPreiAnpa1 / 100
                                neuerPreis1 = End_P - neuerPreis1
                                #print('alt:')
                                #print(End_P)
                                #print("neu")
                                neuerPreis1 = round(neuerPreis1, 3)
                                cursor.execute("UPDATE books SET End_P=? WHERE ID=?", (neuerPreis1, book))
                                connection.commit()






                            direct = "preiChan"
                            window1['proInfo'].update('Alle Buchpreise wurden um ' + str(proPreiAnpa1) + '% gesenkt.', text_color='Green')
                            window1['proPreiAnpa0'].update('')
                        else:

                            direct = "preiChan"
                            window1['proInfo'].update('Der Prozentsatz mus eine Zahl sein z.B.: 10 ohne Prozent Zeichen.', text_color='Red')

                    except:
                        direct = "preiChan"
                        window1['proInfo'].update('Der Prozentsatz mus eine Zahl sein z.B.: 10 ohne Prozent Zeichen.', text_color='Red')







                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"

                    connection.close()
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:

                    direct = "add"
                    connection.close()
                    window1.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()






        while direct == "anzahl":

            count = 0

            open_db()


            inAnzahlBücher = None
            date.today()
            dtNew = date.today()
            dtNew = dtNew.strftime('%Y-%m-%d')

            #if importAnz == "yes":




            books = [books[0] for books in cursor.execute("SELECT ID FROM books")]
            books = list(set(books))
            anzahlB = len(books)
            #print("xxxxxx")
            #print("#######")
            #print(books)

            if anzahlB == 0:
                direct = "add"
                clear_Buch()
                break

            boac = 1
            inAnzahlBücher = None


            layout1 = [

                [sg.Text('')],
                [sg.Button('<<  zurück', size=(11,1)), sg.Button('vorwärts  >>', size=(11,1)), sg.Text(key="boac"), sg.Text("/"), sg.Text(key="anzahlB"), sg.Button('EXCEL Importbestand speichern', size=(25,1)), sg.Button('Bestand neu speichern', size=(25,1)), sg.Text(size=(21,1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                [sg.Text('')],
                [sg.Text(key="outBuech", font=('bold'))],
                [sg.Text('Systembestand (alt):', text_color='Blue'), sg.Text(key="systemBestandKey", text_color='Blue'), sg.Text('EXCEL Importbestand: ', text_color='Blue'), sg.Text(key="csvBestandKey", text_color='Blue'), sg.Text('Differenz: ', text_color='Blue'), sg.Text(key="csvDifferenzKey", text_color='Blue')],
                [sg.Text('')],
                [sg.Text("Bestand neu:"), sg.InputText(inAnzahlBücher, size=(10, 1), key="keyAnzahlBücher"), sg.CalendarButton('Datum ändern', target='-CAL-', size=(12, 1), pad=None, key='_CALENDAR_', format=('%Y-%m-%d')), sg.In(dtNew, key='-CAL-', readonly=True, size=(10, 1)),sg.Text((''), size=(5,1))],
                [sg.Text('')],
                [sg.Button('EXCEL Importbestand Vorlage download', size=(34,1)), sg.FileBrowse('1. EXCEL Anzahl Importdatei auswählen', size=(34,1), key="anzky"), sg.Button('2. EXCEL Importbestand anzeigen', size=(34,1)), sg.Button('3. EXCEL Importbestand nicht mehr anzeigen', size=(34,1))],
                [sg.Text(key="outAnBuZe", size=(144, 3))],
                [sg.Text(key="outAnBu")],
            ]



            open_db()


            #ktStd = [books[0] for books in cursor.execute("SELECT ID, name_of_item, warehouse, language FROM books WHERE ID = ?", (books[1],))]



            ktStd = cursor.execute("SELECT ID, name_of_item, warehouse, language FROM books WHERE ID = ?", (books[count],)).fetchall()

            ktStd = str(ktStd)

            ktStd = ktStd.replace(',', ' -')

            for char in ['(', ')', '\'', ']', '[']:
                if char in ktStd:

                    ktStd = ktStd.replace(char, '')

            quantitiyB = cursor.execute("SELECT quantity FROM books WHERE ID = ?", (books[count],)).fetchall()

            quantitiyB = str(quantitiyB)
            for char in ['(', ')', ',', '\'', ']', '[']:
                if char in quantitiyB:

                    quantitiyB = quantitiyB.replace(char, '')



            window1 = sg.Window("Lagerbestände ändern", layout1, finalize=True)
            window1['outBuech'].update(ktStd)
            window1['anzahlB'].update(anzahlB)
            window1['boac'].update(boac)
            window1['systemBestandKey'].update(quantitiyB)
            #window1['keyAnzahlBücher'].update(quantitiyB)

            if importAnz == "yes":
                #print(dict)
                getKeys = [*dict]
                importAnz = "no"


            while direct == "anzahl":

                if donwload == "yes":

                    random_number = random.randint(1, 10000)
                    date.today()
                    dtNw = date.today()
                    dtNw = dtNw.strftime("%d_%m_%Y")
                    koExc = "Anzahl_Bücher_" + nameLa + "_" + str(random_number) + ".xlsx"
                    title = "Anzahl_Bücher "
                    lagExc = nameLa
                    koExc = os.path.join(repDirPath, koExc)

                    window1['outAnBuZe'].update("Die Datei " + koExc + " wurde erstellt und kann dem entsprechenen Lager zur Verfügung gestellt werden. Das Lager muss in die Spalte Anzahl die Anzahl der jweiligen Bücher schreiben und darf ansonsten nichts verändern. Dann kann die Datei in das System importiert werden.", text_color='Green')




                                        # rows = cursor.execute("SELECT * FROM books WHERE type = ? AND name_of_item = ? AND warehouse = ?", ##(typeDddd, name_of_itemDddd, warehouseDddd,)).fetchall()

                    # cursor.execute("UPDATE kunden_saldo SET kunden_saldo=? WHERE warehouse=? ", (nameLa))
                    open_db()
                    #content = [distributer[0] for distributer in cursor.execute("SELECT Name FROM distributer")]
                    rows = cursor.execute("SELECT ID, name_of_item, type, language FROM books WHERE warehouse = ? ORDER BY language, name_of_item", (nameLa,)).fetchall()
                    colNames = cursor.execute("SELECT ID, name_of_item, type, language FROM books")
                    colNames = [cn[0] for cn in colNames.description]
                    colNames = tuple(colNames)
                    rows.insert(0, colNames)
                    tableBooks = rows

                    df = pd.DataFrame(tableBooks)




                    writer = pd.ExcelWriter(koExc, engine='xlsxwriter')
                    df.to_excel(writer, sheet_name='Anzahl Bücher', index=False)

                    #workbook  = writer.book
                    worksheet = writer.sheets['Anzahl Bücher']


                    #worksheet.write(4, 10, title)


                    writer.save()



                    wb = openpyxl.load_workbook(koExc)
                    sheet = wb['Anzahl Bücher']

                    sheet.delete_rows(1)
                    sheet.insert_rows(idx=0, amount=3)
                    sheet.cell(row=2, column=1).value = title
                    sheet.cell(row=3, column=1).value = lagExc
                    sheet.cell(row=4, column=5).value = "Anzahl:"
                    wb.save(koExc)





                    absolutePath = Path(koExc).resolve()
                    os.system(f'start excel.exe "{absolutePath}"')

                    connection.close()

                    donwload = "no"

                event, values = window1.read()





                if event == "vorwärts  >>" or event == sg.WIN_CLOSED:
                    #print(anzahlB)
                    window1['csvBestandKey'].update("")
                    window1['csvDifferenzKey'].update("")

                    end = anzahlB - 1

                    if count < end:
                        #print(count)
                        count += 1
                        #print(count)
                        boac = count + 1

                    if count == end:
                        #print(count)
                        count = -1
                        #print(count)






                    open_db()

                    ktStd = cursor.execute("SELECT ID, name_of_item, warehouse, language FROM books WHERE ID = ?", (books[count],)).fetchall()

                    ktStd = str(ktStd)

                    ktStd = ktStd.replace(',', ' -')

                    for char in ['(', ')', '\'', ']', '[']:
                        if char in ktStd:

                            ktStd = ktStd.replace(char, '')


                    quantitiyB = cursor.execute("SELECT quantity FROM books WHERE ID = ?", (books[count],)).fetchall()

                    quantitiyB = str(quantitiyB)
                    for char in ['(', ')', ',', '\'', ']', '[']:
                        if char in quantitiyB:

                            quantitiyB = quantitiyB.replace(char, '')


                    if importAnz == "no":

                        window1['csvBestandKey'].update("")
                        window1['csvDifferenzKey'].update("")
                        window1['systemBestandKey'].update(quantitiyB)
                        #window1 = sg.Window("Lagerbestände ändern", layout1, finalize=True)
                        window1['outBuech'].update(ktStd)
                        window1['anzahlB'].update(anzahlB)
                        window1['boac'].update(boac)
                        window1['keyAnzahlBücher'].update("")
                        window1['outAnBu'].update("")



                    if importAnz == "yes":
                        #print("4024 zzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzz")
                        #print(dict)
                        getKeys = [*dict]

                        #print(getKeys)


                        # value = dict.get('wieder')

                        for key in getKeys:

                            pass
                            #print(key)

                        # #print(value)


                        #print("ID wird gedruckt")
                        #print(books)
                        for idcsv in getKeys:
                            #print(idcsv)
                            #print(books[count])
                            if idcsv == books[count]:
                                #print('check yes')
                                value = dict.get(idcsv)

                                try:
                                    value = int(value)
                                except:
                                    pass


                                try:
                                    diffBuecher = int(quantitiyB) - int(value)
                                except:
                                    diffBuecher = "nan"

                                window1['csvBestandKey'].update(value)
                                window1['csvDifferenzKey'].update(diffBuecher)



                    window1['systemBestandKey'].update(quantitiyB)
                    #window1 = sg.Window("Lagerbestände ändern", layout1, finalize=True)
                    window1['outBuech'].update(ktStd)
                    window1['anzahlB'].update(anzahlB)
                    window1['boac'].update(boac)
                    window1['keyAnzahlBücher'].update("")
                    window1['outAnBu'].update("")

                if event == "<<  zurück" or event == sg.WIN_CLOSED:
                    window1['csvBestandKey'].update("")
                    window1['csvDifferenzKey'].update("")

                    #print("nach zurück")
                    #print(count)
                    end = anzahlB - 1

                    if count >= 1:
                        #print(count)
                        count -= 1
                        #print(count)
                        boac = count + 1

                    elif count == -1:
                        #print(count)
                        count = end - 1
                        #print(count)
                        boac = count + 1

                    elif count == 0:
                        count = end
                        #print(count)
                        boac = count + 1





                    open_db()

                    ktStd = cursor.execute("SELECT ID, name_of_item, warehouse, language FROM books WHERE ID = ?", (books[count],)).fetchall()

                    ktStd = str(ktStd)

                    ktStd = ktStd.replace(',', ' -')

                    for char in ['(', ')', '\'', ']', '[']:
                        if char in ktStd:

                            ktStd = ktStd.replace(char, '')


                    quantitiyB = cursor.execute("SELECT quantity FROM books WHERE ID = ?", (books[count],)).fetchall()

                    quantitiyB = str(quantitiyB)
                    for char in ['(', ')', ',', '\'', ']', '[']:
                        if char in quantitiyB:

                            quantitiyB = quantitiyB.replace(char, '')


                    if importAnz == "yes":
                        # #print(dict)
                        getKeys = [*dict]

                        # #print(getKeys)


                        # value = dict.get('wieder')

                        # for key in getKeys:
                        #     #print(key)

                        # #print(value)


                        #print("ID wird gedruckt")
                        #print(books)
                        for idcsv in getKeys:
                            #print(idcsv)
                            #print(books[count])
                            if idcsv == books[count]:
                                #print('check yes')
                                value = dict.get(idcsv)
                                try:
                                    value = int(value)
                                except:
                                    pass

                                try:
                                    diffBuecher = int(quantitiyB) - int(value)
                                except:
                                    diffBuecher = "nan"

                                window1['csvBestandKey'].update(value)
                                window1['csvDifferenzKey'].update(diffBuecher)


                    window1['systemBestandKey'].update(quantitiyB)
                    #window1 = sg.Window("Lagerbestände ändern", layout1, finalize=True)
                    window1['outBuech'].update(ktStd)
                    window1['anzahlB'].update(anzahlB)
                    window1['boac'].update(boac)
                    window1['keyAnzahlBücher'].update("")
                    window1['outAnBu'].update("")


                if event == "EXCEL Importbestand speichern" or event == sg.WIN_CLOSED:

                    try:

                        open_db()

                        BBT_price = cursor.execute("SELECT BBT_price FROM books WHERE ID = ?", (idcsv,)).fetchall()

                        BBT_price = str(BBT_price)

                        for char in ['(', ')', ',', '\'', ']', '[']:
                            if char in BBT_price:

                                BBT_price = BBT_price.replace(char, '')

                        BBT_price = float(BBT_price)

                        value = int(value)

                        total_BBT_price_quantity = BBT_price * value


                        quantitiyB = cursor.execute("SELECT quantity FROM books WHERE ID = ?", (books[count],)).fetchall()

                        quantitiyB = str(quantitiyB)
                        for char in ['(', ')', ',', '\'', ']', '[']:
                            if char in quantitiyB:

                                quantitiyB = quantitiyB.replace(char, '')



                        name_of_item = cursor.execute("SELECT name_of_item FROM books WHERE ID = ?", (books[count],)).fetchall()

                        name_of_item = str(name_of_item)
                        for char in ['(', ')', ',', '\'', ']', '[']:
                            if char in name_of_item:

                                name_of_item = name_of_item.replace(char, '')


                                

                        Entnahme_Zuführung = int(quantitiyB) - int(value)

                        bookDA = values['-CAL-']
                        cursor.execute("INSERT INTO bücherbewegung (book, date, amount, name_of_item, amount_old, Entnahme_Zuführung, Vorgangsnummer) VALUES(?,?,?,?,?,?,?)", (books[count], bookDA, value, name_of_item, quantitiyB, Entnahme_Zuführung, "EXCEL_Importbestand"))                        


                        cursor.execute("UPDATE books SET quantity=?, total_BBT_price_quantity=? WHERE id=? ", (value, total_BBT_price_quantity, idcsv))




                        connection.commit()
                        connection.close()

                        open_db()

                        BBT_price = cursor.execute("SELECT BBT_price FROM books WHERE ID = ?", (books[count],)).fetchall()

                        BBT_price = str(BBT_price)

                        for char in ['(', ')', ',', '\'', ']', '[']:
                            if char in BBT_price:

                                BBT_price = BBT_price.replace(char, '')

                        BBT_price = float(BBT_price)

                        total_BBT_price_quantity = BBT_price * value


                        cursor.execute("UPDATE books SET quantity=?, total_BBT_price_quantity=? WHERE id=? ", (value, total_BBT_price_quantity, books[count]))





                        connection.commit()
                        connection.close()


                        window1['outAnBu'].update("Die Bücheranzahl " + str(value) + " wurde in der sankirtan.db gespeichert.", text_color='Green')


                    except:

                        window1['outAnBu'].update("Anzahl Bücher muss eine Zahl sein z.B.: 108. Datensatz wurde nicht gespeichert.", text_color='Red')





                if event == "Bestand neu speichern" or event == sg.WIN_CLOSED:
                    try:
                        keyAnzahlBücher = values['keyAnzahlBücher']
                        keyAnzahlBücher = int(keyAnzahlBücher)



                        open_db()

                        BBT_price = cursor.execute("SELECT BBT_price FROM books WHERE ID = ?", (books[count],)).fetchall()

                        BBT_price = str(BBT_price)

                        for char in ['(', ')', ',', '\'', ']', '[']:
                            if char in BBT_price:

                                BBT_price = BBT_price.replace(char, '')

                        BBT_price = float(BBT_price)

                        total_BBT_price_quantity = BBT_price * keyAnzahlBücher




                        quantitiyB = cursor.execute("SELECT quantity FROM books WHERE ID = ?", (books[count],)).fetchall()

                        quantitiyB = str(quantitiyB)
                        for char in ['(', ')', ',', '\'', ']', '[']:
                            if char in quantitiyB:

                                quantitiyB = quantitiyB.replace(char, '')



                        name_of_item = cursor.execute("SELECT name_of_item FROM books WHERE ID = ?", (books[count],)).fetchall()

                        name_of_item = str(name_of_item)
                        for char in ['(', ')', ',', '\'', ']', '[']:
                            if char in name_of_item:

                                name_of_item = name_of_item.replace(char, '')


                        bookDA = values['-CAL-']                        

                        Entnahme_Zuführung = int(quantitiyB) - int(keyAnzahlBücher)


                        cursor.execute("INSERT INTO bücherbewegung (book, date, amount, name_of_item, amount_old, Entnahme_Zuführung, Vorgangsnummer) VALUES(?,?,?,?,?,?,?)", (books[count], bookDA, keyAnzahlBücher, name_of_item, quantitiyB, Entnahme_Zuführung, "Bestand_neu_speichern"))                        



                        cursor.execute("UPDATE books SET quantity=?, total_BBT_price_quantity=? WHERE id=? ", (keyAnzahlBücher, total_BBT_price_quantity, books[count]))

                        connection.commit()
                        connection.close()


                        #window1 = sg.Window("Lagerbestände ändern", layout1, finalize=True)
                        window1['outAnBu'].update("Die Bücheranzahl " + str(keyAnzahlBücher) + " wurde in der sankirtan.db gespeichert.", text_color='Green')
                        window1['keyAnzahlBücher'].update("")
                        inAnzahlBücher = 1
                    #                             window1['bezahlt'].update('')
                    # window1['-CAL-'].update('')
                    # window1['outRech'].update("")
                    # window1['outKun'].update("")


                        countBuch += 1
                    except:
                        window1['outAnBu'].update("Anzahl Bücher muss eine Zahl sein z.B.: 108. Datensatz wurde nicht gespeichert.", text_color='Red')


                if event == "Gauranga" or event == sg.WIN_CLOSED:
                    direct = "menu"
                    lineToChange = []
                    connection.close()
                    window1.close()

                if event == "Back" or event == sg.WIN_CLOSED:

                    direct = "add"
                    window1.close()

                if event == "Exit" or event == sg.WIN_CLOSED:
                    db_backup()
                    direct = "exit"
                    connection.close()
                    window1.close()


                if event == "3. EXCEL Importbestand nicht mehr anzeigen" or event == sg.WIN_CLOSED:
                    #print("haha")
                    importAnz = "no"

                    window1['csvBestandKey'].update("")
                    window1['csvDifferenzKey'].update("")
                    open_db()
                    books = [books[0] for books in cursor.execute("SELECT ID FROM books")]
                    #print(books)
                    books = list(set(books))
                    #print(books)

                    anzahlB = len(books)
                    #print(anzahlB)

                    window1['anzahlB'].update(anzahlB)
                    window1['boac'].update(0)
                    count = 0

                    #window1 = sg.Window("Lagerbestände ändern", layout1, finalize=True)
                    window1['outBuech'].update("")
                    window1['keyAnzahlBücher'].update("")
                    window1['outAnBu'].update("Die normal Ansicht in der alle Bücher angezeigt werden ist aktiv und kann mit den zurück und vorwärts Tasten angezeigt werden.", text_color='Green')
                    window1['anzahlB'].update(anzahlB)
                    window1['boac'].update(1)
                    window1['outAnBuZe'].update("")
                    count = 0




                if event == "2. EXCEL Importbestand anzeigen" or event == sg.WIN_CLOSED:

                    try:

                        importAnz = "yes"

                        anzky = values['anzky']
                        #print(anzky)

                        lagerFrFile = openpyxl.load_workbook(anzky)

                        lagerFrFile = lagerFrFile.active

                        lagerFrFile = lagerFrFile['A3']

                        lagerFrFile = lagerFrFile.value

                        #print("warehouse")
                        #print(lagerFrFile)

                        dict = pd.read_excel(anzky, index_col=0, skiprows=4, squeeze=True, header=None, usecols=(0,4))
                        dict = dict.to_dict()
                        #print("dict")
                        #print(dict)

                        open_db()
                        books = [books[0] for books in cursor.execute("SELECT ID FROM books WHERE warehouse = ?", (lagerFrFile,))]
                        #print(books)
                        books = list(set(books))
                        #print(books)

                        anzahlImport = [*dict]
                        anzahlImport = len(anzahlImport)



                        anzahlB = len(books)
                        #print(anzahlB)

                        window1['anzahlB'].update(anzahlB)
                        window1['boac'].update(0)
                        count = 0

                        window1['csvBestandKey'].update("")
                        #window1 = sg.Window("Lagerbestände ändern", layout1, finalize=True)
                        window1['outBuech'].update("")
                        window1['keyAnzahlBücher'].update("")
                        window1['outAnBu'].update("")
                        window1['anzahlB'].update(anzahlB)
                        window1['boac'].update(1)
                        window1['outAnBuZe'].update("Für das Lager " + str(lagerFrFile) + " ist die EXCEL Importanzeige aktiv und kann mit den zurück und vorwärts Tasten angezeigt werden. Die Anzahl verschiedener Bücher für dieses Lager im Systembestand beträgt: " +  str(anzahlB) + ". Die Anzahl im EXCEL Importbestand beträgt: " + str(anzahlImport) + ". Diese EXCEL Importdatei wird verwendet: " + anzky + ". Wenn bei bei Excel Importbestand nan steht bedeutet dass die Anzahl in der Importdatei fehlt. Wird nichts angzeigt heißt das dass dieses Buch nicht in der Importdatei enthalten ist.", text_color='Magenta')
                        count = 0

                    except:
                         window1['outAnBu'].update('Es wurde keine EXCEL Importdatei ausgewählt. Bitte diese zuerst auswählen: 1. EXCEL Anzahl Importdatei auswählen.', text_color='Red')


                if event == "EXCEL Importbestand Vorlage download" or event == sg.WIN_CLOSED:

                    direct = "findLager"
                    donwload = "yes"
                    window1.close()

            while direct == "findLager":

                    open_db()
                    content = [books[0] for books in cursor.execute("SELECT warehouse FROM books")]
                    #print("go")
                    content = list(set(content))
                    #print(content)
                    content = sorted(content)
                    layout1 = [
                        [sg.Text('', size=(37, 1)), sg.Button('Back', size=(9,1)), sg.Button('Gauranga', size=(9,1)), sg.Button('Exit', size=(9,1))],
                        [sg.Text("Lager auswählen:")],
                        [sg.Combo(content, readonly=True, size=(79, 1), key='nameLKey')],
                        [sg.Button('Ok', size=(9,1)), sg.Text(key="outLagerCSV")],
                    ]

                    window1 = sg.Window("Lager finden", layout1)

                    while direct == "findLager":

                        event, values = window1.read()

                        if event == "Ok" or event == sg.WIN_CLOSED:

                            chooseContent = ""

                            try:
                                nameLa = values['nameLKey']
                                if nameLa != "":
                                    direct = "anzahl"
                                    window1.close()


                            except:
                                pass

                        else:
                            direct = "findLager"
                            window1['outLagerCSV'].update('In das Feld oberhalb klicken und auswählen.', text_color='Red')


                        if event == "Gauranga" or event == sg.WIN_CLOSED:
                            direct = "menu"
                            lineToChange = []
                            connection.close()
                            window1.close()

                        if event == "Back" or event == sg.WIN_CLOSED:

                            direct = "eingang"
                            window1.close()

                        if event == "Exit" or event == sg.WIN_CLOSED:
                            db_backup()
                            direct = "exit"
                            connection.close()
                            window1.close()

